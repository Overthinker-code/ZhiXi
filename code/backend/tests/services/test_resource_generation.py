import json
import os
from datetime import datetime, timedelta, timezone
from threading import Barrier
from uuid import uuid4

import pytest

from app.schemas.resource_generation import ResourceGenerationRequest
from app.services.resource_generation_service import ResourceGenerationService
from app.services.resource_generation_service import ChatModelFactory
from app.services.content_quality_service import content_quality_service
from app.core.config import settings
from pypdf import PdfReader
from io import BytesIO
from docx import Document as WordDocument
from docx.oxml.ns import qn


def test_math_topic_uses_math_domain_in_local_fallback():
    service = ResourceGenerationService()
    terms = service._topic_terms("通用学习", "二次函数顶点式参数")
    domain = service._domain_profile("通用学习", "二次函数顶点式参数", terms)

    assert domain["domain"] == "数学课程"
    assert "图像" in domain["scenario"]
    assert "SQL" not in domain["scenario"]


def test_acid_topic_uses_transaction_terms_in_local_fallback():
    terms = ResourceGenerationService._topic_terms(
        "数据库系统原理", "事务 ACID 与转账一致性"
    )

    assert terms[:4] == ["原子性", "一致性", "隔离性", "持久性"]
    assert "关系模型" not in terms


def test_resource_package_contains_closed_loop_artifacts(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))

    response = service.generate(
        ResourceGenerationRequest(
            subject="人工智能导论",
            topic="智能搜索",
            learning_goal="理解智能搜索的状态空间、评价函数和图谱关系",
            difficulty="standard",
            target_minutes=45,
            use_web_search=False,
        )
    )

    names = {artifact.file_name for artifact in response.artifacts}
    assert {
        "lecture.md",
        "practice.md",
        "mind-map.mmd",
        "quality-checklist.md",
    }.issubset(names)

    package_dir = tmp_path / response.package_id
    lecture = (package_dir / "lecture.md").read_text(encoding="utf-8")
    practice = (package_dir / "practice.md").read_text(encoding="utf-8")
    mind_map = (package_dir / "mind-map.mmd").read_text(encoding="utf-8")
    checklist = (package_dir / "quality-checklist.md").read_text(encoding="utf-8")

    assert "学习笔记对齐" in lecture
    assert "概念关系图绑定" in lecture
    assert "未绑定可核验课程资料" in lecture
    assert "标准答案框架" in practice
    assert "AI 追练" in practice
    assert "quality-checklist.md" in mind_map
    assert "质量门槛" in checklist
    assert "继续生成规则" in checklist
    assert all("待填写" not in text for text in [lecture, practice, checklist])


def test_evidence_gate_neutralizes_course_attribution_without_real_sources(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))

    response = service.generate(
        ResourceGenerationRequest(
            subject="通用学习",
            topic="事务 ACID 特性",
            resource_types=["lecture_markdown", "practice_markdown", "reading_list"],
        ),
        runtime_context={
            "citations": [
                {"type": "course", "id": str(uuid4()), "title": "数据库系统"},
                {"type": "learning_evidence", "id": str(uuid4()), "source_id": "quiz-1"},
                {"type": "knowledge_point", "id": "client-supplied-node"},
            ],
            "evidence_summary": "课程元数据和一次练习行为",
        },
    )

    package_dir = tmp_path / response.package_id
    combined = "\n".join(
        (package_dir / name).read_text(encoding="utf-8")
        for name in ("lecture.md", "practice.md", "reading-list.md")
    )
    assert response.local_model_profile["grounding_mode"] == "general_knowledge"
    assert response.local_model_profile["evidence_gate"] == {
        "passed": True,
        "course_claims_allowed": False,
        "verifiable_citation_count": 0,
        "decision": "deny_and_generalize",
        "policy": "course_claims_require_verifiable_source_id",
    }
    assert "未绑定可核验课程资料" in combined
    for unsupported in ("课程讲义", "课堂笔记", "课程图谱", "知识图谱", "课程证据"):
        assert unsupported not in combined
    assert all(
        item["passed"] and item["grounding_mode"] == "general_knowledge"
        for item in response.local_model_profile["quality_results"].values()
    )


def test_evidence_gate_allows_course_wording_with_verifiable_resource_source(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))

    source_id = str(uuid4())
    response = service.generate(
        ResourceGenerationRequest(
            subject="数据库系统",
            topic="关系模型",
            resource_types=["lecture_markdown"],
        ),
        runtime_context={
            "citations": [{"type": "course_resource", "id": source_id, "title": "关系模型讲义"}],
            "evidence_summary": f"已有课程资源：关系模型讲义（source_id={source_id}）",
        },
    )

    lecture = (tmp_path / response.package_id / "lecture.md").read_text(encoding="utf-8")
    assert response.local_model_profile["grounding_mode"] == "course_evidence"
    assert response.local_model_profile["evidence_gate"]["passed"] is True
    assert response.local_model_profile["evidence_gate"]["course_claims_allowed"] is True
    assert response.local_model_profile["evidence_gate"]["verifiable_citation_count"] == 1
    assert "课堂笔记对齐" in lecture
    assert "已有课程资源：关系模型讲义" in lecture
    assert "未绑定可核验课程资料" not in lecture


def test_acronym_completeness_is_repaired_and_recorded():
    service = ResourceGenerationService()
    ctx = service._build_context(
        ResourceGenerationRequest(subject="数据库系统", topic="事务 ACID 特性")
    )
    quality: dict[str, dict] = {}
    incomplete = "# 事务 ACID 特性\n\n## 说明\n" + ("ACID 是数据库事务的重要特性。" * 30)

    repaired = service._finalize_content(
        "lecture_markdown",
        incomplete,
        ctx,
        quality,
        source="deterministic_fallback",
    )

    for component in ("原子性", "一致性", "隔离性", "持久性"):
        assert component in repaired
    assert "completed_acronym_components" in quality["lecture_markdown"]["remediation_actions"]
    assert quality["lecture_markdown"]["passed"] is True


def test_required_field_review_reports_machine_readable_reason():
    result = content_quality_service.review(
        kind="practice_markdown",
        content="# 事务练习\n\n## 任务\n请解释事务。" * 30,
        topic="事务",
        has_course_evidence=False,
    )

    assert result.passed is False
    assert "missing_required_field:answer_or_rubric" in result.reasons
    assert result.checks["required_structure"]["answer_or_rubric"] is False


def test_unsupported_ai_course_claim_retries_then_degrades_with_audit_signal(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    unsupported = (
        "# 智能搜索\n\n## 学习目标\n"
        + "本结论来自课程讲义和课程图谱。智能搜索需要定义状态、动作与评价函数。" * 35
    )
    monkeypatch.setattr(
        service,
        "_generate_ai_contents",
        lambda *_: ({"lecture_markdown": unsupported}, ""),
    )
    monkeypatch.setattr(service, "_generate_single_ai_content", lambda *_: unsupported)

    response = service.generate(
        ResourceGenerationRequest(
            subject="人工智能导论",
            topic="智能搜索",
            resource_types=["lecture_markdown"],
        )
    )

    review = response.local_model_profile["quality_results"]["lecture_markdown"]
    lecture = (tmp_path / response.package_id / "lecture.md").read_text(encoding="utf-8")
    assert response.local_model_profile["fallback_artifacts"] == ["lecture_markdown"]
    assert review["retry_count"] == 2
    assert review["source"] == "deterministic_fallback"
    assert review["degraded"] is True
    assert review["passed"] is True
    assert "removed_unsupported_course_attribution" in review["remediation_actions"]
    assert "课程讲义" not in lecture
    assert "课程图谱" not in lecture


def test_resource_generation_response_hides_paths_and_keeps_context(
    tmp_path, monkeypatch
):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))
    course_id = uuid4()

    response = service.generate(
        ResourceGenerationRequest(
            course_id=course_id,
            resource_id="res-ai-search-01",
            node_id="resource-2",
            node_label="智能搜索",
            map_type="problem",
            source="resource-card",
            subject="人工智能导论",
            topic="智能搜索",
            learning_goal="生成能回流课程图谱的资料包",
            difficulty="standard",
            target_minutes=45,
            resource_types=["lecture_markdown", "mind_map", "quality_checklist"],
        )
    )

    assert response.course_id == course_id
    assert response.resource_id == "res-ai-search-01"
    assert response.node_id == "resource-2"
    assert response.node_label == "智能搜索"
    assert response.map_type == "problem"
    assert response.source == "resource-card"
    assert response.artifacts
    assert all(not hasattr(artifact, "file_path") for artifact in response.artifacts)

    manifest = tmp_path / response.package_id / "manifest.json"
    manifest_text = manifest.read_text(encoding="utf-8")
    assert str(course_id) in manifest_text
    assert "res-ai-search-01" in manifest_text
    assert "resource-2" in manifest_text


def test_legacy_manifest_digest_backfill_is_explicit_and_fail_closed(
    tmp_path, monkeypatch
):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))
    response = service.generate(
        ResourceGenerationRequest(
            subject="数据库系统原理",
            topic="事务完整性",
            resource_types=["lecture_markdown"],
        )
    )
    artifact = response.artifacts[0]
    manifest_path = tmp_path / response.package_id / "manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["artifacts"][0].pop("sha256")
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    with pytest.raises(ValueError, match="digest is missing"):
        service.resolve_verified_artifact_path(
            response.package_id, artifact.file_name
        )

    result = service.backfill_package_artifact_digests(response.package_id)

    assert result == {"verified": 1, "updated": 1}
    assert service.resolve_verified_artifact_path(
        response.package_id, artifact.file_name
    ).is_file()


def test_package_and_manifest_symlinks_are_rejected(tmp_path):
    service = ResourceGenerationService()
    service.output_root = tmp_path / "generated"
    service.output_root.mkdir()

    outside_package = tmp_path / "outside-package"
    outside_package.mkdir()
    (service.output_root / "linked-package").symlink_to(
        outside_package,
        target_is_directory=True,
    )

    with pytest.raises(ValueError, match="package symlinks"):
        service.package_directory("linked-package")

    package_dir = service.output_root / "real-package"
    package_dir.mkdir()
    outside_manifest = tmp_path / "outside-manifest.json"
    outside_manifest.write_text('{"package_id":"real-package"}', encoding="utf-8")
    (package_dir / "manifest.json").symlink_to(outside_manifest)

    with pytest.raises(ValueError, match="manifest symlinks"):
        service.read_package_manifest("real-package")


def test_recent_packages_filters_before_limit(tmp_path):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    target_course_id = str(uuid4())
    other_course_id = str(uuid4())

    for index in range(13):
        folder = tmp_path / f"other_{index}"
        folder.mkdir()
        (folder / "manifest.json").write_text(
            (
                '{"package_id":"other_%s","course_id":"%s","subject":"其他课程",'
                '"topic":"其他主题","generated_at":"2026-06-20T00:00:00",'
                '"artifacts":[]}'
            )
            % (index, other_course_id),
            encoding="utf-8",
        )
        (folder / "lecture.md").write_text("# other", encoding="utf-8")
        timestamp = (datetime.now(timezone.utc) + timedelta(minutes=index)).timestamp()
        os.utime(folder, (timestamp, timestamp))

    unscoped_folder = tmp_path / "unscoped_new"
    unscoped_folder.mkdir()
    (unscoped_folder / "manifest.json").write_text(
        (
            '{"package_id":"unscoped_new","course_id":"","subject":"通用主题",'
            '"topic":"不属于指定课程","generated_at":"2026-06-21T00:00:00",'
            '"artifacts":[]}'
        ),
        encoding="utf-8",
    )
    (unscoped_folder / "lecture.md").write_text("# unscoped", encoding="utf-8")
    newest_timestamp = (datetime.now(timezone.utc) + timedelta(days=1)).timestamp()
    os.utime(unscoped_folder, (newest_timestamp, newest_timestamp))

    target_folder = tmp_path / "target_old"
    target_folder.mkdir()
    (target_folder / "manifest.json").write_text(
        (
            '{"package_id":"target_old","course_id":"%s","resource_id":"res-1",'
            '"node_id":"node-1","node_label":"目标节点","map_type":"problem",'
            '"source":"resource-card","subject":"人工智能导论","topic":"智能搜索",'
            '"generated_at":"2026-06-19T00:00:00","artifacts":[]}'
        )
        % target_course_id,
        encoding="utf-8",
    )
    (target_folder / "lecture.md").write_text("# target", encoding="utf-8")
    old_timestamp = (datetime.now(timezone.utc) - timedelta(days=1)).timestamp()
    os.utime(target_folder, (old_timestamp, old_timestamp))

    packages = service.list_recent_packages(limit=12, course_id=target_course_id)

    assert [package["package_id"] for package in packages] == ["target_old"]
    assert packages[0]["resource_id"] == "res-1"
    assert packages[0]["node_id"] == "node-1"
    artifact = packages[0]["artifacts"][0]
    assert artifact["content_type"] == "text/markdown"
    assert artifact["preview"] == "# target"


def test_recent_packages_restore_manifest_artifact_metadata(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    monkeypatch.setattr(service, "_generate_ai_contents", lambda *_: ({}, "test"))

    response = service.generate(
        ResourceGenerationRequest(
            subject="数据库系统原理",
            topic="关系模型",
            resource_types=["lecture_markdown", "lecture_docx", "lecture_pdf"],
        )
    )

    package = service.list_recent_packages(limit=1)[0]
    artifacts = {item["file_name"]: item for item in package["artifacts"]}

    assert package["package_id"] == response.package_id
    assert artifacts["lecture.md"]["title"] == "关系模型 个性化讲义"
    assert artifacts["lecture.md"]["content_type"] == "text/markdown"
    assert "关系模型" in artifacts["lecture.md"]["preview"]
    assert artifacts["lecture.pdf"]["content_type"] == "application/pdf"
    assert artifacts["lecture.pdf"]["preview"] == ""
    assert artifacts["lecture.docx"]["content_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    word = WordDocument(tmp_path / response.package_id / "lecture.docx")
    assert "关系模型" in "\n".join(paragraph.text for paragraph in word.paragraphs)
    assert word.styles["Normal"]._element.rPr.rFonts.get(qn("w:eastAsia")) == os.getenv(
        "ZHIXI_DOCX_FONT_NAME", "Noto Sans CJK SC"
    )


def test_resource_agents_generate_independent_artifacts_in_parallel(tmp_path, monkeypatch):
    service = ResourceGenerationService()
    service.output_root = tmp_path
    requested = ["lecture_markdown", "practice_markdown", "mind_map"]
    barrier = Barrier(len(requested))

    def generate_one(_context, kind):
        barrier.wait(timeout=2)
        return f"{kind} content"

    monkeypatch.setattr(service, "_generate_single_ai_content", generate_one)

    contents, error = service._generate_ai_contents_parallel({}, requested)

    assert error == ""
    assert set(contents) == set(requested)


def test_chinese_pdf_preserves_text_layout_and_deduplicates_title():
    markdown = """# 事务 ACID 特性讲义

## 核心概念
中文标点：**原子性**、一致性、隔离性、持久性；使用 `COMMIT` 提交事务。

- 支持列表
- 支持分页与页码

| 属性 | 含义 |
| --- | --- |
| 原子性 | 全部成功或全部回滚 |
| 隔离性 | 并发事务相互隔离 |

```sql
SELECT * FROM account WHERE balance > 0;
```
"""
    pdf = ResourceGenerationService._render_pdf_bytes(
        markdown,
        title="事务 ACID 特性讲义",
    )
    reader = PdfReader(BytesIO(pdf))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert text.count("事务 ACID 特性讲义") == 1
    assert "原子性" in text
    assert "COMMIT" in text
    assert "**" not in text
    assert "`COMMIT`" not in text
    assert "全部成功或全部回滚" in text
    assert "SELECT * FROM account" in text
    assert "?" * 4 not in text


def test_chinese_pdf_inline_code_uses_cjk_capable_font():
    markdown = "来源：`ACID事务特性与转账失败分析 讲义 PDF`"

    pdf = ResourceGenerationService._render_pdf_bytes(
        markdown,
        title="来源字体验收",
    )
    reader = PdfReader(BytesIO(pdf))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert "ACID事务特性与转账失败分析 讲义 PDF" in text
    assert "■" not in text


def test_word_ordered_lists_preserve_section_local_markers():
    markdown = """# 编号验收

## 学习目标
1. 第一个目标
2. 第二个目标

## 内容依据
1. 第一条依据
2. 第二条依据
"""

    payload = ResourceGenerationService._render_docx_bytes(
        markdown,
        title="编号验收",
    )
    document = WordDocument(BytesIO(payload))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert "1. 第一个目标" in paragraphs
    assert "2. 第二个目标" in paragraphs
    assert "1. 第一条依据" in paragraphs
    assert "2. 第二条依据" in paragraphs


def test_acid_fallback_lecture_uses_domain_accurate_concept_cards():
    service = ResourceGenerationService()
    context = service._build_context(
        ResourceGenerationRequest(
            subject="数据库系统",
            topic="数据库事务的 ACID 特性与转账失败恢复",
        )
    )

    markdown = service._lecture_markdown(context)

    assert "事务中的操作要么全部完成，要么全部撤销" in markdown
    assert "事务前后都必须满足完整性约束和业务规则" in markdown
    assert "并发事务不应读取彼此未提交的中间状态" in markdown
    assert "事务一旦提交，结果在故障后仍可恢复" in markdown
    assert "用来确定问题对象和基本结构" not in markdown


def test_model_markdown_sanitizes_inline_html_without_breaking_sql_comparisons():
    source = (
        "Student(<u>StudentID</u>, Name)<br>"
        "SELECT * FROM account WHERE balance < 100 AND balance > 0;"
    )

    sanitized = ResourceGenerationService._sanitize_model_markdown(source)

    assert sanitized.startswith("Student(StudentID, Name)\nSELECT")
    assert "<u>" not in sanitized
    assert "</u>" not in sanitized
    assert "balance < 100" in sanitized
    assert "balance > 0" in sanitized


def test_chinese_pdf_long_table_can_split_across_pages():
    rows = "\n".join(f"| 第{i}行 | 中文内容{i} |" for i in range(1, 121))
    markdown = f"""# 长表格验收

| 序号 | 内容 |
| --- | --- |
{rows}
"""
    pdf = ResourceGenerationService._render_pdf_bytes(markdown, title="长表格验收")
    reader = PdfReader(BytesIO(pdf))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert len(reader.pages) >= 2
    assert "第120行" in text


def test_resource_specialist_uses_dedicated_timeout_and_falls_back(monkeypatch):
    captured = {}

    class TimeoutModel:
        def invoke(self, _messages):
            raise TimeoutError("resource generation timeout")

    def create(**kwargs):
        captured.update(kwargs)
        return TimeoutModel()

    monkeypatch.setattr(ChatModelFactory, "create", create)
    service = ResourceGenerationService()
    ctx = service._build_context(
        ResourceGenerationRequest(subject="数据库系统原理", topic="事务 ACID 特性")
    )
    contents, error = service._generate_ai_contents_parallel(ctx, ["lecture_markdown"])

    assert contents == {}
    assert error == "TimeoutError"
    assert captured["timeout_seconds"] == settings.RESOURCE_GENERATION_TIMEOUT_SECONDS == 45
