from __future__ import annotations

import hashlib
import json
import logging
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from threading import BoundedSemaphore
from time import perf_counter
from typing import Any
from uuid import UUID, uuid4

from sqlmodel import Session, select
from sqlalchemy.exc import IntegrityError

from app.core.config import settings
from app.core.db import engine
from app.models import (
    Course,
    GeneratedResourcePackage,
    LearningEvidence,
    ProfileUpdateEvent,
    Resource,
    ResourceGenerationRun,
    ResourceGenerationStep,
    Student,
    StudentTC,
    TC,
    User,
    UserMemoryProfile,
)
from app.schemas.resource_generation import (
    ResourceGenerationRequest,
    ResourceGenerationResponse,
    ResourceRunPublic,
    ResourceRunStepPublic,
)
from app.services.learning_path_service import learning_path_service
from app.services.learning_report_service import learning_report_service
from app.services.knowledge_graph_service import link_generated_resources
from app.services.content_quality_service import content_quality_service
from app.services.content_safety_service import (
    ContentSafetyBlockedError,
    SafetyReview,
    content_safety_service,
)
from app.services.resource_generation_service import (
    RESOURCE_AGENT_LABELS,
    resource_generation_service,
)


logger = logging.getLogger(__name__)
_RESOURCE_RUN_EXECUTOR = ThreadPoolExecutor(max_workers=2, thread_name_prefix="resource-run")
_RESOURCE_RUN_QUEUE_SLOTS = BoundedSemaphore(4)


class ResourcePackagePersistenceError(RuntimeError):
    def __init__(
        self,
        code: str,
        message: str,
        *,
        run_id: str | None = None,
        safety_review: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.run_id = run_id
        self.safety_review = safety_review


class ResourcePackageService:
    """Persistent state-machine for generation, review and course write-back."""

    TERMINAL_STATUSES = {"cancelled", "failed", "partial_success", "completed"}
    RUN_LEASE_SECONDS = 15 * 60

    @staticmethod
    def _resource_request_safety_text(request: ResourceGenerationRequest) -> str:
        """Serialize user-authored fields only; IDs and runtime state add no value."""

        return "\n".join(
            value
            for value in (
                request.subject,
                request.topic,
                request.learning_goal or "",
                request.node_label or "",
                request.source or "",
            )
            if value
        )

    @staticmethod
    def _extract_artifact_text(path: Any) -> str:
        """Extract reviewable text from generated artifacts without executing it."""

        suffix = str(path.suffix).lower()
        if suffix in {".md", ".mmd", ".txt"}:
            return path.read_text(encoding="utf-8", errors="replace")[:250_000]
        if suffix == ".docx":
            from docx import Document

            document = Document(str(path))
            chunks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)[:250_000]
        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)[:250_000]
        return ""

    def _review_resource_input(self, request: ResourceGenerationRequest) -> SafetyReview:
        try:
            return content_safety_service.ensure_safe(
                self._resource_request_safety_text(request),
                direction="input",
            )
        except ContentSafetyBlockedError as exc:
            raise ResourcePackagePersistenceError(
                "CONTENT_SAFETY_BLOCKED",
                str(exc),
                safety_review=exc.review.public_dict(),
            ) from exc

    def _review_resource_output(self, response: ResourceGenerationResponse) -> SafetyReview:
        contents: list[str] = []
        seen_digests: set[str] = set()
        for artifact in response.artifacts:
            path = resource_generation_service.resolve_artifact_path(
                response.package_id,
                artifact.file_name,
            )
            extracted = self._extract_artifact_text(path).strip()
            if not extracted:
                continue
            digest = hashlib.sha256(extracted.encode("utf-8", errors="replace")).hexdigest()
            if digest in seen_digests:
                continue
            seen_digests.add(digest)
            contents.append(extracted)
        package_text = "\n\n".join(contents)
        try:
            return content_safety_service.ensure_safe(
                package_text,
                direction="output",
            )
        except ContentSafetyBlockedError as exc:
            raise ResourcePackagePersistenceError(
                "CONTENT_SAFETY_BLOCKED",
                str(exc),
                safety_review=exc.review.public_dict(),
            ) from exc

    def _validate_artifact_set(
        self,
        response: ResourceGenerationResponse,
        *,
        run_id: str,
    ) -> dict[str, dict[str, Any]]:
        """Require a non-empty, unique and physically readable artifact set."""

        if not response.artifacts:
            raise ResourcePackagePersistenceError(
                "ARTIFACT_QUALITY_FAILED",
                "生成结果不包含可交付文件，资源包未入库",
                run_id=run_id,
            )
        checks: dict[str, dict[str, Any]] = {}
        for artifact in response.artifacts:
            if not artifact.file_name or artifact.file_name in checks:
                raise ResourcePackagePersistenceError(
                    "ARTIFACT_QUALITY_FAILED",
                    "生成结果包含空文件名或重复文件名，资源包未入库",
                    run_id=run_id,
                )
            try:
                path = resource_generation_service.resolve_artifact_path(
                    response.package_id,
                    artifact.file_name,
                )
                actual_size = path.stat().st_size if path.is_file() else 0
            except (FileNotFoundError, OSError, ValueError):
                actual_size = 0
            passed = artifact.file_size > 0 and actual_size > 0
            checks[artifact.file_name] = {
                "exists": actual_size > 0,
                "declared_size": artifact.file_size,
                "size": actual_size,
                "passed": passed,
            }
        if not all(item["passed"] for item in checks.values()):
            raise ResourcePackagePersistenceError(
                "ARTIFACT_QUALITY_FAILED",
                "质量审查发现缺失或空文件，资源包未入库",
                run_id=run_id,
            )
        return checks

    def _assert_attempt_owner(
        self,
        session: Session,
        run: ResourceGenerationRun,
        *,
        attempt_id: str | None,
        step_key: str,
    ) -> None:
        """Fence stale workers before they can commit another attempt's state."""

        attempt_id = attempt_id or getattr(run, "_execution_attempt_id", None)
        session.expire(run)
        session.refresh(run)
        if attempt_id and run.active_attempt_id != attempt_id:
            if run.status == "cancelled" or run.cancel_requested:
                raise ResourcePackagePersistenceError(
                    "RUN_CANCELLED",
                    "资源生成运行已取消",
                    run_id=run.id,
                )
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_LEASE_LOST",
                f"资源运行执行权已转移，旧执行器停止于 {step_key}",
                run_id=run.id,
            )
        if attempt_id:
            run.lease_expires_at = datetime.now(timezone.utc) + timedelta(
                seconds=self.RUN_LEASE_SECONDS
            )

    def _raise_if_cancelled(
        self,
        session: Session,
        run: ResourceGenerationRun,
        *,
        step_key: str,
        attempt_id: str | None = None,
    ) -> None:
        attempt_id = attempt_id or getattr(run, "_execution_attempt_id", None)
        self._assert_attempt_owner(
            session,
            run,
            attempt_id=attempt_id,
            step_key=step_key,
        )
        if not run.cancel_requested and run.status != "cancelled":
            return
        if run.package_id:
            self._finalize_post_persistence_cancel(session, run.id)
            raise ResourcePackagePersistenceError(
                "POST_PERSISTENCE_CANCELLED",
                "资源文件已保存，后续课程关联已停止，可稍后继续",
                run_id=run.id,
            )
        # A cancellation may arrive after generated package/resource rows were
        # flushed but before commit. Roll the whole unit of work back first;
        # otherwise committing the cancelled run would accidentally persist
        # those rows as well.
        session.rollback()
        persisted_run = session.get(ResourceGenerationRun, run.id)
        if not persisted_run:
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_NOT_FOUND", "资源运行不存在", run_id=run.id
            )
        persisted_run.cancel_requested = True
        persisted_run.status = "cancelled"
        persisted_run.current_step = "cancelled"
        persisted_run.finished_at = persisted_run.finished_at or datetime.now(timezone.utc)
        persisted_run.active_attempt_id = None
        persisted_run.lease_expires_at = None
        session.add(persisted_run)
        session.commit()
        raise ResourcePackagePersistenceError(
            "RUN_CANCELLED", "资源生成运行已取消", run_id=run.id
        )

    def _finalize_post_persistence_cancel(
        self, session: Session, run_id: str
    ) -> ResourceGenerationRun:
        """Stop optional write-backs without denying already committed files."""

        session.rollback()
        persisted_run = session.get(ResourceGenerationRun, run_id)
        if not persisted_run or not persisted_run.package_id:
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_NOT_FOUND",
                "未找到已保存的资源运行",
                run_id=run_id,
            )
        state = dict(persisted_run.shared_state or {})
        stage_status = dict(state.get("stage_status") or {})
        post_stages = ("linking_graph", "updating_path", "updating_profile")
        remaining = [
            stage for stage in post_stages if stage_status.get(stage) != "completed"
        ]
        for stage in remaining:
            stage_status[stage] = "skipped"
        state["stage_status"] = stage_status
        state["stage_failures"] = remaining
        degraded_reasons = list(state.get("degraded_reasons") or [])
        if "POST_PERSISTENCE_CANCELLED" not in degraded_reasons:
            degraded_reasons.append("POST_PERSISTENCE_CANCELLED")
        state["degraded_reasons"] = degraded_reasons
        persisted_run.shared_state = state
        persisted_run.cancel_requested = True
        persisted_run.status = "partial_success"
        persisted_run.current_step = "completed"
        persisted_run.error_code = "POST_PERSISTENCE_CANCELLED"
        persisted_run.error_message = "资源文件已保存，后续课程关联已停止，可稍后继续"
        persisted_run.finished_at = datetime.now(timezone.utc)
        persisted_run.active_attempt_id = None
        persisted_run.lease_expires_at = None
        package = session.get(
            GeneratedResourcePackage, persisted_run.package_id
        )
        if package:
            package.status = "partial_success"
            session.add(package)
        session.add(persisted_run)
        session.commit()
        session.refresh(persisted_run)
        return persisted_run

    @staticmethod
    def _digest(payload: Any) -> str:
        encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
        return hashlib.sha256(encoded).hexdigest()

    @staticmethod
    def _course_access_allowed(session: Session, *, user_id: UUID, course_id: UUID) -> bool:
        user = session.get(User, user_id)
        if user and user.is_superuser:
            return True
        student_id = session.exec(select(Student.id).where(Student.user_id == user_id)).first()
        if not student_id:
            return False
        enrolled = session.exec(
            select(StudentTC.id)
            .join(TC, StudentTC.tc_id == TC.id)
            .where(StudentTC.student_id == student_id, TC.course_id == course_id)
        ).first()
        return enrolled is not None

    def _validate_course_access(
        self,
        session: Session,
        *,
        user_id: UUID,
        course_id: UUID | None,
    ) -> None:
        if not course_id:
            return
        if not session.get(Course, course_id):
            raise ResourcePackagePersistenceError("COURSE_NOT_FOUND", "未找到指定课程，资源包未生成")
        if not self._course_access_allowed(session, user_id=user_id, course_id=course_id):
            raise ResourcePackagePersistenceError("COURSE_ACCESS_DENIED", "当前用户无权为该课程生成资源")

    def _start_run(
        self,
        session: Session,
        request: ResourceGenerationRequest,
        owner_id: UUID,
    ) -> ResourceGenerationRun:
        run = self.create_requested_run(session, request=request, owner_id=owner_id)
        run.active_attempt_id = uuid4().hex
        run.attempt_sequence += 1
        run.lease_expires_at = datetime.now(timezone.utc) + timedelta(
            seconds=self.RUN_LEASE_SECONDS
        )
        run.status = "running"
        run.current_step = "profiling"
        run.started_at = datetime.now(timezone.utc)
        session.add(run)
        session.commit()
        session.refresh(run)
        return run

    def create_requested_run(
        self,
        session: Session,
        *,
        request: ResourceGenerationRequest,
        owner_id: UUID,
        idempotency_key: str | None = None,
    ) -> ResourceGenerationRun:
        self._validate_course_access(session, user_id=owner_id, course_id=request.course_id)
        input_safety = self._review_resource_input(request)
        request_payload = request.model_dump(mode="json")
        request_digest = self._digest(request_payload)
        normalized_key = (idempotency_key or "").strip() or None
        if normalized_key and len(normalized_key) > 128:
            raise ResourcePackagePersistenceError(
                "INVALID_IDEMPOTENCY_KEY",
                "幂等键长度不能超过 128 个字符",
            )
        if normalized_key:
            existing = session.exec(
                select(ResourceGenerationRun).where(
                    ResourceGenerationRun.user_id == owner_id,
                    ResourceGenerationRun.idempotency_key == normalized_key,
                )
            ).first()
            if existing:
                if existing.request_digest != request_digest:
                    raise ResourcePackagePersistenceError(
                        "IDEMPOTENCY_CONFLICT",
                        "相同幂等键不能用于不同的资源生成请求",
                        run_id=existing.id,
                    )
                return existing
        active = session.exec(
            select(ResourceGenerationRun).where(
                ResourceGenerationRun.user_id == owner_id,
                ResourceGenerationRun.status.in_(["requested", "running"]),
            )
        ).first()
        if active:
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_ALREADY_ACTIVE",
                f"当前用户已有进行中的资源运行 {active.id}",
                run_id=active.id,
            )
        run = ResourceGenerationRun(
            id=f"rr_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}",
            user_id=owner_id,
            course_id=request.course_id,
            status="requested",
            current_step="requested",
            requested=request_payload,
            idempotency_key=normalized_key,
            request_digest=request_digest,
            shared_state={
                "handoff": "ProfileAgent",
                "stage_status": {"requested": "completed"},
                "max_artifact_retries": 2,
                "content_safety": {"input": input_safety.public_dict()},
            },
        )
        session.add(run)
        try:
            session.commit()
        except IntegrityError as exc:
            session.rollback()
            if normalized_key:
                existing = session.exec(
                    select(ResourceGenerationRun).where(
                        ResourceGenerationRun.user_id == owner_id,
                        ResourceGenerationRun.idempotency_key == normalized_key,
                    )
                ).first()
                if existing and existing.request_digest == request_digest:
                    return existing
            active = session.exec(
                select(ResourceGenerationRun).where(
                    ResourceGenerationRun.user_id == owner_id,
                    ResourceGenerationRun.status.in_(["requested", "running"]),
                )
            ).first()
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_ALREADY_ACTIVE",
                "当前用户已有进行中的资源运行",
                run_id=active.id if active else None,
            ) from exc
        session.refresh(run)
        return run

    def _prepare_attempt(self, run_id: str) -> str | None:
        with Session(engine) as session:
            run = session.exec(
                select(ResourceGenerationRun)
                .where(ResourceGenerationRun.id == run_id)
                .with_for_update()
            ).first()
            if not run or run.status != "requested" or run.active_attempt_id:
                return None
            attempt_id = uuid4().hex
            run.active_attempt_id = attempt_id
            run.attempt_sequence += 1
            run.lease_expires_at = datetime.now(timezone.utc) + timedelta(
                seconds=self.RUN_LEASE_SECONDS
            )
            session.add(run)
            session.commit()
            return attempt_id

    def enqueue_requested_run(self, run_id: str) -> bool:
        attempt_id = self._prepare_attempt(run_id)
        if not attempt_id:
            return False
        if not _RESOURCE_RUN_QUEUE_SLOTS.acquire(blocking=False):
            with Session(engine) as session:
                run = session.get(ResourceGenerationRun, run_id)
                if run and run.active_attempt_id == attempt_id:
                    run.status = "failed"
                    run.error_code = "RESOURCE_RUN_QUEUE_FULL"
                    run.error_message = "资源生成队列已满，请稍后重试"
                    run.finished_at = datetime.now(timezone.utc)
                    run.active_attempt_id = None
                    run.lease_expires_at = None
                    session.add(run)
                    session.commit()
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_QUEUE_FULL",
                "资源生成队列已满，请稍后重试",
                run_id=run_id,
            )
        try:
            future = _RESOURCE_RUN_EXECUTOR.submit(
                self.execute_requested_run,
                run_id,
                attempt_id,
            )
        except Exception:
            _RESOURCE_RUN_QUEUE_SLOTS.release()
            with Session(engine) as session:
                run = session.get(ResourceGenerationRun, run_id)
                if run and run.active_attempt_id == attempt_id:
                    run.active_attempt_id = None
                    run.lease_expires_at = None
                    run.status = "failed"
                    run.error_code = "BACKGROUND_SUBMIT_FAILED"
                    run.error_message = "后台资源运行提交失败"
                    run.finished_at = datetime.now(timezone.utc)
                    session.add(run)
                    session.commit()
            raise
        future.add_done_callback(lambda _future: _RESOURCE_RUN_QUEUE_SLOTS.release())
        return True

    def execute_requested_run(self, run_id: str, attempt_id: str | None = None) -> None:
        with Session(engine) as session:
            run = session.get(ResourceGenerationRun, run_id)
            if not run:
                return
            attempt_id = attempt_id or run.active_attempt_id or ""
            try:
                request = ResourceGenerationRequest.model_validate(run.requested)
                self.generate(
                    session,
                    request,
                    owner_id=run.user_id,
                    existing_run_id=run.id,
                    attempt_id=attempt_id,
                )
            except ResourcePackagePersistenceError:
                logger.warning("resource run %s ended with a controlled failure", run_id)
            except Exception:
                logger.exception("resource run %s crashed", run_id)
                session.rollback()
                persisted = session.get(ResourceGenerationRun, run_id)
                if persisted and persisted.active_attempt_id == attempt_id:
                    persisted.status = "failed"
                    persisted.error_code = "BACKGROUND_EXECUTION_FAILED"
                    persisted.error_message = "后台资源运行发生未处理异常"
                    persisted.finished_at = datetime.now(timezone.utc)
                    persisted.active_attempt_id = None
                    persisted.lease_expires_at = None
                    session.add(persisted)
                    session.commit()

    def _begin_step(
        self,
        session: Session,
        run: ResourceGenerationRun,
        *,
        step_key: str,
        agent_role: str,
        provider: str,
        model: str,
        input_payload: Any,
        input_summary: str,
    ) -> tuple[ResourceGenerationStep, float]:
        self._raise_if_cancelled(session, run, step_key=step_key)
        run.current_step = step_key
        state = dict(run.shared_state or {})
        state["handoff"] = agent_role
        stage_status = dict(state.get("stage_status") or {})
        stage_status[step_key] = "running"
        state["stage_status"] = stage_status
        run.shared_state = state
        step = ResourceGenerationStep(
            run_id=run.id,
            step_key=step_key,
            agent_role=agent_role,
            status="running",
            provider=provider,
            model=model,
            input_digest=self._digest(input_payload),
            input_summary=input_summary[:1000],
        )
        session.add(run)
        session.add(step)
        session.commit()
        session.refresh(step)
        return step, perf_counter()

    def _finish_step(
        self,
        session: Session,
        run: ResourceGenerationRun,
        step: ResourceGenerationStep,
        started: float,
        *,
        output_payload: Any,
        output_summary: str,
        status: str = "completed",
        citations: list[dict[str, Any]] | None = None,
        retry_count: int = 0,
        error_code: str | None = None,
        error_message: str | None = None,
        commit: bool = True,
    ) -> None:
        self._assert_attempt_owner(
            session,
            run,
            attempt_id=getattr(run, "_execution_attempt_id", None),
            step_key=step.step_key,
        )
        finished = datetime.now(timezone.utc)
        step.status = status
        step.output_digest = self._digest(output_payload)
        step.output_summary = output_summary[:1500]
        step.citations = citations or []
        step.retry_count = min(2, max(0, retry_count))
        step.error_code = error_code
        step.error_message = error_message[:2000] if error_message else None
        step.finished_at = finished
        step.duration_ms = max(0, round((perf_counter() - started) * 1000))
        state = dict(run.shared_state or {})
        stage_status = dict(state.get("stage_status") or {})
        stage_status[step.step_key] = status
        state["stage_status"] = stage_status
        state["last_output_digest"] = step.output_digest
        run.shared_state = state
        session.add(step)
        session.add(run)
        if commit:
            session.commit()
        else:
            session.flush([step, run])

    def _load_runtime_context(
        self,
        session: Session,
        run: ResourceGenerationRun,
        request: ResourceGenerationRequest,
    ) -> dict[str, Any]:
        profile_step, started = self._begin_step(
            session, run, step_key="profiling", agent_role="ProfileAgent", provider="database",
            model="user-memory-profile-query-v1", input_payload={"user_id": str(run.user_id)},
            input_summary="查询当前用户学习画像和原子证据统计",
        )
        profile_record = session.exec(select(UserMemoryProfile).where(UserMemoryProfile.user_id == run.user_id)).first()
        profile = dict(profile_record.memory_profile or {}) if profile_record else {}
        weak_points = [str(item) for item in (profile.get("weak_points") or [])[:4]]
        mastery = dict(profile.get("mastery_map") or {})
        profile_summary = ""
        if profile:
            review_priorities = [str(topic) for topic in list(mastery)[:4]]
            profile_summary = (
                f"当前目标：{str(profile.get('current_goal') or '未记录')[:120]}；"
                f"学习偏好：{str(profile.get('learning_style') or '未记录')[:80]}；"
                f"薄弱点：{'、'.join(weak_points) or '未记录'}；"
                f"复习优先方向：{'、'.join(review_priorities) or '未记录'}"
            )
        profile_citations = ([{"type": "user_memory_profile", "id": str(profile_record.id)}] if profile_record else [])
        self._finish_step(
            session, run, profile_step, started,
            output_payload={"profile_found": bool(profile), "weak_points": weak_points, "mastery_topic_count": len(mastery)},
            output_summary=("已读取真实画像并形成精简个性化约束" if profile else "no_profile：当前用户没有持久化学习画像"),
            citations=profile_citations,
        )

        retrieve_step, started = self._begin_step(
            session, run, step_key="retrieving", agent_role="EvidenceRetrieverAgent", provider="database",
            model="course-evidence-query-v1",
            input_payload={"course_id": str(request.course_id or ""), "node_id": request.node_id, "topic": request.topic},
            input_summary="查询课程、图谱节点 ID、已有课程资源和原子 evidence",
        )
        citations: list[dict[str, Any]] = []
        evidence_lines: list[str] = []
        if request.course_id:
            course = session.get(Course, request.course_id)
            if course:
                citations.append({"type": "course", "id": str(course.id), "title": course.name})
                evidence_lines.append(f"课程元数据：{course.name}（{course.identifier}）")
            existing_resources = session.exec(
                select(Resource).where(Resource.course_id == request.course_id).limit(8)
            ).all()
            for resource in existing_resources:
                citations.append({"type": "course_resource", "id": str(resource.id), "title": resource.title})
            if existing_resources:
                evidence_lines.append("已有课程资源：" + "、".join(item.title for item in existing_resources[:5]))
        canonical = learning_report_service.normalize_knowledge_point(request.node_id or request.topic)
        evidence_query = select(LearningEvidence).where(LearningEvidence.user_id == run.user_id)
        if request.course_id:
            evidence_query = evidence_query.where(LearningEvidence.course_id == request.course_id)
        if request.node_id:
            evidence_query = evidence_query.where(LearningEvidence.knowledge_point_id == request.node_id)
        else:
            evidence_query = evidence_query.where(LearningEvidence.knowledge_point == canonical)
        atomic_evidence = session.exec(evidence_query.order_by(LearningEvidence.observed_at.desc()).limit(8)).all()
        for item in atomic_evidence:
            citations.append({"type": "learning_evidence", "id": str(item.id), "source_id": item.source_id})
        if atomic_evidence:
            evidence_lines.append(f"学习证据：{len(atomic_evidence)} 条，来源 {len({item.source_type for item in atomic_evidence})} 类")
        if request.node_id:
            citations.append({"type": "knowledge_point", "id": request.node_id, "title": request.node_label or request.topic})
            evidence_lines.append(f"稳定知识点 ID：{request.node_id}")
        evidence_summary = "；".join(evidence_lines)
        verifiable_citations = content_quality_service.verifiable_evidence(citations)
        self._finish_step(
            session, run, retrieve_step, started,
            output_payload={
                "citation_count": len(citations),
                "verifiable_content_citation_count": len(verifiable_citations),
                "atomic_evidence_count": len(atomic_evidence),
                "canonical_key": canonical,
            },
            output_summary=(
                f"检索到 {len(verifiable_citations)} 个可支撑课程归因的真实来源"
                if verifiable_citations
                else "no_content_evidence：仅有课程元数据/学习行为或无引用，禁止声称来自课程资料"
            ),
            citations=citations,
        )

        plan_step, started = self._begin_step(
            session, run, step_key="planning", agent_role="ResourcePlannerAgent", provider="local",
            model="deterministic-state-machine", input_payload={"types": request.resource_types, "minutes": request.target_minutes, "citation_count": len(citations)},
            input_summary="基于真实画像和检索引用规划专用 Agent、交付物和质量门",
        )
        runtime_context = {
            "profile_summary": profile_summary,
            "evidence_summary": evidence_summary,
            "citations": citations,
            "knowledge_point_id": request.node_id,
            "canonical_key": canonical,
        }
        self._finish_step(
            session, run, plan_step, started,
            output_payload={"resource_types": request.resource_types, "runtime_context_digest": self._digest(runtime_context)},
            output_summary="已生成有界并行计划和最多 2 次单产物返工条件",
            citations=citations,
        )
        state = dict(run.shared_state or {})
        state["runtime_context"] = runtime_context
        run.shared_state = state
        session.add(run)
        session.commit()
        return runtime_context

    def generate(
        self,
        session: Session,
        request: ResourceGenerationRequest,
        *,
        owner_id: UUID,
        existing_run_id: str | None = None,
        attempt_id: str | None = None,
    ) -> ResourceGenerationResponse:
        self._validate_course_access(session, user_id=owner_id, course_id=request.course_id)
        if existing_run_id:
            run = session.get(ResourceGenerationRun, existing_run_id)
            if not run or run.user_id != owner_id:
                raise ResourcePackagePersistenceError("RESOURCE_RUN_NOT_FOUND", "资源运行不存在")
            if not attempt_id or run.active_attempt_id != attempt_id:
                raise ResourcePackagePersistenceError(
                    "RESOURCE_RUN_LEASE_LOST",
                    "资源运行执行权已失效",
                    run_id=run.id,
                )
            setattr(run, "_execution_attempt_id", attempt_id)
            if run.cancel_requested or run.status == "cancelled":
                run.status = "cancelled"
                run.finished_at = datetime.now(timezone.utc)
                session.add(run)
                session.commit()
                raise ResourcePackagePersistenceError("RUN_CANCELLED", "资源生成运行已取消", run_id=run.id)
            run.status = "running"
            run.current_step = "profiling"
            run.started_at = datetime.now(timezone.utc)
            run.error_code = None
            run.error_message = None
            session.add(run)
            session.commit()
            session.refresh(run)
        else:
            run = self._start_run(session, request, owner_id)
            attempt_id = run.active_attempt_id
            setattr(run, "_execution_attempt_id", attempt_id)
        response: ResourceGenerationResponse | None = None
        try:
            state = dict(run.shared_state or {})
            safety_state = dict(state.get("content_safety") or {})
            if not safety_state.get("input"):
                safety_state["input"] = self._review_resource_input(request).public_dict()
                state["content_safety"] = safety_state
                run.shared_state = state
                session.add(run)
                session.commit()
            runtime_context = self._load_runtime_context(session, run, request)
            generation_step, started = self._begin_step(
                session,
                run,
                step_key="generating",
                agent_role="SpecialistAgentTeam",
                provider=settings.CHAT_PROVIDER.lower(),
                model=resource_generation_service._active_chat_model_name(),
                input_payload={"topic": request.topic, "resource_types": request.resource_types},
                input_summary="按资源类型分派不同输入输出契约的专用 Agent",
            )
            response = resource_generation_service.generate(
                request,
                owner_id=owner_id,
                runtime_context=runtime_context,
            )
            file_checks = self._validate_artifact_set(response, run_id=run.id)
            self._raise_if_cancelled(session, run, step_key="generating")
            output_safety = self._review_resource_output(response)
            model_profile = dict(response.local_model_profile or {})
            model_profile["content_safety"] = {
                "input": safety_state.get("input"),
                "output": output_safety.public_dict(),
            }
            response = response.model_copy(update={"local_model_profile": model_profile})
            state = dict(run.shared_state or {})
            state["content_safety"] = model_profile["content_safety"]
            run.shared_state = state
            self._assert_attempt_owner(
                session,
                run,
                attempt_id=attempt_id,
                step_key="generating",
            )
            session.add(run)
            session.commit()
            # Provider calls may be blocking. Re-read the durable cancellation
            # flag before any generated file can be persisted or linked.
            self._raise_if_cancelled(session, run, step_key="generating")
            self._record_specialist_steps(session, run, request, response, runtime_context)
            retries = response.local_model_profile.get("artifact_retries") or {}
            self._finish_step(
                session, run, generation_step, started,
                output_payload={"package_id": response.package_id, "artifact_count": len(response.artifacts)},
                output_summary=f"生成 {len(response.artifacts)} 个文件；模型与回退来源逐产物记录",
                retry_count=max([int(value) for value in retries.values()] or [0]),
            )

            review_step, started = self._begin_step(
                session, run, step_key="reviewing", agent_role="CriticSafetyAgent",
                provider="local", model="deterministic-quality-gates",
                input_payload=response.local_model_profile.get("quality_results") or {},
                input_summary="逐产物检查结构、主题绑定、引用提示和安全边界",
            )
            quality = response.local_model_profile.get("quality_results") or {}
            file_checks = self._validate_artifact_set(response, run_id=run.id)
            fallback_count = len(response.local_model_profile.get("fallback_artifacts") or [])
            self._finish_step(
                session, run, review_step, started,
                output_payload={"content_quality": quality, "file_checks": file_checks},
                output_summary=f"质量审查完成；{fallback_count} 个产物使用确定性本地回退",
                retry_count=max([int(item.get("retry_count", 0)) for item in quality.values() if isinstance(item, dict)] or [0]),
            )

            persist_step, started = self._begin_step(
                session, run, step_key="persisting", agent_role="PackageAssemblerPersistenceAgent",
                provider="local", model="sqlmodel-filesystem-transaction",
                input_payload={"package_id": response.package_id, "artifacts": [item.file_name for item in response.artifacts]},
                input_summary="核验全部文件后原子写入 package 与 resource 行",
            )
            package = GeneratedResourcePackage(
                id=response.package_id, user_id=owner_id, course_id=request.course_id,
                subject=response.subject, topic=response.topic, source=response.source,
                resource_id=response.resource_id, node_id=response.node_id,
                node_label=response.node_label, map_type=response.map_type,
                status="completed", persistence_status="resources_persisted",
                model_profile=response.local_model_profile, agent_trace=response.agent_trace,
                quality_notes=response.quality_notes, generated_at=response.generated_at,
            )
            session.add(package)
            session.flush([package])
            resources: list[Resource] = []
            for artifact in response.artifacts:
                artifact_path = resource_generation_service.resolve_artifact_path(
                    response.package_id, artifact.file_name
                )
                resource = Resource(
                    title=artifact.title,
                    type=artifact.kind,
                    file_name=artifact.file_name,
                    file_path=f"generated_resources/{response.package_id}/{artifact.file_name}",
                    file_size=artifact_path.stat().st_size,
                    content_type=artifact.content_type,
                    subject=response.subject,
                    knowledge_point=response.node_label or response.topic,
                    difficulty=request.difficulty,
                    source="agent",
                    course_id=request.course_id,
                    package_id=response.package_id,
                    uploader_id=owner_id,
                )
                session.add(resource)
                resources.append(resource)
            session.flush()
            self._raise_if_cancelled(session, run, step_key="persisting")
            resource_ids = [item.id for item in resources]
            resource_generation_service.update_package_manifest(
                response.package_id,
                {"owner_id": str(owner_id), "run_id": run.id, "persistence_status": package.persistence_status,
                 "persisted_resource_ids": [str(item) for item in resource_ids]},
            )
            self._finish_step(
                session, run, persist_step, started,
                output_payload={"package_id": response.package_id, "resource_ids": resource_ids},
                output_summary=f"package、{len(resource_ids)} 条 resource 与文件已一致落库",
                commit=False,
            )
            # _finish_step fences the worker by expiring and refreshing ``run``.
            # Assign the package only after that fence so the durable run cannot
            # lose its core-persistence marker before commit.
            run.package_id = response.package_id
            session.add(run)
            session.flush([run])
            session.commit()

            stage_failures: list[str] = []
            if request.course_id:
                try:
                    self._link_graph(session, run, request, response, resources)
                except ResourcePackagePersistenceError:
                    raise
                except Exception as exc:
                    session.rollback(); stage_failures.append("linking_graph")
                    self._record_failed_stage(session, run, "linking_graph", "GraphLinkAgent", exc)
            else:
                self._record_skipped_stage(session, run, "linking_graph", "GraphLinkAgent", "全局资源包无 course_id")
            try:
                self._update_path(session, run, request, response)
            except ResourcePackagePersistenceError:
                raise
            except Exception as exc:
                session.rollback(); stage_failures.append("updating_path")
                self._record_failed_stage(session, run, "updating_path", "LearningPathUpdateAgent", exc)
            try:
                self._update_profile(session, run, request, response, owner_id)
            except ResourcePackagePersistenceError:
                raise
            except Exception as exc:
                session.rollback(); stage_failures.append("updating_profile")
                self._record_failed_stage(session, run, "updating_profile", "ProfileUpdateAgent", exc)

            self._raise_if_cancelled(session, run, step_key="finalizing")
            fallback_artifacts = list(
                response.local_model_profile.get("fallback_artifacts") or []
            )
            degraded_reasons = (["MODEL_FALLBACK"] if fallback_artifacts else [])
            run.status = (
                "partial_success"
                if stage_failures or degraded_reasons
                else "completed"
            )
            run.current_step = "completed"
            run.finished_at = datetime.now(timezone.utc)
            run.package_id = response.package_id
            run.active_attempt_id = None
            run.lease_expires_at = None
            state = dict(run.shared_state or {})
            state["stage_failures"] = stage_failures
            state["degraded_reasons"] = degraded_reasons
            state["fallback_artifacts"] = fallback_artifacts
            state["resource_ids"] = [str(item) for item in resource_ids]
            run.shared_state = state
            package.status = run.status
            session.add(package)
            session.add(run)
            session.commit()
            persisted_artifacts = [
                artifact.model_copy(
                    update={
                        "resource_id": resource.id,
                        "knowledge_point": resource.knowledge_point,
                        "difficulty": resource.difficulty,
                        "generated_at": response.generated_at,
                        "course_id": resource.course_id,
                    }
                )
                for artifact, resource in zip(response.artifacts, resources, strict=True)
            ]
            return response.model_copy(update={
                "run_id": run.id,
                "run_status": run.status,
                "stage_status": dict((run.shared_state or {}).get("stage_status") or {}),
                "persistence_status": package.persistence_status,
                "persisted_resource_ids": resource_ids,
                "artifacts": persisted_artifacts,
            })
        except ResourcePackagePersistenceError as exc:
            persisted_package = (
                session.get(GeneratedResourcePackage, response.package_id)
                if response
                else None
            )
            if response and not persisted_package:
                resource_generation_service.delete_package(response.package_id)
            persisted_run = session.get(ResourceGenerationRun, run.id)
            owns_attempt = bool(
                persisted_run
                and persisted_run.active_attempt_id == attempt_id
            )
            if (
                persisted_run
                and owns_attempt
                and persisted_run.status != "cancelled"
                and exc.code != "RESOURCE_RUN_LEASE_LOST"
            ):
                persisted_run.status = "failed"
                persisted_run.error_code = exc.code
                persisted_run.error_message = str(exc)[:2000]
                if exc.safety_review:
                    state = dict(persisted_run.shared_state or {})
                    safety_state = dict(state.get("content_safety") or {})
                    safety_state[str(exc.safety_review.get("direction") or "review")] = exc.safety_review
                    state["content_safety"] = safety_state
                    persisted_run.shared_state = state
                persisted_run.finished_at = datetime.now(timezone.utc)
                persisted_run.active_attempt_id = None
                persisted_run.lease_expires_at = None
                session.add(persisted_run)
                session.commit()
            raise
        except Exception as exc:
            session.rollback()
            persisted_package = (
                session.get(GeneratedResourcePackage, response.package_id)
                if response
                else None
            )
            if response and not persisted_package:
                resource_generation_service.delete_package(response.package_id)
            run = session.get(ResourceGenerationRun, run.id)
            if run and run.active_attempt_id == attempt_id:
                run.status = "failed"; run.error_code = "RESOURCE_RUN_FAILED"; run.error_message = str(exc)[:2000]
                run.finished_at = datetime.now(timezone.utc)
                run.active_attempt_id = None; run.lease_expires_at = None
                session.add(run); session.commit()
            raise ResourcePackagePersistenceError(
                "RESOURCE_RUN_FAILED", "资源运行失败；未完成的临时文件已清理", run_id=run.id if run else None
            ) from exc

    def _record_specialist_steps(
        self,
        session: Session,
        run: ResourceGenerationRun,
        request: ResourceGenerationRequest,
        response: ResourceGenerationResponse,
        runtime_context: dict[str, Any],
    ) -> None:
        ai_generated = set(response.local_model_profile.get("ai_generated_artifacts") or [])
        retries = response.local_model_profile.get("artifact_retries") or {}
        quality = response.local_model_profile.get("quality_results") or {}
        contracts = response.local_model_profile.get("agent_contracts") or {}
        artifacts_by_kind = {artifact.kind: artifact for artifact in response.artifacts}
        for kind, contract in contracts.items():
            if kind not in artifacts_by_kind:
                continue
            role = RESOURCE_AGENT_LABELS.get(kind, (f"{kind}Agent", kind))[0]
            provider = settings.CHAT_PROVIDER.lower() if kind in ai_generated else "local_fallback"
            model = resource_generation_service._active_chat_model_name() if kind in ai_generated else "deterministic-template-v1"
            step, started = self._begin_step(
                session,
                run,
                step_key=f"specialist:{kind}",
                agent_role=role,
                provider=provider,
                model=model,
                input_payload={"contract": contract, "runtime_context_digest": self._digest(runtime_context)},
                input_summary=f"执行 {kind} 专用输入输出契约与质量门",
            )
            artifact = artifacts_by_kind[kind]
            result = quality.get(kind) if isinstance(quality.get(kind), dict) else {}
            if kind not in ai_generated:
                result = {
                    "passed": True,
                    "mode": "validated_local_template",
                    "retry_count": int(retries.get(kind, 0)),
                }
            self._finish_step(
                session,
                run,
                step,
                started,
                output_payload={"file_name": artifact.file_name, "file_size": artifact.file_size, "quality": result},
                output_summary=f"生成 {artifact.file_name}；provider={provider}；质量门已记录",
                citations=list(runtime_context.get("citations") or []),
                retry_count=int(retries.get(kind, 0)),
                error_code=(
                    str(response.local_model_profile.get("fallback_reason") or "MODEL_FALLBACK")
                    if kind not in ai_generated else None
                ),
                error_message=(
                    "模型调用未产出通过质量门的内容，已使用确定性本地回退"
                    if kind not in ai_generated else None
                ),
            )

    def _link_graph(self, session: Session, run: ResourceGenerationRun, request: ResourceGenerationRequest, response: ResourceGenerationResponse, resources: list[Resource]) -> None:
        step, started = self._begin_step(session, run, step_key="linking_graph", agent_role="GraphLinkAgent", provider="local", model="normalized-knowledge-linker", input_payload={"topic": request.topic, "resources": [str(item.id) for item in resources]}, input_summary="将本次课程资源关联到规范化知识点")
        point = learning_report_service.normalize_knowledge_point(request.node_label or request.topic)
        node, link_count = link_generated_resources(
            session,
            run_id=run.id,
            package_id=response.package_id,
            course_id=request.course_id,
            knowledge_point=point,
            resources=resources,
        )
        self._finish_step(
            session,
            run,
            step,
            started,
            output_payload={
                "knowledge_point": point,
                "knowledge_node_id": str(node.id),
                "link_count": link_count,
            },
            output_summary=f"已建立 {link_count} 条资源-知识节点关系",
            commit=False,
        )
        session.commit()

    def _update_path(self, session: Session, run: ResourceGenerationRun, request: ResourceGenerationRequest, response: ResourceGenerationResponse) -> None:
        step, started = self._begin_step(session, run, step_key="updating_path", agent_role="LearningPathUpdateAgent", provider="local", model="event-sourced-path-update", input_payload={"topic": request.topic, "package_id": response.package_id}, input_summary="以本次资源包为可追踪来源更新学习路径")
        event = learning_path_service.update_from_resource_run(session, user_id=run.user_id, course_id=request.course_id, run_id=run.id, subject=request.subject, topic=request.topic, package_id=response.package_id)
        self._finish_step(session, run, step, started, output_payload={"event_id": str(event.id)}, output_summary=event.summary, commit=False)
        session.commit()

    def _update_profile(self, session: Session, run: ResourceGenerationRun, request: ResourceGenerationRequest, response: ResourceGenerationResponse, owner_id: UUID) -> None:
        step, started = self._begin_step(session, run, step_key="updating_profile", agent_role="ProfileUpdateAgent", provider="local", model="evidence-confidence-v1", input_payload={"topic": request.topic, "run_id": run.id}, input_summary="记录资源暴露证据；不把生成资源误当作掌握证明")
        before = learning_report_service.evidence_confidence(session, owner_id, course_id=request.course_id)
        evidence = learning_report_service.record_evidence(session, user_id=owner_id, course_id=request.course_id, run_id=run.id, knowledge_point=request.node_label or request.topic, knowledge_point_id=request.node_id, source_type="resource_run", source_id=response.package_id, event_type="resource_generated", weight=0.25, score=None, payload={"artifact_count": len(response.artifacts)})
        after = learning_report_service.evidence_confidence(session, owner_id, course_id=request.course_id)
        event = ProfileUpdateEvent(run_id=run.id, user_id=owner_id, course_id=request.course_id, status="completed", before_state=before, after_state=after, evidence_ids=[str(evidence.id)], summary="新增资源暴露证据；掌握度不因资源生成自动上调")
        session.add(event); session.flush([event])
        self._finish_step(session, run, step, started, output_payload={"event_id": str(event.id), "evidence_id": str(evidence.id)}, output_summary=event.summary, commit=False)
        session.commit()

    def _record_failed_stage(self, session: Session, run: ResourceGenerationRun, step_key: str, role: str, exc: Exception) -> None:
        step, started = self._begin_step(session, run, step_key=step_key, agent_role=role, provider="local", model="post-persistence", input_payload={}, input_summary="恢复阶段执行")
        self._finish_step(session, run, step, started, output_payload={}, output_summary="阶段失败，可按 run_id 重试", status="failed", error_code=exc.__class__.__name__, error_message=str(exc))

    def _record_skipped_stage(self, session: Session, run: ResourceGenerationRun, step_key: str, role: str, reason: str) -> None:
        step, started = self._begin_step(session, run, step_key=step_key, agent_role=role, provider="local", model="deterministic", input_payload={}, input_summary=reason)
        self._finish_step(session, run, step, started, output_payload={"reason": reason}, output_summary=reason, status="skipped")

    def get_run(self, session: Session, *, run_id: str, user_id: UUID, is_superuser: bool = False) -> ResourceRunPublic | None:
        # Polling may reuse a request/session identity map while the bounded
        # worker commits through another Session. Expire cached ORM state so
        # every poll observes the durable database status.
        session.expire_all()
        run = session.get(ResourceGenerationRun, run_id)
        if not run or (not is_superuser and run.user_id != user_id):
            return None
        steps = session.exec(select(ResourceGenerationStep).where(ResourceGenerationStep.run_id == run_id).order_by(ResourceGenerationStep.started_at)).all()
        return ResourceRunPublic(
            run_id=run.id, package_id=run.package_id,
            result_url=(f"/api/v1/resource-generation/packages/{run.package_id}" if run.package_id else None),
            course_id=run.course_id, status=run.status,
            current_step=run.current_step, cancel_requested=run.cancel_requested,
            attempt_sequence=run.attempt_sequence,
            lease_expires_at=run.lease_expires_at,
            requested=run.requested, shared_state=run.shared_state, error_code=run.error_code,
            error_message=run.error_message, created_at=run.created_at, started_at=run.started_at,
            finished_at=run.finished_at,
            steps=[ResourceRunStepPublic.model_validate(step, from_attributes=True) for step in steps],
        )

    def get_package(
        self,
        session: Session,
        *,
        package_id: str,
        user_id: UUID,
        is_superuser: bool = False,
    ) -> dict[str, Any] | None:
        package = session.get(GeneratedResourcePackage, package_id)
        if not package or (not is_superuser and package.user_id != user_id):
            return None
        try:
            manifest = resource_generation_service.read_package_manifest(package.id)
            artifacts = resource_generation_service.get_package_payload(package.id).get("artifacts", [])
            expected = {
                str(item.get("file_name")): int(item.get("file_size") or 0)
                for item in (manifest.get("artifacts") or [])
                if isinstance(item, dict) and item.get("file_name")
            }
            actual = {
                str(item.get("file_name")): int(item.get("file_size") or 0)
                for item in artifacts
                if isinstance(item, dict) and item.get("file_name")
            }
            integrity_ok = bool(expected) and set(actual) == set(expected)
            integrity_ok = integrity_ok and all(
                expected[name] > 0 and actual.get(name, 0) == expected[name]
                for name in expected
            )
            if integrity_ok:
                for item in manifest.get("artifacts") or []:
                    if not isinstance(item, dict):
                        integrity_ok = False
                        break
                    file_name = str(item.get("file_name") or "")
                    expected_digest = str(item.get("sha256") or "").strip().lower()
                    if not expected_digest:
                        integrity_ok = False
                        break
                    target = resource_generation_service.resolve_artifact_path(
                        package.id, file_name
                    )
                    if resource_generation_service._file_sha256(target) != expected_digest:
                        integrity_ok = False
                        break
            if not integrity_ok:
                raise ValueError("artifact manifest does not match files on disk")
        except (FileNotFoundError, ValueError, OSError, TypeError) as exc:
            raise ResourcePackagePersistenceError(
                "ARTIFACT_INTEGRITY_FAILED",
                "资源包文件不完整或已损坏，无法作为成功结果返回",
                run_id=(
                    session.exec(
                        select(ResourceGenerationRun.id).where(
                            ResourceGenerationRun.package_id == package.id
                        )
                    ).first()
                    or None
                ),
            ) from exc
        resources = session.exec(
            select(Resource).where(Resource.package_id == package.id)
        ).all()
        resources_by_file = {
            item.file_name: item for item in resources if item.file_name
        }
        hydrated_artifacts: list[dict[str, Any]] = []
        for artifact in artifacts:
            if not isinstance(artifact, dict):
                continue
            public_artifact = dict(artifact)
            resource = resources_by_file.get(str(artifact.get("file_name") or ""))
            if resource is not None:
                public_artifact.update(
                    {
                        "resource_id": str(resource.id),
                        "knowledge_point": resource.knowledge_point,
                        "difficulty": resource.difficulty,
                        "generated_at": package.generated_at.isoformat(),
                        "course_id": str(resource.course_id) if resource.course_id else None,
                    }
                )
            hydrated_artifacts.append(public_artifact)
        resource_ids = [item.id for item in resources]
        run = session.exec(select(ResourceGenerationRun).where(ResourceGenerationRun.package_id == package.id)).first()
        steps = (
            session.exec(
                select(ResourceGenerationStep)
                .where(ResourceGenerationStep.run_id == run.id)
                .order_by(ResourceGenerationStep.started_at)
            ).all()
            if run
            else []
        )
        return {
            "package_id": package.id,
            "run_id": run.id if run else "",
            "run_status": run.status if run else package.status,
            "stage_status": dict((run.shared_state or {}).get("stage_status") or {}) if run else {},
            "cancel_requested": run.cancel_requested if run else False,
            "steps": [
                ResourceRunStepPublic.model_validate(step, from_attributes=True).model_dump(mode="json")
                for step in steps
            ],
            "requested": dict(run.requested or {}) if run else {},
            "course_id": str(package.course_id) if package.course_id else "",
            "resource_id": package.resource_id or "",
            "node_id": package.node_id or "",
            "node_label": package.node_label or "",
            "map_type": package.map_type or "",
            "source": package.source or "",
            "subject": package.subject,
            "topic": package.topic,
            "generated_at": package.generated_at.isoformat(),
            "status": package.status,
            "persistence_status": package.persistence_status,
            "persisted_resource_ids": [str(item) for item in resource_ids],
            "local_model_profile": package.model_profile,
            "agent_trace": package.agent_trace,
            "quality_notes": package.quality_notes,
            "artifacts": hydrated_artifacts,
        }

    def request_cancel(self, session: Session, *, run_id: str, user_id: UUID) -> ResourceRunPublic | None:
        run = session.get(ResourceGenerationRun, run_id)
        if not run or run.user_id != user_id:
            return None
        if run.status == "requested":
            run.cancel_requested = True
            run.status = "cancelled"
            run.current_step = "cancelled"
            run.finished_at = datetime.now(timezone.utc)
            run.active_attempt_id = None
            run.lease_expires_at = None
            session.add(run)
            session.commit()
            session.refresh(run)
        elif run.status == "running":
            # A blocking provider call cannot be force-killed safely. Keep the
            # execution lease until that worker reaches its next fence, then
            # let _raise_if_cancelled publish the terminal cancelled state.
            lease_expires_at = run.lease_expires_at
            if lease_expires_at and lease_expires_at.tzinfo is None:
                lease_expires_at = lease_expires_at.replace(tzinfo=timezone.utc)
            run.cancel_requested = True
            if (
                run.package_id
                and (not lease_expires_at or lease_expires_at <= datetime.now(timezone.utc))
            ):
                self._finalize_post_persistence_cancel(session, run.id)
                return self.get_run(session, run_id=run_id, user_id=user_id)
            elif not lease_expires_at or lease_expires_at <= datetime.now(timezone.utc):
                run.status = "cancelled"
                run.current_step = "cancelled"
                run.finished_at = datetime.now(timezone.utc)
                run.active_attempt_id = None
                run.lease_expires_at = None
            else:
                run.current_step = "cancelling"
            session.add(run)
            session.commit()
            session.refresh(run)
        return self.get_run(session, run_id=run_id, user_id=user_id)

    def resume(self, session: Session, *, run_id: str, user_id: UUID) -> ResourceRunPublic | None:
        run = session.get(ResourceGenerationRun, run_id)
        if not run or run.user_id != user_id:
            return None
        if run.status == "running" and run.cancel_requested:
            lease_expires_at = run.lease_expires_at
            if lease_expires_at and lease_expires_at.tzinfo is None:
                lease_expires_at = lease_expires_at.replace(tzinfo=timezone.utc)
            if (
                run.package_id
                and (not lease_expires_at or lease_expires_at <= datetime.now(timezone.utc))
            ):
                self._finalize_post_persistence_cancel(session, run.id)
            elif not lease_expires_at or lease_expires_at <= datetime.now(timezone.utc):
                run.status = "cancelled"
                run.current_step = "cancelled"
                run.finished_at = datetime.now(timezone.utc)
                run.active_attempt_id = None
                run.lease_expires_at = None
                session.add(run)
                session.commit()
            return self.get_run(session, run_id=run_id, user_id=user_id)
        if run.status in {"requested", "cancelled", "failed"} and not run.package_id:
            run.status = "requested"
            run.current_step = "requested"
            run.cancel_requested = False
            run.error_code = None
            run.error_message = None
            run.finished_at = None
            run.active_attempt_id = None
            run.lease_expires_at = None
            session.add(run)
            session.commit()
            self.enqueue_requested_run(run.id)
            return self.get_run(session, run_id=run_id, user_id=user_id)
        if run.status == "running":
            lease_expires_at = run.lease_expires_at
            if lease_expires_at and lease_expires_at.tzinfo is None:
                lease_expires_at = lease_expires_at.replace(tzinfo=timezone.utc)
            if not lease_expires_at or lease_expires_at <= datetime.now(timezone.utc):
                run.status = "requested"
                run.current_step = "requested"
                run.error_code = "STALE_RUN_RECOVERED"
                run.error_message = "检测到执行租约已过期，已重新排队"
                run.active_attempt_id = None
                run.lease_expires_at = None
                session.add(run)
                session.commit()
                self.enqueue_requested_run(run.id)
            return self.get_run(session, run_id=run_id, user_id=user_id)
        if run.status != "partial_success" or not run.package_id:
            return self.get_run(session, run_id=run_id, user_id=user_id)
        package = session.get(GeneratedResourcePackage, run.package_id)
        if not package:
            return self.get_run(session, run_id=run_id, user_id=user_id)
        request = ResourceGenerationRequest.model_validate(run.requested)
        payload = resource_generation_service.get_package_payload(run.package_id)
        response = ResourceGenerationResponse.model_validate({
            **payload,
            "package_id": run.package_id,
            "subject": package.subject,
            "topic": package.topic,
            "generated_at": package.generated_at,
            "local_model_profile": package.model_profile,
            "agent_trace": package.agent_trace,
            "quality_notes": package.quality_notes,
        })
        resources = session.exec(select(Resource).where(Resource.package_id == run.package_id)).all()
        failures = list((run.shared_state or {}).get("stage_failures") or [])
        degraded_reasons = [
            reason
            for reason in list((run.shared_state or {}).get("degraded_reasons") or [])
            if reason != "POST_PERSISTENCE_CANCELLED"
        ]
        if not failures and degraded_reasons:
            return self.get_run(session, run_id=run_id, user_id=user_id)
        run.cancel_requested = False
        run.status = "running"
        run.current_step = "resuming"
        run.error_code = None
        run.error_message = None
        state = dict(run.shared_state or {})
        state["degraded_reasons"] = degraded_reasons
        run.shared_state = state
        session.add(run)
        session.commit()
        remaining: list[str] = []
        for stage in failures:
            try:
                if stage == "linking_graph" and request.course_id:
                    self._link_graph(session, run, request, response, resources)
                elif stage == "updating_path":
                    self._update_path(session, run, request, response)
                elif stage == "updating_profile":
                    self._update_profile(session, run, request, response, user_id)
            except Exception:
                session.rollback()
                remaining.append(stage)
        run.status = "partial_success" if remaining or degraded_reasons else "completed"
        run.current_step = "completed"
        run.finished_at = datetime.now(timezone.utc)
        state = dict(run.shared_state or {})
        state["stage_failures"] = remaining
        state["degraded_reasons"] = degraded_reasons
        run.shared_state = state
        package.status = run.status
        session.add(package)
        session.add(run)
        session.commit()
        return self.get_run(session, run_id=run_id, user_id=user_id)

    def list_recent(self, session: Session, *, owner_id: UUID, course_id: UUID | None = None, limit: int = 12) -> list[dict[str, Any]]:
        query = select(GeneratedResourcePackage).where(GeneratedResourcePackage.user_id == owner_id)
        if course_id:
            query = query.where(GeneratedResourcePackage.course_id == course_id)
        packages = session.exec(
            query.order_by(
                GeneratedResourcePackage.generated_at.desc(),
                GeneratedResourcePackage.id.desc(),
            ).limit(limit)
        ).all()
        output: list[dict[str, Any]] = []
        for package in packages:
            try:
                payload = self.get_package(session, package_id=package.id, user_id=owner_id)
            except ResourcePackagePersistenceError as exc:
                if exc.code != "ARTIFACT_INTEGRITY_FAILED":
                    raise
                logger.warning("skipping corrupt generated package %s", package.id)
                continue
            if payload:
                output.append(payload)
        return output

    @staticmethod
    def can_access(session: Session, *, package_id: str, user_id: UUID, is_superuser: bool) -> bool:
        package = session.get(GeneratedResourcePackage, package_id)
        return bool(package and (is_superuser or package.user_id == user_id))


resource_package_service = ResourcePackageService()
