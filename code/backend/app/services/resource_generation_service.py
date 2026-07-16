from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
import hashlib
import json
import os
import re
import shutil
from typing import Any
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID, uuid4

from langchain_core.messages import HumanMessage

from app.core.config import settings
from app.schemas.resource_generation import (
    GeneratedResourceArtifact,
    ResourceGenerationRequest,
    ResourceGenerationResponse,
    ResourceKind,
)
from app.services.chat_model_factory import ChatModelFactory
from app.services.content_quality_service import content_quality_service


DEFAULT_RESOURCE_TYPES: list[ResourceKind] = [
    "lecture_markdown",
    "lecture_docx",
    "lecture_pdf",
    "practice_markdown",
    "practice_docx",
    "practice_pdf",
    "mind_map",
    "reading_list",
    "case_project",
    "video_script",
    "quality_checklist",
]

AI_RESOURCE_KINDS: frozenset[ResourceKind] = frozenset(
    {
        "lecture_markdown",
        "practice_markdown",
        "mind_map",
        "reading_list",
        "case_project",
        "video_script",
        "quality_checklist",
    }
)

RESOURCE_AGENT_LABELS: dict[ResourceKind, tuple[str, str]] = {
    "lecture_markdown": ("LectureAgent", "个性化讲义"),
    "practice_markdown": ("ExerciseAgent", "分层练习"),
    "mind_map": ("MindMapAgent", "知识导图"),
    "reading_list": ("ReadingAgent", "拓展阅读"),
    "case_project": ("CaseAgent", "实操案例"),
    "video_script": ("ScriptAgent", "讲解脚本"),
    "quality_checklist": ("QualityAgent", "质量清单"),
}


class ResourceGenerationService:
    """Local-first resource producer for the course resource center."""

    def __init__(self) -> None:
        self.output_root = Path(settings.BASE_PATH) / "generated_resources"

    def generate(
        self,
        request: ResourceGenerationRequest,
        *,
        owner_id: UUID | None = None,
        runtime_context: dict[str, Any] | None = None,
    ) -> ResourceGenerationResponse:
        package_id = (
            f"rg_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}"
        )
        target_dir = self.output_root / package_id
        target_dir.mkdir(parents=True, exist_ok=True)

        kinds = request.resource_types or DEFAULT_RESOURCE_TYPES
        requested_ai_kinds = [kind for kind in kinds if kind in AI_RESOURCE_KINDS]
        # Word/PDF are renderings of the same reviewed semantic resource. When a
        # caller requests only a binary format, still generate and review its
        # Markdown source before rendering it.
        if any(kind in kinds for kind in ("lecture_docx", "lecture_pdf")) and "lecture_markdown" not in requested_ai_kinds:
            requested_ai_kinds.append("lecture_markdown")
        if any(kind in kinds for kind in ("practice_docx", "practice_pdf")) and "practice_markdown" not in requested_ai_kinds:
            requested_ai_kinds.append("practice_markdown")
        context = self._build_context(request, runtime_context=runtime_context)
        ai_contents, ai_generation_error = self._generate_ai_contents(context, requested_ai_kinds)
        ai_contents, artifact_retries, quality_results = self._review_and_rework_ai_contents(
            context,
            requested_ai_kinds,
            ai_contents,
        )
        fallback_artifacts = sorted(set(requested_ai_kinds) - set(ai_contents))
        finalized_contents: dict[str, str] = {}
        for content_kind in requested_ai_kinds:
            raw_content = ai_contents.get(content_kind) or self._fallback_content(content_kind, context)
            finalized_contents[content_kind] = self._finalize_content(
                content_kind,
                raw_content,
                context,
                quality_results,
                source="ai" if content_kind in ai_contents else "deterministic_fallback",
            )
        artifacts: list[GeneratedResourceArtifact] = []
        markdown_cache: dict[str, str] = {}

        for kind in kinds:
            if kind == "lecture_markdown":
                markdown_cache["lecture"] = finalized_contents["lecture_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 个性化讲义",
                        "lecture.md",
                        markdown_cache["lecture"],
                        "text/markdown",
                    )
                )
            elif kind == "lecture_docx":
                lecture = markdown_cache.get("lecture") or finalized_contents["lecture_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 讲义 Word",
                        "lecture.docx",
                        self._render_docx_bytes(lecture, title=f"{request.topic} 个性化讲义"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                )
            elif kind == "lecture_pdf":
                lecture = markdown_cache.get("lecture") or finalized_contents["lecture_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 讲义 PDF",
                        "lecture.pdf",
                        self._render_pdf_bytes(lecture, title=f"{request.topic} 个性化讲义"),
                        "application/pdf",
                    )
                )
            elif kind == "practice_markdown":
                markdown_cache["practice"] = finalized_contents["practice_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 分层练习",
                        "practice.md",
                        markdown_cache["practice"],
                        "text/markdown",
                    )
                )
            elif kind == "practice_docx":
                practice = markdown_cache.get("practice") or finalized_contents["practice_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 练习 Word",
                        "practice.docx",
                        self._render_docx_bytes(practice, title=f"{request.topic} 分层练习"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                )
            elif kind == "practice_pdf":
                practice = markdown_cache.get("practice") or finalized_contents["practice_markdown"]
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 练习 PDF",
                        "practice.pdf",
                        self._render_pdf_bytes(practice, title=f"{request.topic} 分层练习"),
                        "application/pdf",
                    )
                )
            elif kind == "mind_map":
                mind_map = finalized_contents[kind]
                mind_map = self._ensure_mind_map_resource_links(mind_map, kinds)
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 思维导图",
                        "mind-map.mmd",
                        mind_map,
                        "text/plain",
                    )
                )
            elif kind == "reading_list":
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 拓展阅读",
                        "reading-list.md",
                        finalized_contents[kind],
                        "text/markdown",
                    )
                )
            elif kind == "case_project":
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 实操案例",
                        "case-project.md",
                        finalized_contents[kind],
                        "text/markdown",
                    )
                )
            elif kind == "video_script":
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 数字人脚本",
                        "video-script.md",
                        finalized_contents[kind],
                        "text/markdown",
                    )
                )
            elif kind == "quality_checklist":
                artifacts.append(
                    self._write_artifact(
                        target_dir,
                        kind,
                        f"{request.topic} 使用审查清单",
                        "quality-checklist.md",
                        finalized_contents[kind],
                        "text/markdown",
                    )
                )

        self.write_manifest(
            target_dir,
            request=request,
            package_id=package_id,
            artifacts=artifacts,
            owner_id=owner_id,
        )

        return ResourceGenerationResponse(
            package_id=package_id,
            course_id=request.course_id,
            resource_id=request.resource_id,
            node_id=request.node_id,
            node_label=request.node_label,
            map_type=request.map_type,
            source=request.source,
            subject=request.subject,
            topic=request.topic,
            generated_at=datetime.now(timezone.utc),
            local_model_profile={
                "chat_provider": settings.CHAT_PROVIDER,
                "chat_model": self._active_chat_model_name(),
                "embedding_provider": settings.EMBEDDINGS_PROVIDER,
                "multimodal_model": settings.MULTIMODAL_MODEL,
                "deployment": "api-first"
                if settings.CHAT_PROVIDER.lower() == "mimo"
                else "provider-configured",
                "mode": "课程上下文 + MiMo 结构化生成 + 质量审查"
                if ai_contents
                else "课程画像 + 本地结构化回退 + 质量审查",
                "content_provider": "mimo"
                if len(ai_contents) == len(requested_ai_kinds)
                else ("mimo_partial" if ai_contents else "local_fallback"),
                "ai_generated_artifacts": sorted(ai_contents.keys()),
                "fallback_artifacts": fallback_artifacts,
                "domain": context["domain"],
                "fallback_reason": ai_generation_error or "",
                "runtime_context_digest": self._context_digest(runtime_context or {}),
                "grounding_mode": context["grounding_mode"],
                "evidence_gate": {
                    "passed": True,
                    "course_claims_allowed": context["has_course_evidence"],
                    "verifiable_citation_count": context["citation_count"],
                    "decision": "allow_grounded_claims" if context["has_course_evidence"] else "deny_and_generalize",
                    "policy": "course_claims_require_verifiable_source_id",
                },
                "artifact_retries": artifact_retries,
                "quality_results": quality_results,
                "agent_contracts": self._agent_contracts(requested_ai_kinds),
            },
            agent_trace=self._build_agent_trace(
                context,
                requested_ai_kinds,
                ai_contents,
            ),
            quality_notes=[
                f"已按“{context['scenario']}”组织案例，不输出空泛学习建议。",
                (
                    "讲义、练习、导图、阅读和案例均绑定可核验课程来源与后续学习动作。"
                    if context["has_course_evidence"]
                    else "讲义、练习、导图、阅读和案例按通用知识模式生成，未声称绑定课程资料。"
                ),
                "资源包包含 quality-checklist.md，可用于下载后逐项验收与学习闭环追踪。",
                "中文 PDF 与 Markdown 均为完整产物；PDF 支持标题、段落、列表、表格、代码、分页和页码。",
                "联网搜索默认受控关闭；若启用，外部资料必须在内容中单独标注来源。",
                (
                    f"已使用 {context['citation_count']} 条可核验来源约束课程依据。"
                    if context["has_course_evidence"]
                    else "未找到可核验课程来源；已降级为通用知识模式并移除课程资料归因。"
                ),
            ],
            artifacts=artifacts,
        )

    def list_recent_packages(
        self, limit: int = 12, course_id: str | None = None
    ) -> list[dict[str, object]]:
        self.output_root.mkdir(parents=True, exist_ok=True)
        packages: list[dict[str, object]] = []
        for folder in sorted(
            [item for item in self.output_root.iterdir() if item.is_dir()],
            key=lambda item: item.stat().st_mtime,
            reverse=True,
        ):
            artifacts = sorted(folder.iterdir(), key=lambda item: item.name)
            manifest = next((item for item in artifacts if item.suffix == ".json" and item.name == "manifest.json"), None)
            payload: dict[str, object] = {}
            if manifest and manifest.exists():
                try:
                    payload = json.loads(manifest.read_text(encoding="utf-8"))
                except Exception:
                    payload = {}
            if not payload:
                payload = self._infer_manifest_from_artifacts(folder)
            package_course_id = str(payload.get("course_id") or "")
            if course_id and package_course_id != course_id:
                continue
            packages.append(
                {
                    "package_id": folder.name,
                    "course_id": package_course_id,
                    "resource_id": payload.get("resource_id") or "",
                    "node_id": payload.get("node_id") or "",
                    "node_label": payload.get("node_label") or "",
                    "map_type": payload.get("map_type") or "",
                    "source": payload.get("source") or "",
                    "subject": payload.get("subject") or "",
                    "topic": payload.get("topic") or folder.name,
                    "generated_at": payload.get("generated_at")
                    or datetime.fromtimestamp(
                        folder.stat().st_mtime, timezone.utc
                    ).isoformat(),
                    "artifacts": self._recent_artifact_payloads(
                        folder,
                        artifacts,
                        payload,
                    ),
                }
            )
            if len(packages) >= limit:
                break
        return packages

    def read_package_manifest(self, package_id: str) -> dict[str, Any]:
        manifest = self.package_directory(package_id) / "manifest.json"
        if manifest.is_symlink():
            raise ValueError(f"Package manifest symlinks are not allowed: {package_id}")
        if not manifest.is_file():
            raise FileNotFoundError(f"Missing manifest for package {package_id}")
        payload = json.loads(manifest.read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            raise ValueError(f"Invalid manifest for package {package_id}")
        return payload

    def get_package_payload(self, package_id: str) -> dict[str, Any]:
        package_dir = self.package_directory(package_id)
        manifest = self.read_package_manifest(package_id)
        artifacts = sorted(package_dir.iterdir(), key=lambda item: item.name)
        return {
            **manifest,
            "package_id": package_id,
            "artifacts": self._recent_artifact_payloads(
                package_dir,
                artifacts,
                manifest,
            ),
        }

    def update_package_manifest(
        self,
        package_id: str,
        updates: dict[str, Any],
    ) -> None:
        package_dir = self.package_directory(package_id)
        manifest = package_dir / "manifest.json"
        payload = self.read_package_manifest(package_id)
        payload.update(updates)
        temporary = package_dir / "manifest.json.tmp"
        temporary.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        temporary.replace(manifest)

    def backfill_package_artifact_digests(self, package_id: str) -> dict[str, int]:
        """Establish SHA-256 baselines for an audited legacy package.

        Existing size metadata must match every file before any digest is
        written. The manifest replacement is atomic, so a partial backfill can
        never leave a package looking verified.
        """

        package_dir = self.package_directory(package_id)
        manifest_path = package_dir / "manifest.json"
        payload = self.read_package_manifest(package_id)
        raw_artifacts = payload.get("artifacts") or []
        if not isinstance(raw_artifacts, list) or not raw_artifacts:
            raise ValueError("Package manifest has no artifacts to verify")
        updated = 0
        verified = 0
        for metadata in raw_artifacts:
            if not isinstance(metadata, dict):
                raise ValueError("Package manifest contains invalid artifact metadata")
            file_name = str(metadata.get("file_name") or "")
            target = self.resolve_artifact_path(package_id, file_name)
            expected_size = int(metadata.get("file_size") or 0)
            if expected_size <= 0 or target.stat().st_size != expected_size:
                raise ValueError("Artifact size does not match package manifest")
            actual_digest = self._file_sha256(target)
            expected_digest = str(metadata.get("sha256") or "").strip().lower()
            if expected_digest and expected_digest != actual_digest:
                raise ValueError("Artifact digest does not match package manifest")
            if not expected_digest:
                metadata["sha256"] = actual_digest
                updated += 1
            verified += 1
        payload["artifact_digest_scheme"] = "sha256-v1"
        temporary = package_dir / "manifest.json.tmp"
        temporary.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        temporary.replace(manifest_path)
        return {"verified": verified, "updated": updated}

    def package_directory(self, package_id: str) -> Path:
        if not package_id or self._safe_file_name(package_id) != package_id:
            raise ValueError("Invalid generated package id")
        root = self.output_root.resolve()
        candidate = root / package_id
        if candidate.is_symlink():
            raise ValueError("Generated package symlinks are not allowed")
        target = candidate.resolve()
        if target.parent != root:
            raise ValueError("Invalid generated package path")
        return target

    def resolve_artifact_path(self, package_id: str, file_name: str) -> Path:
        if not file_name or self._safe_file_name(file_name) != file_name:
            raise ValueError("Invalid generated artifact name")
        package_dir = self.package_directory(package_id)
        candidate = package_dir / file_name
        if candidate.is_symlink():
            raise ValueError("Generated artifact symlinks are not allowed")
        target = candidate.resolve()
        if target.parent != package_dir or not target.is_file():
            raise FileNotFoundError(file_name)
        return target

    @staticmethod
    def _file_sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    def resolve_verified_artifact_path(
        self, package_id: str, file_name: str
    ) -> Path:
        target = self.resolve_artifact_path(package_id, file_name)
        manifest = self.read_package_manifest(package_id)
        metadata = next(
            (
                item
                for item in (manifest.get("artifacts") or [])
                if isinstance(item, dict) and item.get("file_name") == file_name
            ),
            None,
        )
        if not metadata:
            raise ValueError("Artifact is not declared in package manifest")
        expected_size = int(metadata.get("file_size") or 0)
        if expected_size <= 0 or target.stat().st_size != expected_size:
            raise ValueError("Artifact size does not match package manifest")
        expected_sha256 = str(metadata.get("sha256") or "").strip().lower()
        if not expected_sha256:
            raise ValueError("Artifact digest is missing from package manifest")
        if self._file_sha256(target) != expected_sha256:
            raise ValueError("Artifact digest does not match package manifest")
        return target

    def delete_package(self, package_id: str) -> None:
        package_dir = self.package_directory(package_id)
        if package_dir.is_dir():
            shutil.rmtree(package_dir)

    @staticmethod
    def _active_chat_model_name() -> str:
        provider = settings.CHAT_PROVIDER.lower()
        if provider == "ollama":
            return settings.OLLAMA_MODEL
        if provider == "mimo":
            return settings.MIMO_CHAT_MODEL or settings.CHAT_MODEL
        return settings.CHAT_MODEL

    @staticmethod
    def _artifact_kind_from_name(file_name: str) -> str:
        name = file_name.lower()
        if "practice" in name:
            if name.endswith(".pdf"):
                return "practice_pdf"
            if name.endswith(".docx"):
                return "practice_docx"
            return "practice_markdown"
        if "mind" in name:
            return "mind_map"
        if "reading" in name:
            return "reading_list"
        if "case" in name:
            return "case_project"
        if "video" in name or "script" in name:
            return "video_script"
        if "quality" in name or "check" in name:
            return "quality_checklist"
        if name.endswith(".pdf"):
            return "lecture_pdf"
        if name.endswith(".docx"):
            return "lecture_docx"
        return "lecture_markdown"

    @staticmethod
    def _infer_manifest_from_artifacts(folder: Path) -> dict[str, object]:
        lecture = folder / "lecture.md"
        subject = ""
        topic = folder.name
        if lecture.exists():
            try:
                text = lecture.read_text(encoding="utf-8", errors="ignore")
                for line in text.splitlines():
                    if line.startswith("课程："):
                        subject = line.replace("课程：", "", 1).strip()
                    elif line.startswith("# "):
                        topic = (
                            line.replace("# ", "", 1)
                            .replace(" 个性化讲义", "")
                            .strip()
                            or topic
                        )
            except Exception:
                pass
        return {
            "package_id": folder.name,
            "subject": subject,
            "topic": topic,
            "generated_at": datetime.fromtimestamp(
                folder.stat().st_mtime, timezone.utc
            ).isoformat(),
        }

    @classmethod
    def _recent_artifact_payloads(
        cls,
        folder: Path,
        artifacts: list[Path],
        manifest: dict[str, object],
    ) -> list[dict[str, object]]:
        raw_manifest_artifacts = manifest.get("artifacts", [])
        if not isinstance(raw_manifest_artifacts, list):
            raw_manifest_artifacts = []
        manifest_artifacts = {
            str(item.get("file_name")): item
            for item in raw_manifest_artifacts
            if isinstance(item, dict) and item.get("file_name")
        }
        payloads: list[dict[str, object]] = []
        for artifact in artifacts:
            if not artifact.is_file() or artifact.name == "manifest.json":
                continue
            metadata = manifest_artifacts.get(artifact.name, {})
            content_type = str(
                metadata.get("content_type") or cls._artifact_content_type(artifact)
            )
            payloads.append(
                {
                    "kind": metadata.get("kind")
                    or cls._artifact_kind_from_name(artifact.name),
                    "title": metadata.get("title")
                    or artifact.stem.replace("-", " ").replace("_", " "),
                    "file_name": artifact.name,
                    "download_url": metadata.get("download_url")
                    or f"/api/v1/resource-generation/artifacts/{folder.name}/{artifact.name}",
                    "file_size": artifact.stat().st_size,
                    "content_type": content_type,
                    "preview": metadata.get("preview")
                    or cls._artifact_preview(artifact, content_type),
                }
            )
        return payloads

    @staticmethod
    def _artifact_content_type(path: Path) -> str:
        if path.suffix.lower() == ".pdf":
            return "application/pdf"
        if path.suffix.lower() == ".docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if path.suffix.lower() in {".md", ".mmd", ".txt"}:
            return "text/markdown" if path.suffix.lower() == ".md" else "text/plain"
        return "application/octet-stream"

    @staticmethod
    def _artifact_preview(path: Path, content_type: str) -> str:
        if not content_type.startswith("text/"):
            return ""
        try:
            return path.read_text(encoding="utf-8", errors="ignore")[:500]
        except OSError:
            return ""

    @staticmethod
    def _context_digest(payload: dict[str, Any]) -> str:
        import hashlib
        return hashlib.sha256(
            json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
        ).hexdigest()

    def _build_context(
        self,
        request: ResourceGenerationRequest,
        runtime_context: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        goal = request.learning_goal or f"掌握 {request.topic} 的核心概念、典型题型和应用方法"
        difficulty_label = {
            "foundation": "基础巩固",
            "standard": "标准提升",
            "challenge": "挑战拓展",
        }[request.difficulty]
        terms = self._topic_terms(request.subject, request.topic)
        domain = self._domain_profile(request.subject, request.topic, terms)
        note_blocks = [
            f"概念边界：用自己的话解释 {request.topic}，并圈出 {terms[0]} 的判定条件",
            f"方法步骤：把 {terms[1]} 拆成可检查的 3-5 个动作",
            f"结果校验：用 {terms[2]}、反例或评价指标验证结论不过度推广",
            f"迁移记录：把课堂案例改写成一个来自 {domain['domain']} 的新问题",
        ]
        graph_nodes = [
            f"先修节点：{terms[0]}",
            f"核心节点：{request.topic}",
            f"方法节点：{terms[1]}",
            f"校验节点：{terms[2]}",
            f"应用节点：{domain['transfer']}",
        ]
        learning_sequence = [
            "先读讲义第 1-4 节，补齐定义、证据和案例链路",
            "再打开 mind-map.mmd，把每个节点对应到课堂笔记或资源文件",
            "完成 practice.md 的基础题和标准题，并把错题送入 AI 批改",
            "用 case-project.md 输出一份可提交任务，再用 quality-checklist.md 自查",
        ]
        quality_gate = [
            "每个结论必须能回指课堂笔记、知识图谱节点或练习解析",
            "每个练习必须有评分点、错因判断和下一步追练动作",
            "每个案例必须包含输入材料、操作步骤、验收标准和可迁移边界",
            "每个资源必须说明何时进入 AI 伴学、何时进入资源再生成",
        ]
        context = {
            "subject": request.subject.strip(),
            "topic": request.topic.strip(),
            "goal": goal.strip(),
            "difficulty": difficulty_label,
            "minutes": str(request.target_minutes),
            "terms": "、".join(terms),
            "primary": terms[0],
            "secondary": terms[1],
            "third": terms[2],
            "concept_cards": self._topic_concept_cards(request.subject, request.topic, terms),
            "profile": self._course_profile(request.subject, request.topic),
            "domain": domain["domain"],
            "scenario": domain["scenario"],
            "case": domain["case"],
            "evidence": "；".join(domain["evidence"]),
            "rubric": "；".join(domain["rubric"]),
            "mistakes": "；".join(domain["mistakes"]),
            "transfer": domain["transfer"],
            "note_blocks": "\n".join(f"- {item}" for item in note_blocks),
            "graph_nodes": "\n".join(f"- {item}" for item in graph_nodes),
            "learning_sequence": "\n".join(f"{index}. {item}" for index, item in enumerate(learning_sequence, 1)),
            "quality_gate": "\n".join(f"- {item}" for item in quality_gate),
            "resource_contract": (
                f"讲义负责讲清 {request.topic}，练习负责暴露错因，思维导图负责定位关系，"
                f"案例负责迁移应用，质量清单负责把学习动作闭环到 AI 批改和课程图谱。"
            ),
        }
        runtime = runtime_context or {}
        citations = [item for item in (runtime.get("citations") or []) if isinstance(item, dict)]
        verifiable_citations = content_quality_service.verifiable_evidence(citations)
        has_course_evidence = bool(verifiable_citations)
        profile_summary = str(runtime.get("profile_summary") or "").strip()
        evidence_summary = str(runtime.get("evidence_summary") or "").strip()
        if profile_summary:
            context["profile"] = profile_summary
        else:
            context["profile"] = "no_profile：当前用户暂无可用学习画像"
        if evidence_summary and has_course_evidence:
            context["evidence"] = evidence_summary
        else:
            context["evidence"] = "no_evidence：未检索到课程文档、原子证据或图谱节点引用"
        context["citations"] = citations
        context["citation_count"] = len(verifiable_citations)
        context["has_course_evidence"] = has_course_evidence
        context["grounding_mode"] = "course_evidence" if has_course_evidence else "general_knowledge"
        context["evidence_instruction"] = (
            "只能使用证据清单中具有真实 source id 的课程来源；不得扩展或编造章节、讲义、图谱节点。"
            if has_course_evidence
            else "未提供可核验课程来源。不得声称内容来自课程讲义、课堂笔记、课程图谱或知识图谱节点；必须明确标为通用学科知识。"
        )
        if not has_course_evidence:
            context["graph_nodes"] = context["graph_nodes"].replace("节点", "概念").replace("课程", "主题")
            context["learning_sequence"] = (
                "1. 先读生成的主题说明，补齐定义、条件和案例链路\n"
                "2. 再查看概念关系图，核对每条关系是否能由通用定义或示例支持\n"
                "3. 完成分层练习，并把错误按概念、条件、步骤和结论分类\n"
                "4. 将生成内容与正式教材或教师资料交叉核验后再纳入学习记录"
            )
            context["quality_gate"] = (
                "- 每个事实性结论必须能由通用定义、例子或反例解释\n"
                "- 不得声称引用了未提供的讲义、笔记、教材章节或图谱节点\n"
                "- 每个练习必须有评分点、错因判断和下一步追练动作\n"
                "- 使用前必须提示与正式课程资料交叉核验"
            )
            context["resource_contract"] = (
                f"讲解负责说明 {request.topic}，练习负责暴露错因，概念关系图负责定位关系，"
                "案例负责迁移应用；当前未绑定可核验课程资料。"
            )
        return context

    def _generate_ai_contents(
        self,
        ctx: dict[str, Any],
        kinds: list[ResourceKind],
    ) -> tuple[dict[str, str], str]:
        provider = settings.CHAT_PROVIDER.lower()
        if (
            not settings.RESOURCE_GENERATION_AI_ENABLED
            or provider != "mimo"
            or not settings.MIMO_API_KEY
        ):
            return {}, "mimo_not_configured"
        requested = [kind for kind in kinds if kind in AI_RESOURCE_KINDS]
        if not requested:
            return {}, ""
        return self._generate_ai_contents_parallel(ctx, requested)

    def _generate_ai_contents_parallel(
        self,
        ctx: dict[str, Any],
        requested: list[ResourceKind],
    ) -> tuple[dict[str, str], str]:
        contents: dict[str, str] = {}
        errors: list[str] = []
        worker_count = min(4, len(requested))
        with ThreadPoolExecutor(
            max_workers=worker_count,
            thread_name_prefix="resource-agent",
        ) as executor:
            pending = {
                executor.submit(self._generate_single_ai_content, ctx, kind): kind
                for kind in requested
            }
            for future in as_completed(pending):
                kind = pending[future]
                try:
                    content = future.result()
                except Exception as exc:
                    errors.append(exc.__class__.__name__)
                    continue
                if content:
                    contents[kind] = content
        if not contents:
            return {}, errors[0] if errors else "empty_ai_contents"
        if len(contents) != len(requested):
            return contents, errors[0] if errors else "partial_ai_contents"
        return contents, ""

    @staticmethod
    def _build_agent_trace(
        ctx: dict[str, Any],
        requested: list[ResourceKind],
        ai_contents: dict[str, str],
    ) -> list[str]:
        trace = [
            "ProfileAgent: 读取学习画像和目标难度",
            f"DomainAgent: 识别课程域为 {ctx['domain']}",
            (
                f"EvidenceAgent: 验证 {ctx['citation_count']} 条带 source id 的课程依据"
                if ctx["has_course_evidence"]
                else "EvidenceAgent: 未发现可核验课程依据，切换为通用知识模式"
            ),
            f"ResourcePlannerAgent: 规划 {len(requested)} 类资源并并行执行",
        ]
        for kind in requested:
            agent, label = RESOURCE_AGENT_LABELS[kind]
            provider = "MiMo 生成完成" if kind in ai_contents else "本地结构化回退完成"
            trace.append(f"{agent}: {label}，{provider}")
        trace.extend(
            [
                "SafetyReviewAgent: 检查事实边界、缩写完整性、证据归因和输出格式",
                "FinalizerAgent: 汇总并写入可下载资源包",
            ]
        )
        return trace

    def _generate_single_ai_content(
        self,
        ctx: dict[str, Any],
        kind: str,
    ) -> str:
        max_tokens = {
            "lecture_markdown": 2400,
            "practice_markdown": 2600,
            "mind_map": 1400,
            "reading_list": 1400,
            "case_project": 2200,
            "video_script": 2000,
            "quality_checklist": 2200,
        }.get(kind, 1800)
        model = ChatModelFactory.create(
            temperature=0.2,
            max_tokens=max_tokens,
            top_p=0.9,
            model_name=settings.MIMO_FAST_MODEL or settings.MIMO_CHAT_MODEL,
            timeout_seconds=settings.RESOURCE_GENERATION_TIMEOUT_SECONDS,
        )
        response = model.invoke(
            [
                HumanMessage(
                    content=self._resource_generation_prompt(ctx, [kind])
                )
            ]
        )
        raw = getattr(response, "content", response)
        raw_text = str(raw)
        payload = self._parse_structured_payload(raw_text, [kind])
        try:
            if not payload:
                payload = self._parse_json_payload(raw_text)
        except Exception:
            if not payload:
                payload = {kind: raw_text.strip()}
        contents = self._validate_ai_contents(payload, [kind], ctx)
        return contents.get(kind, "")

    @staticmethod
    def _agent_contracts(kinds: list[ResourceKind]) -> dict[str, dict[str, Any]]:
        gates: dict[str, str] = {
            "lecture_markdown": "完整讲义；至少两个小节；包含概念、例子与边界",
            "practice_markdown": "至少三个层次的练习；包含答案或评分框架",
            "mind_map": "合法 Mermaid mindmap/graph；包含核心主题与关系",
            "reading_list": "每条阅读给出来源、用途与引用核验提示",
            "case_project": "包含场景、任务、交付物与验收量规",
            "video_script": "包含分段讲解、互动停顿与镜头提示",
            "quality_checklist": "包含事实、引用、难度与安全审查项",
        }
        return {
            kind: {
                "agent_role": RESOURCE_AGENT_LABELS[kind][0],
                "input_schema": ["subject", "topic", "goal", "difficulty", "course_evidence"],
                "output_schema": {"kind": kind, "content": "markdown_or_mermaid", "citations": "list"},
                "quality_gate": gates[kind],
                "max_retries": 2,
            }
            for kind in kinds
        }

    def _review_and_rework_ai_contents(
        self,
        ctx: dict[str, Any],
        requested: list[ResourceKind],
        contents: dict[str, str],
    ) -> tuple[dict[str, str], dict[str, int], dict[str, dict[str, Any]]]:
        accepted = dict(contents)
        retry_counts: dict[str, int] = {kind: 0 for kind in requested}
        results: dict[str, dict[str, Any]] = {}
        for kind in requested:
            content = accepted.get(kind, "")
            passed, reasons = self._artifact_quality_gate(
                kind,
                content,
                ctx["topic"],
                has_course_evidence=bool(ctx["has_course_evidence"]),
            )
            while content and not passed and retry_counts[kind] < 2:
                retry_counts[kind] += 1
                try:
                    content = self._generate_single_ai_content(ctx, kind)
                except Exception as exc:
                    reasons = [f"rework_error:{exc.__class__.__name__}"]
                    content = ""
                    break
                passed, reasons = self._artifact_quality_gate(
                    kind,
                    content,
                    ctx["topic"],
                    has_course_evidence=bool(ctx["has_course_evidence"]),
                )
            if passed:
                accepted[kind] = content
            else:
                accepted.pop(kind, None)
            results[kind] = {
                "passed": passed,
                "reasons": reasons,
                "retry_count": retry_counts[kind],
                "reviewer": "CriticSafetyAgent",
            }
        return accepted, retry_counts, results

    @staticmethod
    def _artifact_quality_gate(
        kind: ResourceKind,
        content: str,
        topic: str,
        *,
        has_course_evidence: bool = False,
    ) -> tuple[bool, list[str]]:
        result = content_quality_service.review(
            kind=kind,
            content=content,
            topic=topic,
            has_course_evidence=has_course_evidence,
        )
        return result.passed, list(result.reasons)

    @staticmethod
    def _fallback_content(kind: ResourceKind, ctx: dict[str, Any]) -> str:
        factories = {
            "lecture_markdown": ResourceGenerationService._lecture_markdown,
            "practice_markdown": ResourceGenerationService._practice_markdown,
            "mind_map": ResourceGenerationService._mind_map,
            "reading_list": ResourceGenerationService._reading_list,
            "case_project": ResourceGenerationService._case_project,
            "video_script": ResourceGenerationService._video_script,
            "quality_checklist": ResourceGenerationService._quality_checklist,
        }
        factory = factories.get(kind)
        if factory is None:
            raise ValueError(f"No semantic fallback for resource kind: {kind}")
        return factory(ctx)

    @staticmethod
    def _finalize_content(
        kind: ResourceKind,
        content: str,
        ctx: dict[str, Any],
        quality_results: dict[str, dict[str, Any]],
        *,
        source: str,
    ) -> str:
        initial = content_quality_service.review(
            kind=kind,
            content=content,
            topic=str(ctx["topic"]),
            has_course_evidence=bool(ctx["has_course_evidence"]),
        )
        remediated = content
        actions: list[str] = []
        remediated, acronym_actions = content_quality_service.repair_acronym_completeness(
            remediated,
            str(ctx["topic"]),
            kind=kind,
        )
        actions.extend(acronym_actions)
        if not ctx["has_course_evidence"]:
            remediated, replacements = content_quality_service.neutralize_ungrounded_course_claims(
                remediated,
                kind=kind,
            )
            if replacements:
                actions.append("removed_unsupported_course_attribution")
        final = content_quality_service.review(
            kind=kind,
            content=remediated,
            topic=str(ctx["topic"]),
            has_course_evidence=bool(ctx["has_course_evidence"]),
        )
        previous = dict(quality_results.get(kind) or {})
        quality_results[kind] = {
            **previous,
            "passed": final.passed,
            "source": source,
            "degraded": source != "ai" or bool(actions),
            "initial_reasons": list(initial.reasons),
            "reasons": list(final.reasons),
            "remediation_actions": actions,
            "checks": final.checks,
            "grounding_mode": ctx["grounding_mode"],
            "reviewer": "CriticSafetyAgent",
        }
        if not final.passed:
            raise ValueError(f"Content quality gate failed for {kind}: {','.join(final.reasons)}")
        return remediated

    @staticmethod
    def _resource_generation_prompt(ctx: dict[str, Any], kinds: list[str]) -> str:
        tags = "\n".join(
            f"<{kind}>\n请在这里输出 {kind} 正文\n</{kind}>" for kind in kinds
        )
        grounding_requirements = (
            "课程依据只允许引用证据清单中已有的真实来源；每项课程归因都必须可回指 source id。"
            if ctx["has_course_evidence"]
            else "当前没有可核验课程来源。内容必须标明基于通用学科知识生成，不得出现课程讲义、课堂笔记、课程图谱或知识图谱节点等来源归因。"
        )
        return f"""你是教育 SaaS 平台的课程资源生成器。请只输出下面这些 XML 风格标签，不要输出 Markdown 代码围栏，不要添加标签之外的解释。

输出结构:
{tags}

生成要求：
1. 每个请求的标签必须完整出现，开始标签和结束标签必须完全匹配。
2. 内容必须围绕课程《{ctx['subject']}》和知识点“{ctx['topic']}”，不得泛泛而谈。
3. 必须绑定学习目标、错因诊断和后续学习动作；课程归因必须经过证据门禁。
4. 不得编造外部文献、教材章节或课程来源；阅读清单只能列可核验的资料类型和核验动作。
5. 每个 Markdown 类资源至少包含标题、学习目标、内容依据、概念关系、学习任务、质量自查。
6. mind_map 字段必须输出 Mermaid mindmap 文本，根节点是“{ctx['topic']}”。
7. practice_markdown 必须包含基础题、标准题、挑战题、答案框架和错因追练。
8. quality_checklist 必须能用于验收资源是否真实服务学习闭环。
9. 不允许出现“第X章”“第Y次课”“某教材”“待补充”“占位”等占位文案；资料无法确定时必须明确写“未提供可核验来源”，不得用看似具体的课程表述代替证据。
10. {grounding_requirements}

课程上下文：
- 课程：{ctx['subject']}
- 知识点：{ctx['topic']}
- 学习目标：{ctx['goal']}
- 难度：{ctx['difficulty']}
- 建议时长：{ctx['minutes']} 分钟
- 课程域：{ctx['domain']}
- 课程画像：{ctx['profile']}
- 应用场景：{ctx['scenario']}
- 案例线索：{ctx['case']}
- 证据清单：{ctx['evidence']}
- 证据模式：{ctx['grounding_mode']}
- 证据门禁：{ctx['evidence_instruction']}
- 评分量规：{ctx['rubric']}
- 常见错因：{ctx['mistakes']}
- 迁移目标：{ctx['transfer']}
- 图谱节点：
{ctx['graph_nodes']}
- 学习路径：
{ctx['learning_sequence']}
- 质量门禁：
{ctx['quality_gate']}
"""

    @staticmethod
    def _parse_structured_payload(raw: str, requested: list[str]) -> dict[str, Any]:
        payload: dict[str, Any] = {}
        for kind in requested:
            match = re.search(
                rf"<{re.escape(kind)}>\s*(.*?)\s*</{re.escape(kind)}>",
                raw,
                flags=re.DOTALL | re.IGNORECASE,
            )
            if match:
                payload[kind] = match.group(1).strip()
        return payload

    @staticmethod
    def _parse_json_payload(raw: str) -> dict[str, Any]:
        text = raw.strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text).strip()
        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            text = text[start : end + 1]
        payload = json.loads(text)
        return payload if isinstance(payload, dict) else {}

    @staticmethod
    def _validate_ai_contents(
        payload: dict[str, Any],
        requested: list[str],
        ctx: dict[str, Any],
    ) -> dict[str, str]:
        contents: dict[str, str] = {}
        for kind in requested:
            value = payload.get(kind)
            if not isinstance(value, str):
                continue
            text = ResourceGenerationService._strip_artifact_tags(value.strip(), kind)
            text = ResourceGenerationService._sanitize_model_markdown(text)
            if len(text) < 240 and kind != "mind_map":
                continue
            if ResourceGenerationService._contains_placeholder(text):
                continue
            if ResourceGenerationService._contains_protocol_markup(text):
                continue
            if kind == "mind_map":
                if "mindmap" not in text.lower() or ctx["topic"] not in text:
                    continue
            elif ctx["topic"] not in text:
                continue
            elif ctx["subject"] not in text:
                text = f"课程：{ctx['subject']}\n\n{text}"
            contents[kind] = text
        return contents

    @staticmethod
    def _sanitize_model_markdown(text: str) -> str:
        """Normalize a tiny HTML subset that models sometimes mix into Markdown.

        Downloadable artifacts deliberately render model-authored HTML as inert
        text. Removing only harmless inline-formatting tags here prevents
        strings such as ``<u>StudentID</u>`` from leaking into Word/PDF while
        leaving SQL comparison operators and code blocks untouched.
        """
        normalized = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        return re.sub(
            r"</?(?:u|em|strong|b|i)(?:\s+[^>]*)?>",
            "",
            normalized,
            flags=re.IGNORECASE,
        ).strip()

    @staticmethod
    def _contains_placeholder(text: str) -> bool:
        patterns = [
            r"第\s*[XxYyZz]\s*(章|节|次|讲)",
            r"某(教材|资料|章节|案例|文件)",
            r"待(补充|填写|完善|确认)",
            r"TBD",
            r"占位",
            r"xxx+",
        ]
        return any(re.search(pattern, text) for pattern in patterns)

    @staticmethod
    def _strip_artifact_tags(text: str, kind: str) -> str:
        text = re.sub(
            rf"^\s*<{re.escape(kind)}>\s*",
            "",
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            rf"\s*</{re.escape(kind)}>\s*$",
            "",
            text,
            flags=re.IGNORECASE,
        )
        return text.strip()

    @staticmethod
    def _contains_protocol_markup(text: str) -> bool:
        return bool(
            re.search(
                r"</?(lecture_markdown|practice_markdown|mind_map|reading_list|case_project|video_script|quality_checklist|resource)\b",
                text,
                flags=re.IGNORECASE,
            )
        )

    @staticmethod
    def _ensure_mind_map_resource_links(
        content: str, kinds: list[ResourceKind]
    ) -> str:
        """Keep the generated map aligned with files that exist in the package."""
        file_names: dict[ResourceKind, str] = {
            "lecture_markdown": "lecture.md",
            "lecture_docx": "lecture.docx",
            "lecture_pdf": "lecture.pdf",
            "practice_markdown": "practice.md",
            "practice_docx": "practice.docx",
            "practice_pdf": "practice.pdf",
            "reading_list": "reading-list.md",
            "case_project": "case-project.md",
            "video_script": "video-script.md",
            "quality_checklist": "quality-checklist.md",
        }
        missing = [
            file_names[kind]
            for kind in kinds
            if kind in file_names and file_names[kind] not in content
        ]
        if not missing:
            return content
        package_index = "\n".join(
            ["  资源包索引", *(f"    {file_name}" for file_name in missing)]
        )
        return f"{content.rstrip()}\n{package_index}\n"

    def _write_artifact(
        self,
        target_dir: Path,
        kind: ResourceKind,
        title: str,
        file_name: str,
        content: str | bytes,
        content_type: str,
    ) -> GeneratedResourceArtifact:
        safe_name = self._safe_file_name(file_name)
        path = target_dir / safe_name
        if isinstance(content, bytes):
            path.write_bytes(content)
            preview = ""
        else:
            path.write_text(content, encoding="utf-8")
            preview = content[:500]
        stat = path.stat()
        return GeneratedResourceArtifact(
            kind=kind,
            title=title,
            file_name=safe_name,
            download_url=f"/api/v1/resource-generation/artifacts/{target_dir.name}/{safe_name}",
            content_type=content_type,
            file_size=stat.st_size,
            preview=preview,
        )

    def write_manifest(
        self,
        target_dir: Path,
        *,
        request: ResourceGenerationRequest,
        package_id: str,
        artifacts: list[GeneratedResourceArtifact],
        owner_id: UUID | None = None,
    ) -> None:
        manifest = {
            "package_id": package_id,
            "owner_id": str(owner_id) if owner_id else "",
            "course_id": str(request.course_id) if request.course_id else "",
            "resource_id": request.resource_id or "",
            "node_id": request.node_id or "",
            "node_label": request.node_label or "",
            "map_type": request.map_type or "",
            "source": request.source or "",
            "subject": request.subject,
            "topic": request.topic,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "persistence_status": "file_only",
            "persisted_resource_ids": [],
            "artifacts": [
                {
                    "kind": artifact.kind,
                    "title": artifact.title,
                    "file_name": artifact.file_name,
                    "download_url": artifact.download_url,
                    "content_type": artifact.content_type,
                    "file_size": artifact.file_size,
                    "sha256": self._file_sha256(
                        target_dir / artifact.file_name
                    ),
                    "preview": artifact.preview,
                }
                for artifact in artifacts
            ],
        }
        (target_dir / "manifest.json").write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    @staticmethod
    def _topic_terms(subject: str, topic: str) -> list[str]:
        text = f"{subject} {topic}".lower()
        if any(term in text for term in ("数学", "函数", "方程", "几何", "概率", "导数", "抛物线")):
            return ["数学定义", "参数与图像关系", "解题条件", "数形结合", "结果检验"]
        if any(term in text for term in ("acid", "事务", "转账一致性")):
            return ["原子性", "一致性", "隔离性", "持久性", "日志与恢复"]
        if "数据库" in text or "sql" in text:
            return ["关系模型", "SQL 查询", "完整性约束", "事务并发", "规范化"]
        if "数据结构" in text or "算法" in text:
            return ["抽象数据类型", "复杂度分析", "遍历策略", "存储结构", "边界条件"]
        if "人工智能" in text or "ai" in text:
            return ["问题建模", "搜索策略", "知识表示", "模型训练", "评估指标"]
        if "宏观" in text or "经济" in text:
            return ["总量指标", "总需求", "政策传导", "经济周期", "长期增长"]
        if "审计" in text:
            return ["审计目标", "重大错报风险", "内部控制", "审计证据", "审计意见"]
        if "金融" in text:
            return ["时间价值", "现金流折现", "风险收益", "投资组合", "资本成本"]
        return [topic, "定义边界", "适用条件", "典型案例", "自测反馈"]

    @staticmethod
    def _topic_concept_cards(subject: str, topic: str, terms: list[str]) -> str:
        """Return accurate fallback concept rows when a specialist times out."""
        text = f"{subject} {topic}".lower()
        if any(term in text for term in ("acid", "事务", "转账一致性")):
            rows = [
                ("原子性", "事务中的操作要么全部完成，要么全部撤销", "能解释扣款成功而入账失败时为何必须回滚"),
                ("一致性", "事务前后都必须满足完整性约束和业务规则", "能检查转账前后总金额与余额约束是否成立"),
                ("隔离性", "并发事务不应读取彼此未提交的中间状态", "能识别脏读等并发异常及其隔离级别"),
                ("持久性", "事务一旦提交，结果在故障后仍可恢复", "能说明日志、REDO 与持久化的关系"),
            ]
        else:
            rows = [
                (terms[0], "界定核心对象、定义与适用边界", "能否说清定义、输入和输出"),
                (terms[1], "组织主要推理或操作步骤", "能否列出至少三个判断条件"),
                (terms[2], "检验结果、反例和边界条件", "能否解释一个成立与不成立的例子"),
            ]
        return "\n".join(f"| {name} | {role} | {check} |" for name, role, check in rows)

    @staticmethod
    def _course_profile(subject: str, topic: str) -> str:
        return (
            f"课程《{subject}》当前围绕“{topic}”组织资料。资源包默认面向学生自学，"
            "先补概念边界，再给案例和练习，最后用检查清单确认是否能迁移应用。"
        )

    @staticmethod
    def _domain_profile(subject: str, topic: str, terms: list[str]) -> dict[str, object]:
        text = f"{subject} {topic}".lower()
        primary, secondary, third = terms[:3]
        if any(term in text for term in ("数学", "函数", "方程", "几何", "概率", "导数", "抛物线")):
            return {
                "domain": "数学课程",
                "scenario": f"围绕 {topic} 识别条件、建立表达式、分析图像或数量关系，并用代入与反例检验结论",
                "case": f"给出与 {topic} 相关的表达式、图像或实际情境，要求说明参数作用、完成推导并验证答案。",
                "evidence": [
                    "定义、符号和取值条件是否完整",
                    "推导步骤是否能回到公式或图像依据",
                    "结果是否通过代入、特殊值或反例检验",
                ],
                "rubric": ["概念与条件 30%", "推导过程 35%", "结果检验 20%", "表达规范 15%"],
                "mistakes": ["忽略参数取值条件", "平移方向或符号判断错误", "只写结果不保留推导与检验"],
                "transfer": "能在表达式、图像和实际情境之间转换，并判断同一方法的适用边界。",
            }
        if "数据库" in text or "sql" in text:
            return {
                "domain": "数据库课程",
                "scenario": "把校园选课、成绩登记或图书借阅业务转化为关系模式与 SQL 查询任务",
                "case": "给出学生、课程、选课三张表，要求识别主键/外键、写出连接查询，并说明事务提交失败时如何恢复。",
                "evidence": [
                    "关系模式是否标出主键、外键和属性域",
                    "SQL 是否能从投影、筛选、连接和聚合四步追溯",
                    "事务分析是否说明 ACID 与并发调度依据",
                ],
                "rubric": [
                    "概念边界 25%",
                    "SQL 或模式设计步骤 35%",
                    "约束与异常处理 25%",
                    "可读性与验证 15%",
                ],
                "mistakes": [
                    "把外键写成普通属性",
                    "JOIN 条件缺失导致笛卡尔积",
                    "只写最终 SQL 不解释业务约束",
                ],
                "transfer": "能把新的业务描述拆成实体、联系、约束、查询和事务五类对象。",
            }
        if "数据结构" in text or "算法" in text:
            return {
                "domain": "数据结构课程",
                "scenario": "为一个检索、排队或路径问题选择合适的数据结构并分析复杂度",
                "case": "根据操作频率选择数组、链表、堆或图存储，写出关键操作伪代码并计算最坏复杂度。",
                "evidence": [
                    "问题操作是否拆成访问、插入、删除、遍历",
                    "结构选择是否匹配复杂度需求",
                    "边界条件是否覆盖空结构、单元素和重复值",
                ],
                "rubric": [
                    "抽象建模 25%",
                    "结构选择 25%",
                    "算法步骤 30%",
                    "复杂度与边界 20%",
                ],
                "mistakes": [
                    "只背结构定义，不说明适用操作",
                    "平均复杂度和最坏复杂度混用",
                    "忽略空指针或越界条件",
                ],
                "transfer": "能从题目中的操作频率反推结构选择，而不是凭关键词套模板。",
            }
        if "人工智能" in text or "ai" in text:
            return {
                "domain": "人工智能课程",
                "scenario": "把真实问题转化为状态空间、知识表示或模型训练任务",
                "case": "为路径规划、诊断或分类任务定义状态、动作、评价指标，并比较搜索与学习方法的适用边界。",
                "evidence": [
                    "是否明确输入、输出、状态或样本标签",
                    "是否说明启发函数、知识表示或模型假设",
                    "是否用准确率、召回率或代价函数评价结果",
                ],
                "rubric": [
                    "问题建模 30%",
                    "方法选择 25%",
                    "评价指标 25%",
                    "风险与边界 20%",
                ],
                "mistakes": [
                    "把 AI 方法名称当成答案",
                    "不区分训练数据、验证数据和测试数据",
                    "忽略偏差、方差或不可解释性风险",
                ],
                "transfer": "能判断一个任务更适合搜索、规则推理、传统机器学习还是深度学习。",
            }
        if "宏观" in text or "经济" in text:
            return {
                "domain": "宏观经济学课程",
                "scenario": "用总量指标和模型解释政策变化对产出、价格和就业的影响",
                "case": "分析降准、财政扩张或外需下降对 AD-AS、利率和就业的传导路径。",
                "evidence": [
                    "变量是否区分名义量与实际量",
                    "传导链是否包含部门、市场和时间滞后",
                    "结论是否说明短期与长期差异",
                ],
                "rubric": [
                    "指标解释 25%",
                    "模型推理 35%",
                    "政策边界 25%",
                    "图形表达 15%",
                ],
                "mistakes": [
                    "把相关关系当因果关系",
                    "只说政策方向，不写传导链",
                    "忽略价格水平和实际产出的区别",
                ],
                "transfer": "能把新闻中的宏观政策转写为变量变化和模型移动。",
            }
        if "审计" in text:
            return {
                "domain": "审计学课程",
                "scenario": "围绕重大错报风险设计审计程序并评价证据是否充分适当",
                "case": "对收入确认或存货跌价风险设计询问、观察、函证和重新计算程序。",
                "evidence": [
                    "是否从认定出发识别风险",
                    "程序是否能回应对应风险",
                    "证据是否同时评价充分性与适当性",
                ],
                "rubric": [
                    "风险识别 30%",
                    "程序设计 30%",
                    "证据评价 25%",
                    "职业判断 15%",
                ],
                "mistakes": [
                    "把审计目标和具体程序混写",
                    "只列程序不说明能验证哪项认定",
                    "证据数量和证据质量不区分",
                ],
                "transfer": "能针对新的业务循环选择匹配的审计程序和证据标准。",
            }
        if "金融" in text:
            return {
                "domain": "金融学课程",
                "scenario": "用现金流、风险收益和资本成本分析投资或融资决策",
                "case": "根据债券现金流、股票估值或项目 NPV 判断是否投资，并说明风险调整依据。",
                "evidence": [
                    "现金流时点是否列清",
                    "贴现率是否与风险水平匹配",
                    "结论是否包含敏感性或情景分析",
                ],
                "rubric": [
                    "现金流建模 30%",
                    "折现与估值 30%",
                    "风险解释 25%",
                    "决策建议 15%",
                ],
                "mistakes": [
                    "混淆现值和终值",
                    "贴现率随意取值",
                    "只算结果不解释风险来源",
                ],
                "transfer": "能把任意资产或项目拆成现金流、折现率、风险和决策标准。",
            }
        return {
            "domain": "通用课程",
            "scenario": f"围绕 {topic} 完成概念解释、案例拆解和迁移练习",
            "case": f"选择一个与 {topic} 相关的课堂案例，标出条件、方法、结论和边界。",
            "evidence": [
                f"是否能定义 {primary}",
                f"是否能说明 {secondary} 的适用条件",
                f"是否能用 {third} 做结果校验",
            ],
            "rubric": ["概念准确 30%", "步骤完整 30%", "应用迁移 25%", "表达清晰 15%"],
            "mistakes": ["只背结论", "条件遗漏", "缺少边界校验"],
            "transfer": "能把本节主题迁移到新的题目或真实问题。",
        }

    @staticmethod
    def _safe_file_name(value: str) -> str:
        name = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip()).strip("-")
        return name or f"artifact-{uuid4().hex[:8]}.md"

    @staticmethod
    def _lecture_markdown(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 个性化讲义

课程：{ctx['subject']}
目标：{ctx['goal']}
难度：{ctx['difficulty']}
建议学习时长：{ctx['minutes']} 分钟
学习侧重点：{ctx['profile']}
课程域：{ctx['domain']}
资源契约：{ctx['resource_contract']}

## 1. 一句话定位
{ctx['topic']} 是本节学习的核心对象。学习时先弄清它解决什么问题，再看它和 {ctx['primary']}、{ctx['secondary']}、{ctx['third']} 的关系。不要只背术语，必须能把定义、证据、步骤和适用边界串成一条可复述的链路。

## 2. 核心概念卡
| 概念 | 课堂定位 | 学习检查 |
| --- | --- | --- |
{ctx['concept_cards']}

## 3. 课堂笔记对齐
把课堂笔记整理成下面四个块，后续练习、导图和案例都从这里取证：

{ctx['note_blocks']}

## 4. 课程证据清单
{ctx['evidence']}

使用资料时先把每个结论对应到上述证据项。若某一步没有证据，只能标为“推断”，并在 AI 伴学中要求补证或换例。

## 5. 课堂案例拆解
场景：{ctx['scenario']}

案例：{ctx['case']}

1. 标出题干或材料中的已知条件。
2. 判断这些条件分别对应 {ctx['primary']}、{ctx['secondary']} 还是 {ctx['third']}。
3. 写出推理链：条件 -> 方法 -> 中间结果 -> 结论。
4. 用一个反例或边界条件检查结论是否过度推广。

## 6. 课程图谱绑定
学习时按下面节点在课程图谱中逐个点开，确认每个节点都有笔记、资源和练习：

{ctx['graph_nodes']}

## 7. 易错点
{ctx['mistakes']}

## 8. 学习路径
{ctx['learning_sequence']}

## 9. 迁移目标
{ctx['transfer']}

## 10. 自我检查
- 我能否不用教材原句解释 {ctx['topic']}？
- 我能否指出 {ctx['primary']} 在案例中的证据？
- 我能否说出一个不适用 {ctx['topic']} 的场景？
- 我能否把本资源中的一个节点拖入课程图谱，并说明它连接了哪份练习？
"""

    @staticmethod
    def _practice_markdown(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 分层练习

课程：{ctx['subject']}
匹配场景：{ctx['scenario']}
练习目标：通过分层题暴露“定义不清、条件遗漏、步骤跳跃、结论无证据”四类错因。

## 基础题
1. 用 80 字以内解释 {ctx['topic']}，并写出它和 {ctx['primary']} 的关系。
2. 判断题：只要题干出现关键词，就一定可以套用 {ctx['topic']}。请说明理由，并指出需要补充的条件。
3. 填空：解决这类问题时，第一步应先识别 ______，第二步再选择 ______。

## 标准题
4. 给定一个课程案例，列出已知条件、适用概念、推理步骤、最终结论和证据来源。
5. 设计一道同类变式题，并写出标准答案和评分点。

## 挑战题
6. 比较 {ctx['topic']} 与 {ctx['secondary']} 的差异，至少列出 3 个判断标准。
7. 写一个容易出错的答案，并说明它错在定义、条件、步骤还是结论。

## 标准答案框架
| 题号 | 合格答案应包含 | 常见错因 | AI 追练指令 |
| --- | --- | --- | --- |
| 1 | 定义、适用对象、与 {ctx['primary']} 的关系 | 只背原句，没有边界 | 要求 AI 追问一个反例 |
| 2 | 判断为否，并说明关键词不足以替代条件 | 关键词套模板 | 要求 AI 生成两个相似但不适用的题 |
| 3 | 先识别条件，再选择方法或表示方式 | 步骤跳跃 | 要求 AI 按步骤逐格批改 |
| 4 | 条件表、概念表、步骤表、结论表、证据表 | 结论没有证据 | 要求 AI 用评分量规打分 |
| 5 | 题干、参考答案、评分点、变式说明 | 只换数字不换能力点 | 要求 AI 判断变式是否有效 |
| 6 | 至少 3 个标准，并对应示例 | 概念混淆 | 要求 AI 做概念对照卡 |
| 7 | 错误答案、错因分类、订正答案 | 只说“粗心” | 要求 AI 继续生成错因同类题 |

## 评分量规
{ctx['rubric']}

## 课堂笔记回填
完成练习后，把错因回填到课堂笔记：

{ctx['note_blocks']}

## 完成后的闭环动作
1. 把低于 80 分的题目送入 AI 批改，要求输出“错因 -> 订正 -> 追练”。
2. 在课程图谱里标记对应节点为“待巩固”，并关联本练习题号。
3. 若连续两题错在同一类原因，重新生成 15 分钟微资源包。
"""

    @staticmethod
    def _mind_map(ctx: dict[str, Any]) -> str:
        return f"""mindmap
  root(({ctx['topic']}))
    课程定位
      课程::{ctx['subject']}
      课程域::{ctx['domain']}
      学习目标::{ctx['goal']}
    课堂笔记
      概念边界::{ctx['primary']}
      方法步骤::{ctx['secondary']}
      结果校验::{ctx['third']}
      迁移记录
    知识图谱节点
      先修节点::{ctx['primary']}
      核心节点::{ctx['topic']}
      方法节点::{ctx['secondary']}
      校验节点::{ctx['third']}
      应用节点::{ctx['transfer']}
    资源文件
      lecture.md
        概念卡
        课堂案例
        自我检查
      practice.md
        基础题
        标准题
        挑战题
        评分量规
      case-project.md
        输入材料
        操作步骤
        验收标准
      quality-checklist.md
        文件验收
        图谱回填
        AI批改闭环
    易错与追练
      定义不清
      条件遗漏
      步骤跳跃
      结论无证据
      AI追练
"""

    @staticmethod
    def _reading_list(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 拓展阅读清单

## 课程内必读
- 当前章节讲义：优先阅读定义、例题和课后练习部分，阅读后补全“概念边界”和“方法步骤”两块课堂笔记。
- 课堂笔记：重点检查 {ctx['primary']}、{ctx['secondary']}、{ctx['third']} 的边界，并标记无法解释的句子。
- 知识图谱：查看该主题的先修节点、方法节点、校验节点和关联资源，确认每个节点至少有一条证据。

## 拓展阅读
- 与 {ctx['topic']} 相邻的概念对比材料，阅读时只记录“差异判断标准”。
- 一个真实应用案例或工程实践说明，阅读时标出输入、方法、输出和限制。
- 一组同类题解析，阅读时关注评分点，而不是只看最终答案。

## 阅读证据模板
| 资料 | 支撑的结论 | 关键页/片段 | 是否可直接引用 |
| --- | --- | --- | --- |
| 课程讲义 | {ctx['primary']} 的定义和边界 | 记录讲义标题、章节名和首句关键词 | 是 |
| 课堂案例 | {ctx['secondary']} 的应用步骤 | 记录案例名称和关键条件 | 是 |
| 练习解析 | {ctx['third']} 的校验方式 | 记录题号、评分点和错因 | 需要复核 |

## 阅读后产物
1. 一张概念对照表：{ctx['topic']}、{ctx['primary']}、{ctx['secondary']} 的差异。
2. 一条图谱边：把“{ctx['topic']} -> {ctx['third']}”写成可解释关系。
3. 一个 AI 追问：要求 AI 根据阅读材料生成一道边界判断题。

## 阅读任务
读完后写下 3 个问题：一个定义问题、一个应用问题、一个易错边界问题。每个问题都要标注来自讲义、图谱还是练习。
"""

    @staticmethod
    def _case_project(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 实操案例

## 任务背景
围绕 {ctx['subject']} 中的 {ctx['topic']}，完成一个 20-30 分钟的小任务，用来验证是否能把概念迁移到真实情境。

## 真实情境
{ctx['scenario']}

## 输入材料
- 一段课程案例或题干。
- {ctx['primary']}、{ctx['secondary']}、{ctx['third']} 的定义和约束条件。
- 一份最终产物模板：条件表、步骤表、结论表。

## 操作步骤
1. 提取已知条件。
2. 判断是否适用 {ctx['topic']}。
3. 写出推理或实现步骤，每一步标注依据。
4. 给出结论并做自检。
5. 写出一个边界情况，说明本方法何时不适用。
6. 将结果回填到课程图谱：新增一个“案例证据”节点，并连接到 {ctx['topic']}。

## 验收标准
- 每一步都有依据。
- 结论能回扣题目目标。
- 能说明一个可能出错的地方。
- 能把错误修改成一版更符合课程术语的答案。
- 能说明该案例如何触发下一份练习或下一次 AI 批改。

## 提交物模板
| 模块 | 内容 | 依据 |
| --- | --- | --- |
| 条件识别 | 写出题干条件 | 引用讲义或案例片段 |
| 方法选择 | 说明为何选择该方法 | 对应 {ctx['primary']} / {ctx['secondary']} |
| 结果验证 | 写出校验或反例 | 对应 {ctx['third']} |
| 图谱回填 | 写出新增节点和关系 | 对应 mind-map.mmd |

## AI 复核提示词
请按“条件识别、方法选择、结果验证、图谱回填、迁移边界”五项检查我的案例提交物。每项给 0-20 分，指出证据是否来自课堂笔记或资源包文件。
"""

    @staticmethod
    def _video_script(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 数字人讲解脚本

大家好，这节课我们用 3 分钟讲清楚 {ctx['topic']}。

第一步，先看它解决什么问题。不要急着背结论，要先知道它适合处理哪类场景：{ctx['scenario']}。

第二步，记住核心判断条件：{ctx['primary']}、{ctx['secondary']}、{ctx['third']}。遇到题目时，先圈出已知条件，再判断是否满足这些条件。

第三步，用一个例子检查理解：{ctx['case']}。如果你能把定义、步骤、证据和结论讲给同学听，说明已经初步掌握。

最后提醒，最常见的错误是把相邻概念混用，或者只写结论不写依据。做题后建议进入 AI 批改模式，让系统根据答案更新掌握度。

镜头提示：在“证据清单”处展示 {ctx['primary']}、{ctx['secondary']}、{ctx['third']} 三张卡片；在案例处展示“条件 -> 方法 -> 结论 -> 校验 -> 图谱回填”的流程。

互动停顿：
1. 请学生用 20 秒说出 {ctx['topic']} 的适用条件。
2. 请学生判断一个反例是否满足 {ctx['third']}。
3. 请学生把错因归类为定义、条件、步骤或证据。

课后动作：把自己的答案复制到 AI 陪练，要求它按评分量规检查：{ctx['rubric']}。
"""

    @staticmethod
    def _quality_checklist(ctx: dict[str, Any]) -> str:
        return f"""# {ctx['topic']} 资源包使用审查清单

课程：{ctx['subject']}
目标：{ctx['goal']}
闭环说明：{ctx['resource_contract']}

## 1. 文件完整性
| 文件 | 必须完成的动作 | 验收标准 |
| --- | --- | --- |
| lecture.md | 阅读概念卡、课堂案例和图谱绑定 | 能用自己的话复述 {ctx['topic']}，并指出证据来源 |
| practice.md | 完成基础题、标准题和挑战题 | 每题都有错因分类、订正答案和下一步追练 |
| mind-map.mmd | 在课程图谱中定位先修、核心、方法、校验、应用节点 | 每个节点都能连接到讲义、练习或案例 |
| reading-list.md | 完成阅读证据模板和阅读后产物 | 每条引用都能说明支撑的结论 |
| case-project.md | 提交条件表、方法表、验证表和图谱回填 | 结果能被 AI 按量规复核 |
| video-script.md | 用 3 分钟复述主题并完成互动停顿 | 能发现至少一个易错边界 |

## 2. 课堂笔记对齐
{ctx['note_blocks']}

## 3. 课程图谱绑定
{ctx['graph_nodes']}

## 4. 个性化学习路径
{ctx['learning_sequence']}

## 5. 质量门槛
{ctx['quality_gate']}

## 6. 继续生成规则
- 如果定义题低于 80 分：重新生成“基础巩固”资源包，只保留讲义、导图和基础练习。
- 如果案例题低于 80 分：重新生成“标准提升”资源包，增加案例项目和评分量规。
- 如果能完成迁移任务：生成“挑战拓展”资源包，加入跨章节对比和开放题。
- 如果引用无法指向文件：回到阅读清单补充资料、片段和支撑结论，再进入 AI 伴学核验。
"""

    @staticmethod
    def _render_docx_bytes(markdown: str, *, title: str) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        def plain(value: str) -> str:
            value = re.sub(r"`([^`\n]+)`", r"\1", value)
            value = re.sub(r"\*\*([^*\n]+)\*\*", r"\1", value)
            value = re.sub(r"__([^_\n]+)__", r"\1", value)
            return value.strip()

        document = Document()
        document.core_properties.title = title
        document.core_properties.author = "智屿"
        docx_font_name = os.getenv("ZHIXI_DOCX_FONT_NAME", "Noto Sans CJK SC")
        section = document.sections[0]
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)

        def set_style_font(style: Any, font_name: str) -> None:
            style.font.name = font_name
            fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
            for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                fonts.attrib.pop(qn(f"w:{theme_attr}"), None)
            for font_attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{font_attr}"), font_name)

        normal = document.styles["Normal"]
        set_style_font(normal, docx_font_name)
        normal.font.size = Pt(10.5)
        for style_name, size, color in (
            ("Title", 20, "132238"),
            ("Heading 1", 16, "183B66"),
            ("Heading 2", 13, "245A8D"),
            ("Heading 3", 11.5, "344054"),
        ):
            style = document.styles[style_name]
            set_style_font(style, docx_font_name)
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)

        title_paragraph = document.add_paragraph(style="Title")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.add_run(title)

        lines = markdown.splitlines()
        index = 0
        first_h1_seen = False
        in_code = False
        code_lines: list[str] = []
        while index < len(lines):
            raw = lines[index].rstrip()
            stripped = raw.strip()
            if stripped.startswith("```"):
                if in_code:
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run("\n".join(code_lines))
                    run.font.name = "Menlo"
                    run.font.size = Pt(9)
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                index += 1
                continue
            if in_code:
                code_lines.append(raw)
                index += 1
                continue
            if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
                rows: list[list[str]] = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    cells = [plain(cell) for cell in lines[index].strip().strip("|").split("|")]
                    if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
                        rows.append(cells)
                    index += 1
                if rows:
                    width = max(len(row) for row in rows)
                    table = document.add_table(rows=len(rows), cols=width)
                    table.style = "Table Grid"
                    for row_index, row in enumerate(rows):
                        for col_index, cell in enumerate(row):
                            table.cell(row_index, col_index).text = cell
                continue
            if not stripped:
                document.add_paragraph()
            elif stripped.startswith("### "):
                document.add_heading(plain(stripped[4:]), level=3)
            elif stripped.startswith("## "):
                document.add_heading(plain(stripped[3:]), level=2)
            elif stripped.startswith("# "):
                heading = plain(stripped[2:])
                if first_h1_seen or re.sub(r"\s+", "", heading) != re.sub(r"\s+", "", title):
                    document.add_heading(heading, level=1)
                first_h1_seen = True
            elif re.match(r"^[-*+]\s+", stripped):
                document.add_paragraph(plain(re.sub(r"^[-*+]\s+", "", stripped)), style="List Bullet")
            elif re.match(r"^\d+[.)]\s+", stripped):
                # Preserve the marker emitted by Markdown.  Word's built-in
                # List Number style otherwise keeps one counter across
                # unrelated sections (for example learning goals 1-2 followed
                # by sources 3-5), which changes the meaning of the document.
                document.add_paragraph(plain(stripped))
            else:
                document.add_paragraph(plain(stripped))
            index += 1
        if code_lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Menlo"
            run.font.size = Pt(9)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_pdf_bytes(markdown: str, *, title: str) -> bytes:
        """Render complete Chinese Markdown into a paginated PDF.

        A configured or commonly installed CJK TrueType/OpenType font is used
        when available. ReportLab's bundled ``STSong-Light`` CID font is the
        deterministic fallback, so missing host fonts never degrade Chinese to
        Latin-1 replacement glyphs.
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import (
                BaseDocTemplate,
                Frame,
                PageTemplate,
                Paragraph,
                Preformatted,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError as exc:  # pragma: no cover - deployment guard
            raise RuntimeError(
                "PDF generation requires reportlab; install code/requirements.txt"
            ) from exc

        font_name = "ZhiXiCJK"
        font_candidates = [
            os.getenv("ZHIXI_CJK_FONT_PATH", ""),
            str(Path.home() / "Library/Fonts/AlibabaPuHuiTi-2-55-Regular.ttf"),
            str(Path.home() / "Library/Fonts/思源黑体-Normal_0.otf"),
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        ]
        for candidate in font_candidates:
            if not candidate or not Path(candidate).is_file():
                continue
            try:
                pdfmetrics.registerFont(TTFont(font_name, candidate))
                break
            except Exception:
                continue
        else:
            font_name = "STSong-Light"
            try:
                pdfmetrics.getFont(font_name)
            except KeyError:
                pdfmetrics.registerFont(UnicodeCIDFont(font_name))

        buffer = BytesIO()
        page_width, page_height = A4
        frame = Frame(19 * mm, 18 * mm, page_width - 38 * mm, page_height - 36 * mm, id="body")

        def draw_page(canvas: Any, doc: Any) -> None:
            canvas.saveState()
            canvas.setFont(font_name, 8.5)
            canvas.setFillColor(colors.HexColor("#667085"))
            canvas.drawRightString(page_width - 19 * mm, 10 * mm, f"第 {doc.page} 页")
            canvas.restoreState()

        doc = BaseDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=19 * mm,
            rightMargin=19 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=title,
            author="智屿",
        )
        doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=draw_page)])
        base = getSampleStyleSheet()
        styles = {
            "title": ParagraphStyle("CJKTitle", parent=base["Title"], fontName=font_name, fontSize=20, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#132238"), spaceAfter=12),
            "h1": ParagraphStyle("CJKH1", parent=base["Heading1"], fontName=font_name, fontSize=16, leading=23, textColor=colors.HexColor("#183B66"), spaceBefore=10, spaceAfter=6),
            "h2": ParagraphStyle("CJKH2", parent=base["Heading2"], fontName=font_name, fontSize=13, leading=20, textColor=colors.HexColor("#245A8D"), spaceBefore=8, spaceAfter=5),
            "h3": ParagraphStyle("CJKH3", parent=base["Heading3"], fontName=font_name, fontSize=11.5, leading=18, textColor=colors.HexColor("#344054"), spaceBefore=6, spaceAfter=4),
            "body": ParagraphStyle("CJKBody", parent=base["BodyText"], fontName=font_name, fontSize=10.5, leading=18, textColor=colors.HexColor("#1D2939"), spaceAfter=5, wordWrap="CJK"),
            "list": ParagraphStyle("CJKList", parent=base["BodyText"], fontName=font_name, fontSize=10.2, leading=17, leftIndent=12, firstLineIndent=-8, textColor=colors.HexColor("#1D2939"), spaceAfter=3, wordWrap="CJK"),
            "code": ParagraphStyle("CJKCode", parent=base["Code"], fontName=font_name, fontSize=8.5, leading=13, leftIndent=8, rightIndent=8, backColor=colors.HexColor("#F2F4F7"), borderColor=colors.HexColor("#D0D5DD"), borderWidth=0.5, borderPadding=7, spaceBefore=4, spaceAfter=7, wordWrap="CJK"),
        }

        def escape(value: str) -> str:
            return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        def inline_markup(value: str) -> str:
            """Translate the small inline Markdown subset emitted by resource agents.

            ReportLab ``Paragraph`` accepts a constrained XML-like markup. Escape
            user/model text first, then add only formatting tags that we control.
            This keeps literal HTML inert while avoiding visible ``**`` and
            backticks in the generated teaching PDF.
            """
            rendered = escape(value)
            # Inline code may contain Chinese course titles or source labels.
            # Courier has no CJK glyphs and previously rendered those labels as
            # black squares in otherwise valid PDFs.  Keep inline code on the
            # registered CJK-capable font; fenced code blocks still use their
            # dedicated code style and remain visually separated.
            rendered = re.sub(
                r"`([^`\n]+)`",
                lambda match: f'<font name="{font_name}">{match.group(1)}</font>',
                rendered,
            )
            rendered = re.sub(r"\*\*([^*\n]+)\*\*", r"<b>\1</b>", rendered)
            rendered = re.sub(r"__([^_\n]+)__", r"<b>\1</b>", rendered)
            rendered = re.sub(
                r"(?<!\*)\*([^*\n]+)\*(?!\*)",
                r"<i>\1</i>",
                rendered,
            )
            return rendered

        story: list[Any] = [Paragraph(escape(title), styles["title"]), Spacer(1, 3 * mm)]
        lines = markdown.splitlines()
        index = 0
        first_h1_seen = False
        in_code = False
        code_lines: list[str] = []
        while index < len(lines):
            raw = lines[index].rstrip()
            if raw.strip().startswith("```"):
                if in_code:
                    story.append(Preformatted("\n".join(code_lines), styles["code"]))
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                index += 1
                continue
            if in_code:
                code_lines.append(raw)
                index += 1
                continue
            if raw.strip().startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
                table_rows: list[list[str]] = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                    if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
                        table_rows.append([Paragraph(inline_markup(cell), styles["body"]) for cell in cells])
                    index += 1
                if table_rows:
                    column_count = max(len(row) for row in table_rows)
                    for row in table_rows:
                        row.extend([""] * (column_count - len(row)))
                    table = Table(table_rows, colWidths=[frame._width / column_count] * column_count, repeatRows=1)
                    table.setStyle(TableStyle([
                        ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 8.8),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#183B66")),
                        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#B8C4CE")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 3 * mm))
                continue
            stripped = raw.strip()
            if not stripped:
                story.append(Spacer(1, 2.5 * mm))
            elif stripped.startswith("### "):
                story.append(Paragraph(inline_markup(stripped[4:]), styles["h3"]))
            elif stripped.startswith("## "):
                story.append(Paragraph(inline_markup(stripped[3:]), styles["h2"]))
            elif stripped.startswith("# "):
                heading = stripped[2:].strip()
                normalized_heading = re.sub(r"[\s\-_—–:：]+", "", heading).lower()
                normalized_title = re.sub(r"[\s\-_—–:：]+", "", title).lower()
                if first_h1_seen or normalized_heading != normalized_title:
                    story.append(Paragraph(inline_markup(heading), styles["h1"]))
                first_h1_seen = True
            elif re.match(r"^[-*+]\s+", stripped):
                story.append(Paragraph("• " + inline_markup(re.sub(r"^[-*+]\s+", "", stripped)), styles["list"]))
            elif re.match(r"^\d+[.)]\s+", stripped):
                marker, content = stripped.split(maxsplit=1)
                story.append(Paragraph(escape(marker) + " " + inline_markup(content), styles["list"]))
            else:
                story.append(Paragraph(inline_markup(stripped), styles["body"]))
            index += 1
        if code_lines:
            story.append(Preformatted("\n".join(code_lines), styles["code"]))
        doc.build(story)
        return buffer.getvalue()


resource_generation_service = ResourceGenerationService()
