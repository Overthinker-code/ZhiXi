from __future__ import annotations

import os
from datetime import date
from pathlib import Path

_EARLY_ROOT = Path(__file__).resolve().parents[1]
os.environ.setdefault(
    "MPLCONFIGDIR", str(_EARLY_ROOT / "code" / "backend" / "uploads" / ".matplotlib")
)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
ASSETS = OUT / "overview_design_assets"
DOCX = OUT / "智屿概要设计说明书_V1.0.docx"
BLUE = "315EFB"
DEEP_BLUE = "173A8F"
LIGHT_BLUE = "EAF0FF"
TEAL = "12B8A6"
GRAY = "667085"


def configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "DejaVu Sans"
    ]
    plt.rcParams["axes.unicode_minus"] = False


def box(ax, x, y, w, h, title, subtitle="", color="#315EFB", fill="#EEF2FF", size=12):
    patch = FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.018",
        linewidth=1.5, edgecolor=color, facecolor=fill
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h * 0.62, title, ha="center", va="center",
            fontsize=size, weight="bold", color="#172B4D")
    if subtitle:
        ax.text(x + w / 2, y + h * 0.30, subtitle, ha="center", va="center",
                fontsize=size - 3, color="#52607A", wrap=True)
    return patch


def arrow(ax, start, end, color="#7181A6", rad=0.0):
    ax.add_patch(FancyArrowPatch(
        start, end, arrowstyle="-|>", mutation_scale=12, linewidth=1.4,
        color=color, connectionstyle=f"arc3,rad={rad}"
    ))


def save_architecture(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 8), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .96, "智屿系统总体技术架构", ha="center", fontsize=20, weight="bold", color="#173A8F")
    layers = [
        (.79, "交互表现层", "Vue 3 · Vite · Pinia · Arco Design · SSE/Markdown/多模态卡片", "#EEF2FF"),
        (.60, "业务服务层", "AI伴学 · 课程中心 · 资料中心 · 错题本 · 学情档案 · 数字人", "#EAF8F6"),
        (.41, "多智能体编排层", "Supervisor · Tutor · Retrieval · Planner · Profile · Grading · Safety Review", "#FFF4E8"),
        (.22, "能力与数据层", "DeepSeek/Qwen/Seedance · RAG/Chroma · PostgreSQL · Redis · 文件对象", "#F3EEFF"),
    ]
    for y, title, subtitle, fill in layers:
        box(ax, .08, y, .70, .13, title, subtitle, fill=fill, size=13)
    for y in [.79, .60, .41]:
        arrow(ax, (.43, y), (.43, y-.055))
    box(ax, .82, .66, .15, .17, "外部模型服务", "百炼 Qwen/VL/万相\nSeedance 视频\nEdge TTS", color="#8B5CF6", fill="#F7F2FF", size=12)
    box(ax, .82, .39, .15, .17, "外部资源", "国家智慧教育\n中国大学MOOC\n学堂在线/B站", color="#12B8A6", fill="#ECFBF8", size=12)
    box(ax, .82, .13, .15, .17, "运维支撑", "健康检查·日志\n限流·重试·迁移\n安全与质量审查", color="#F59E0B", fill="#FFF8E8", size=12)
    arrow(ax, (.78, .68), (.82, .745)); arrow(ax, (.78, .49), (.82, .475)); arrow(ax, (.78, .28), (.82, .215))
    ax.text(.08, .08, "设计原则：画像驱动、证据优先、能力路由、异步生成、可回退、全链路可追踪", fontsize=11, color="#52607A")
    fig.tight_layout(); fig.savefig(path, bbox_inches="tight", facecolor="white"); plt.close(fig)


def save_agents(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 8), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .95, "多智能体协同与个人学习闭环", ha="center", fontsize=20, weight="bold", color="#173A8F")
    box(ax, .39, .70, .22, .13, "Supervisor Agent", "意图识别·任务拆解·结果汇总", color="#315EFB", fill="#EAF0FF", size=13)
    nodes = [
        (.05, .49, "Retrieval Agent", "课程知识库与上传资料证据"),
        (.28, .49, "Tutor Agent", "多模态理解与分步辅导"),
        (.51, .49, "Resource Agent", "PPT/题目/图表/视频生成"),
        (.74, .49, "Safety Review", "事实、来源与内容安全审查"),
        (.12, .23, "Grading Agent", "批改·错因·掌握度证据"),
        (.39, .23, "Profile Agent", "八维画像随学随新"),
        (.66, .23, "Planner Agent", "路径规划·任务重排·推送"),
    ]
    for x, y, t, s in nodes:
        box(ax, x, y, .20, .13, t, s, color="#12B8A6" if y < .3 else "#8B5CF6", fill="#F7FAFF", size=11)
        arrow(ax, (.50, .70), (x+.10, y+.13))
    arrow(ax, (.22, .23), (.39, .295)); arrow(ax, (.59, .295), (.66, .295)); arrow(ax, (.76, .23), (.59, .23), rad=-.25)
    box(ax, .33, .04, .34, .10, "学习闭环状态", "行为/练习 → 证据 → 画像 → 路径/资源 → 新行为", color="#F59E0B", fill="#FFF8E8", size=12)
    arrow(ax, (.49, .23), (.49, .14)); arrow(ax, (.67, .09), (.86, .49), rad=-.30)
    fig.tight_layout(); fig.savefig(path, bbox_inches="tight", facecolor="white"); plt.close(fig)


def save_resource_flow(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 6.2), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .92, "画像驱动的个性化资源生成流程", ha="center", fontsize=19, weight="bold", color="#173A8F")
    stages = [
        (.03, "需求输入", "文本/图片/文件\n课程与目标"),
        (.19, "画像融合", "掌握度/偏好\n薄弱点/节奏"),
        (.35, "任务编排", "类型路由\nAgent协作"),
        (.51, "内容生成", "文档/PPT/题目\n图表/视频/代码"),
        (.67, "质量审查", "RAG证据\n结构/安全检查"),
        (.83, "交付反馈", "卡片/下载/进度\n收藏/评价/练习"),
    ]
    colors = ["#EAF0FF", "#ECFBF8", "#FFF4E8", "#F4EEFF", "#FFF8E8", "#EAF8F6"]
    for i, (x, t, s) in enumerate(stages):
        box(ax, x, .48, .135, .22, t, s, fill=colors[i], size=12)
        if i < len(stages)-1: arrow(ax, (x+.135, .59), (stages[i+1][0], .59))
    ax.text(.5, .31, "失败回退链", ha="center", fontsize=12, weight="bold", color="#B54708")
    box(ax, .22, .10, .56, .13, "视频：Seedance → Qwen 分镜 + Manim + 中文 TTS　｜　图表：确定性绘图 → Mermaid　｜　模型：主模型 → 备用模型", color="#F59E0B", fill="#FFF9ED", size=11)
    arrow(ax, (.74, .48), (.70, .23), color="#F59E0B", rad=.18)
    fig.tight_layout(); fig.savefig(path, bbox_inches="tight", facecolor="white"); plt.close(fig)


def save_deployment(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 7), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .94, "部署拓扑与运行依赖", ha="center", fontsize=20, weight="bold", color="#173A8F")
    box(ax, .05, .57, .18, .18, "浏览器客户端", "Vue SPA\nlocalhost:5174 / HTTPS", fill="#EAF0FF", size=13)
    box(ax, .31, .57, .20, .18, "FastAPI 应用", "REST + SSE + WebSocket\nUvicorn :8001", fill="#ECFBF8", size=13)
    box(ax, .60, .65, .16, .14, "PostgreSQL", "业务与画像数据", color="#336791", fill="#EEF5FA", size=12)
    box(ax, .60, .45, .16, .14, "Redis / Celery", "任务队列与结果状态", color="#D82C20", fill="#FFF1F0", size=12)
    box(ax, .60, .25, .16, .14, "Chroma / 文件存储", "向量、附件、生成物", color="#8B5CF6", fill="#F7F2FF", size=12)
    box(ax, .82, .57, .15, .18, "模型与资源服务", "DeepSeek/Qwen\nSeedance/Edge TTS\n公开课程目录", color="#F59E0B", fill="#FFF8E8", size=12)
    arrow(ax, (.23, .66), (.31, .66)); arrow(ax, (.51, .66), (.60, .72)); arrow(ax, (.51, .63), (.60, .52)); arrow(ax, (.51, .60), (.60, .32)); arrow(ax, (.76, .66), (.82, .66))
    box(ax, .22, .07, .56, .10, "统一配置与安全边界：.env 密钥隔离 · JWT 鉴权 · CORS/TrustedHost · 上传大小限制 · 并发预算 · 审计日志", color="#344054", fill="#F5F7FA", size=11)
    fig.tight_layout(); fig.savefig(path, bbox_inches="tight", facecolor="white"); plt.close(fig)


def save_data_model(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 8), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .95, "核心数据对象关系（概要级）", ha="center", fontsize=20, weight="bold", color="#173A8F")
    entities = {
        "用户与画像": (.05, .64, "user\nstudent_profile\nuser_memory_profile"),
        "会话与任务": (.31, .64, "chatthread / conversation_message\nagent_task / chatartifact"),
        "课程知识": (.57, .64, "course\ncourse_knowledge_node / edge\nresource"),
        "练习评估": (.05, .30, "question / quiz_attempt\nwrong_question\nlearning_evidence"),
        "路径与推送": (.37, .30, "learning_path / learning_task\npersonalized_resource_recommendation"),
        "资源生成": (.69, .30, "resource_generation_run / step\ngenerated_resource_package\nexternal_resource"),
    }
    for i, (title, (x, y, sub)) in enumerate(entities.items()):
        box(ax, x, y, .25, .18, title, sub, color="#315EFB" if y>.5 else "#12B8A6", fill="#F8FAFF", size=12)
    arrow(ax, (.30,.73),(.31,.73)); arrow(ax, (.56,.73),(.57,.73)); arrow(ax, (.17,.64),(.17,.48)); arrow(ax, (.44,.64),(.46,.48)); arrow(ax, (.69,.64),(.78,.48));
    arrow(ax, (.30,.39),(.37,.39)); arrow(ax, (.62,.39),(.69,.39)); arrow(ax, (.57,.67),(.54,.48), rad=.15)
    ax.text(.5,.12,"主线：用户身份 → 学习行为/对话 → 证据与画像 → 路径/推荐 → 资源使用与练习 → 新证据",ha="center",fontsize=11,color="#52607A")
    fig.tight_layout(); fig.savefig(path, bbox_inches="tight", facecolor="white"); plt.close(fig)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold=False, color=None, size=9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], BLUE)
        set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", font_size)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, None, font_size)
            if r_idx % 2:
                set_cell_shading(cells[i], "F7F9FC")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    paragraph.add_run(" 页")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar"); fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "目录将在 WPS/Word 中打开时自动更新"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GRAY)


def add_picture(doc, path, caption, width=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def page_break(doc):
    doc.add_page_break()


def build_document() -> None:
    OUT.mkdir(exist_ok=True); ASSETS.mkdir(exist_ok=True)
    configure_matplotlib()
    figures = {
        "architecture": ASSETS / "01_system_architecture.png",
        "agents": ASSETS / "02_multi_agent_loop.png",
        "resource": ASSETS / "03_resource_generation_flow.png",
        "deployment": ASSETS / "04_deployment_topology.png",
        "data": ASSETS / "05_core_data_model.png",
    }
    save_architecture(figures["architecture"]); save_agents(figures["agents"])
    save_resource_flow(figures["resource"]); save_deployment(figures["deployment"])
    save_data_model(figures["data"])

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3); sec.bottom_margin = Cm(2.2); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.3)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"; normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for name, size, color in [("Title", 28, DEEP_BLUE), ("Heading 1", 18, DEEP_BLUE), ("Heading 2", 14, BLUE), ("Heading 3", 11.5, "344054")]:
        st = doc.styles[name]; st.font.name = "Microsoft YaHei"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)

    header = sec.header.paragraphs[0]
    header.text = "智屿——基于大模型的个性化资源生成与学习多智能体系统　概要设计说明书"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(sec.footer.paragraphs[0])

    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(70)
    r = p.add_run("第十五届“中国软件杯”大学生软件设计大赛\nA3 赛题参赛项目")
    r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(BLUE); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(40)
    r = p.add_run("智　屿"); r.font.size = Pt(38); r.bold = True; r.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于大模型的个性化资源生成与学习多智能体系统"); r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string("344054")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(42)
    r = p.add_run("概要设计说明书"); r.font.size = Pt(30); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(80)
    p.add_run("文档版本：V1.0\n编制单位：智屿项目团队\n完成日期：2026 年 7 月 20 日").font.size = Pt(12)
    page_break(doc)

    heading(doc, "文档控制", 1)
    add_table(doc, ["项目", "内容"], [
        ("文档名称", "智屿概要设计说明书"), ("赛题", "A3—基于大模型的个性化资源生成与学习多智能体系统开发"),
        ("出题企业", "科大讯飞股份有限公司"), ("版本/状态", "V1.0 / 提交评审版"),
        ("编制/审核", "智屿项目团队 / 待团队负责人签署"), ("保密级别", "参赛材料；密钥、口令和个人敏感信息不进入文档"),
    ], [4, 12], 10)
    heading(doc, "修改情况记录", 2)
    add_table(doc, ["版本", "日期", "修改人", "修改说明", "批准人"], [
        ("V0.1", "2026-07-18", "智屿项目团队", "形成系统概要设计初稿", "—"),
        ("V0.9", "2026-07-19", "智屿项目团队", "补充多模态、课程知识库、学习画像与资源生成设计", "—"),
        ("V1.0", "2026-07-20", "智屿项目团队", "按赛题要求和现有实现完成提交评审版", "待签署"),
    ], font_size=9)
    heading(doc, "目录", 1); add_toc(doc.add_paragraph())
    body(doc, "提示：在 WPS/Word 中右击目录并选择“更新域”，即可生成带页码的正式目录。")
    page_break(doc)

    heading(doc, "1 引言", 1)
    heading(doc, "1.1 编写目的", 2)
    body(doc, "本文档说明智屿系统的总体技术方案、模块划分、接口、运行机制、数据结构和异常处理策略，为后续详细设计、编码实现、系统测试、部署验收和比赛答辩提供统一依据。读者包括项目开发与测试成员、指导教师、竞赛评审专家以及后续运维人员。")
    body(doc, "文档重点回答三个问题：系统如何把学生画像转化为个性化学习行动；不同角色智能体如何协作完成证据检索、内容生成、批改评估和安全审查；系统如何以可追踪、可回退的工程机制稳定交付文档、PPT、题库、图表、视频与代码案例等多模态资源。")
    heading(doc, "1.2 范围", 2)
    add_table(doc, ["范围项", "说明"], [
        ("软件名称", "智屿——基于大模型的个性化资源生成与学习多智能体系统"),
        ("任务来源", "第十五届中国软件杯 A3 赛题"),
        ("开发者", "智屿项目团队"), ("目标用户", "高校学生为主，兼顾教师、课程建设者和系统管理员"),
        ("运行单位", "参赛演示环境及具备常规浏览器、Python、Node.js、PostgreSQL 的高校教学环境"),
        ("本期课程", "以“软件工程导论”为完整专业课程样板，并保留多课程扩展能力"),
        ("不在范围", "学校教务收费、排课、学籍审批等行政业务；通用社交平台内容抓取；替代教师作出高风险学术评价"),
    ], font_size=9)
    heading(doc, "1.3 定义", 2)
    add_table(doc, ["术语/缩写", "定义"], [
        ("LLM", "Large Language Model，大语言模型。"), ("Agent", "具有角色目标、上下文、工具和状态的智能体执行单元。"),
        ("Supervisor", "协作主管智能体，负责任务识别、分解、路由和最终汇总。"), ("RAG", "Retrieval-Augmented Generation，检索增强生成。"),
        ("SSE", "Server-Sent Events，服务器向浏览器持续推送生成状态和文本片段。"), ("动态画像", "由对话、练习、资源行为等证据持续更新的多维学生模型。"),
        ("学习证据", "能够支持掌握度、偏好或风险判断的可追溯行为与评测记录。"), ("资源包", "一次生成任务产生的文档、PPT、练习、图表、视频等成果集合。"),
        ("能力路由", "依据用户意图和输入模态选择文本、视觉、绘图、视频或确定性工具链。"), ("降级/回退", "主服务不可用时切换到备用模型或本地确定性工具，维持核心可用性。"),
    ], font_size=9)
    heading(doc, "1.4 参考资料", 2)
    refs = [
        "[1] 第十五届中国软件杯 A3 赛题《基于大模型的个性化资源生成与学习多智能体系统开发》，软件杯大赛官网，2026-04-01，https://www.cnsoftbei.com/content-3-1286-1.html。",
        "[2] 《概要设计说明书编写规范》，用户提供的国家标准软件开发规范模板。",
        "[3] 智屿项目源代码、数据库迁移、测试用例与运行配置，ZhiYu-main 仓库，核对日期 2026-07-20。",
        "[4] 《软件工程导论课程知识库》及原始资料集：课程大纲、13章笔记、13章习题、PPT、教材PDF和结构化切片。",
        "[5] Vue、FastAPI、LangGraph、PostgreSQL、Chroma、Mermaid、Manim、python-pptx 等项目官方文档与许可证。",
    ]
    for item in refs: bullet(doc, item)
    heading(doc, "1.5 赛题约束与设计原则", 2)
    body(doc, "赛题要求构建不少于6维的动态学生画像、采用多智能体架构、协同生成至少5类个性化资源、规划动态学习路径并精确推送资源；同时要求至少提供一门完整高校专业课程知识库，具备防幻觉、内容安全、合理响应时间和现代AI产品交互。智屿将这些约束转化为以下设计原则：")
    for text in [
        "画像驱动：所有路径、推荐和生成任务均可读取学生画像，但仅使用有证据支撑的维度。",
        "证据优先：优先引用课程知识库和上传资料；证据不足时明确边界，不伪造来源。",
        "多智能体可解释协作：保留任务、步骤、状态和产物记录，前端展示进度而不暴露内部思维链。",
        "多模态与确定性工具结合：自然语言生成交给模型，严谨图表、PPT排版和教学动画交给可验证工具链。",
        "安全与最小权限：密钥服务端保存，上传受类型和大小限制，访问使用JWT与对象归属校验。",
        "可恢复：异步任务支持查询、取消、恢复和回退；外部API故障不影响已有资料和基础学习功能。",
    ]: bullet(doc, text)

    heading(doc, "2 总体设计", 1)
    heading(doc, "2.1 需求规定", 2)
    heading(doc, "2.1.1 功能需求", 3)
    functional = [
        ("FR-01", "动态学习画像", "从自然语言对话、目标、练习和资源行为中更新知识掌握、认知风格、学习偏好、学习行为、动机、目标、问题解决能力、知识网络状态等八维画像。", "核心"),
        ("FR-02", "多智能体协作", "主管智能体按意图调度检索、辅导、画像、规划、批改、联网研究、安全审查等角色，并统一输出。", "核心"),
        ("FR-03", "多模态输入", "支持输入文字、粘贴或拖拽图片与文件；视觉模型联合课程上下文识别题目、手写内容和版式。", "核心"),
        ("FR-04", "个性化资源生成", "生成讲解文档、PPT、练习题、思维导图/流程图、科学图表、视频/动画和代码案例，成果卡片化展示并可下载。", "核心"),
        ("FR-05", "学习路径与推送", "根据画像、掌握度和前置关系生成阶段目标、任务顺序，推荐站内资源及经约束的外部学习资源。", "核心"),
        ("FR-06", "课程知识库", "维护软件工程导论完整资料集，支持切片、向量检索、题库、PPT播放、知识图谱和来源追踪。", "核心"),
        ("FR-07", "练习与错题闭环", "答题、自动批改、错因定位、按学科归档错题，并根据错题生成针对性练习。", "核心"),
        ("FR-08", "智能辅导", "结合课程证据、图片和学习画像提供分步讲解、图解和短视频答疑。", "加分"),
        ("FR-09", "学习效果评估", "聚合学习证据形成能力趋势、风险提示、资源使用分布和下一步建议，反向更新路径。", "加分"),
        ("FR-10", "小智悬浮助手", "在资料库和AI聊天等学习场景提供提醒与快捷交互，在做题界面禁用，避免干扰测评。", "创新"),
    ]
    add_table(doc, ["编号", "功能", "概要说明", "级别"], functional, font_size=8)
    heading(doc, "2.1.2 输入、处理与输出", 3)
    add_table(doc, ["类别", "主要输入", "关键处理", "输出"], [
        ("对话辅导", "问题、课程、会话历史、画像", "意图路由→RAG→专业Agent→安全审查", "流式文字、引用、任务轨迹、建议"),
        ("图片/文件", "PNG/JPEG/PDF/DOCX/PPTX等", "校验→存储→OCR/视觉理解/解析→证据化", "内容摘要、题目识别、解答、入库索引"),
        ("资源生成", "主题、类型、难度、画像、知识证据", "异步编排→多Agent生成→质量门禁→打包", "文档、PPTX、题库、PNG、MP4、代码"),
        ("学习评估", "答题、错题、访问、反馈、对话", "证据聚合→画像更新→路径重排", "掌握度、趋势、风险、学习任务"),
        ("外部推荐", "画像薄弱点、课程主题", "限定来源发现→元数据校验→排序去重", "B站/中国大学MOOC/学堂在线等链接"),
    ], font_size=8)
    heading(doc, "2.1.3 非功能需求", 3)
    add_table(doc, ["编号", "属性", "设计指标/约束"], [
        ("NFR-01", "易用性", "界面简洁一致；支持流式输出、Markdown、公式、资源卡片和生成进度；关键操作有空态、加载态和失败态。"),
        ("NFR-02", "性能", "普通查询优先流式首字反馈；同步AI每用户并发1、SSE每用户并发2；长任务异步执行并可查询进度。"),
        ("NFR-03", "可靠性", "外部模型超时、重试和回退；资源任务具备步骤状态、租约、取消/恢复与幂等。"),
        ("NFR-04", "安全", "JWT鉴权、密钥不下发、上传白名单与大小限制、TrustedHost/CORS、访问归属检查、内容安全审核。"),
        ("NFR-05", "可维护性", "前后端分层、Service封装、统一配置、Alembic迁移、类型模型、自动测试与健康检查。"),
        ("NFR-06", "可扩展性", "课程、Agent、模型供应商、资源类型和外部来源均采用注册/适配器式扩展。"),
        ("NFR-07", "防幻觉", "课程证据优先、来源边界、结构化输出校验、SafetyReview复核；无证据时降低确定性。"),
        ("NFR-08", "合规", "标注开源依赖、AI服务和AI Coding使用边界；不在文档、代码库和前端暴露API密钥。"),
    ], font_size=8)

    heading(doc, "2.2 运行环境", 2)
    add_table(doc, ["层次", "推荐环境", "说明"], [
        ("客户端", "Windows 10/11、macOS或Linux；Edge/Chrome最新版；≥1366×768", "支持JavaScript、SSE、文件拖拽和本地下载。"),
        ("前端", "Node.js 18+，Vue 3、TypeScript、Vite", "开发端口5174；生产环境输出静态资源。"),
        ("后端", "Python 3.11+，FastAPI、Uvicorn", "REST/SSE/WebSocket；默认端口8001。"),
        ("数据库", "PostgreSQL 14+", "结构化业务、画像、任务、学习证据与推荐记录。"),
        ("检索与队列", "Chroma；Redis/Celery（生产推荐）", "向量检索、异步任务、状态与结果。"),
        ("多媒体", "Pillow、Matplotlib、python-pptx、Manim、PyAV、Edge TTS", "确定性图表、PPT、动画、音视频合成。"),
        ("外部AI", "DeepSeek兼容接口；阿里云百炼Qwen/VL/万相；Seedance", "通过环境变量配置，配置缺失时按能力回退。"),
        ("硬件", "开发机≥4核CPU、16GB内存、20GB可用空间；本地Manim建议更高CPU", "云端模型为主，不强制独立GPU。"),
    ], font_size=8)
    add_picture(doc, figures["deployment"], "图2-1 部署拓扑与运行依赖")

    heading(doc, "2.3 基本设计概念和处理流程", 2)
    body(doc, "系统采用前后端分离、业务服务分层和多智能体状态图编排。浏览器只处理交互与可视化；FastAPI负责身份、业务规则、任务状态和能力路由；多智能体层根据意图调用知识检索、模型与确定性工具；PostgreSQL、Chroma和文件目录分别保存结构化状态、向量证据和二进制产物。")
    add_picture(doc, figures["architecture"], "图2-2 系统总体技术架构")
    body(doc, "一次学习请求首先形成统一上下文，包括用户身份、课程、会话、附件、当前画像和允许使用的工具。Supervisor根据规则与模型判断选择专业Agent；Retrieval Agent先收集课程或文件证据；Tutor/Resource/Grading等Agent完成专业任务；Safety Review检查事实边界、引用和安全；最终由主管统一成面向学生的流式回答。答题、阅读和反馈进一步沉淀为学习证据，驱动画像与路径更新。")
    add_picture(doc, figures["agents"], "图2-3 多智能体协同与个人学习闭环")
    heading(doc, "2.3.1 个性化资源生成流程", 3)
    body(doc, "资源生成不是把提示词直接交给单一模型，而是由需求解析、画像融合、课程证据检索、结构规划、能力生成、质量审查和交付反馈组成。文字解释使用通用文本模型；图片题目使用Qwen-VL；PPT由Qwen输出结构化JSON后由python-pptx排版；严谨数据图由Matplotlib绘制；知识图可由Mermaid渲染；视频优先Seedance，额度或服务不可用时回退到Qwen分镜、Manim本地渲染和中文TTS合成。")
    add_picture(doc, figures["resource"], "图2-4 画像驱动的个性化资源生成流程")

    heading(doc, "2.4 结构", 2)
    heading(doc, "2.4.1 功能模块划分", 3)
    modules = [
        ("M01", "身份与会话", "登录、JWT、会话、消息、附件、历史与反馈", "login/chat_threads/ai_chat"),
        ("M02", "多智能体编排", "主管路由、专业Agent、工具权限、状态流转、最终汇总", "app.ai.chat_engine/chat_runtime"),
        ("M03", "课程与知识库", "课程、资料、RAG、知识图谱、PPT播放、知识节点动作", "rag/resource_hub/knowledge_graph"),
        ("M04", "多模态理解", "图片和文件上传、类型识别、视觉理解、文档检索", "vision_client/document_processor"),
        ("M05", "资源工坊", "个性化资源包、生成任务、步骤进度、成果预览与下载", "resource_generation/resource_workshop"),
        ("M06", "练习与错题", "测验、批改、尝试记录、错题分类、错题衍生练习", "quiz_service/quiz endpoints"),
        ("M07", "画像与学情", "八维数字分身、能力趋势、证据可信度、分析建议", "student_profile_agent/learning_report"),
        ("M08", "路径与推荐", "学习路径、当前任务、站内外资源推荐、反馈与排序", "learning_path/resource_recommendation"),
        ("M09", "小智与数字人", "悬浮提醒、语音、文本/PPT转视频、场景控制", "digital_human/classroom_assistant"),
        ("M10", "安全与运维", "安全过滤、质量评估、限流、健康检查、日志和迁移", "content_safety/health/http_security"),
    ]
    add_table(doc, ["标识", "模块", "职责", "主要实现"], modules, font_size=8)
    heading(doc, "2.4.2 智能体角色结构", 3)
    agents = [
        ("Supervisor", "任务拆解、选择Agent、控制轮次、汇总结果"), ("Code Tutor", "代码报错、调试、原理与最小修复"),
        ("Knowledge Mentor", "跨学科概念讲解、例题与迁移"), ("Planner", "目标拆解、计划重排、里程碑与下一步"),
        ("Analyst", "行为分析、风险识别和数据解释"), ("Doc Researcher", "上传文档检索、引用与问答"),
        ("Quiz Master", "苏格拉底式测验与逐题引导"), ("Profile Agent", "画像信号提取、版本更新与协同输出"),
        ("Retrieval Agent", "课程知识库和文件证据整理"), ("Web Research Agent", "需要时检索最新官方或公开资料"),
        ("Tutor Agent", "图像+文本联合理解与分步答疑"), ("Grading Agent", "评分、错因、订正与掌握度建议"),
        ("Safety Review Agent", "事实、来源、敏感内容和不当建议审查"),
    ]
    add_table(doc, ["角色", "主要职责"], agents, [5, 11], 8)
    heading(doc, "2.4.3 课程知识库结构", 3)
    body(doc, "软件工程导论作为本期完整专业课程样板。原始资料集共127个文件、约497 MB，包括86份PDF、28份Markdown、5份DOCX、2份XLSX、2份JSONL及其他结构化材料；内容覆盖课程大纲、13章笔记、13章课后习题、课堂PPT、教材和RAG切片。入库使用稳定file_id，重复执行先清理旧版本，避免重复切片；大型PPT和教材先写入资料索引，按需解析以控制启动和检索成本。")
    add_table(doc, ["层次", "内容", "用途"], [
        ("原始层", "PPT、教材PDF、笔记、题目、表格", "保留原件、下载、PPT播放和人工追溯"),
        ("结构层", "课程、章节、知识节点、先修边、题目", "导航、图谱、练习和路径规划"),
        ("检索层", "chunks.jsonl、向量、关键词索引", "RAG召回与引用证据"),
        ("行为层", "节点访问、学习证据、答题与反馈", "画像更新和推荐排序"),
    ], font_size=8)

    heading(doc, "2.5 功能需求与程序的关系", 2)
    matrix_headers = ["功能", "聊天编排", "知识/RAG", "资源工坊", "练习错题", "画像报告", "路径推荐", "数字人"]
    matrix_rows = [
        ("动态画像", "√", "○", "○", "√", "√", "√", "○"), ("多智能体", "√", "√", "√", "√", "√", "√", "○"),
        ("多模态输入", "√", "√", "√", "", "○", "", "○"), ("资源生成", "√", "√", "√", "○", "○", "√", "√"),
        ("路径与推送", "○", "√", "○", "√", "√", "√", "○"), ("课程知识库", "○", "√", "○", "√", "○", "√", ""),
        ("错题闭环", "○", "○", "○", "√", "√", "√", ""), ("智能辅导", "√", "√", "○", "√", "○", "○", "√"),
        ("学习评估", "○", "○", "", "√", "√", "√", "○"),
    ]
    add_table(doc, matrix_headers, matrix_rows, font_size=7)
    body(doc, "注：√表示主要实现模块，○表示协同或数据消费模块。矩阵体现赛题功能不是孤立页面，而是由对话、证据、画像、路径和资源共同闭环。")

    heading(doc, "2.6 人工处理过程", 2)
    add_table(doc, ["环节", "人工职责", "系统支持"], [
        ("课程资料初始化", "教师确认课程版权、章节边界与题目答案", "预览、稳定ID入库、重复检测和来源字段"),
        ("高风险内容复核", "教师复核关键答案、评分标准和外部来源适用性", "证据引用、质量状态和反馈入口"),
        ("模型/API配置", "管理员申请密钥、充值额度并填写服务端环境变量", "配置探测、健康检查和错误提示"),
        ("画像纠偏", "学生或教师发现画像误判时确认或更正", "画像版本、更新事件和可解释证据"),
        ("比赛材料发布", "团队检查开源协议、敏感信息和演示数据", "示例环境文件、密钥隔离和提交清单"),
    ], font_size=8)
    heading(doc, "2.7 尚未解决的问题", 2)
    add_table(doc, ["问题", "影响", "当前策略", "后续计划"], [
        ("外部视频服务按量计费", "额度不足时无法调用Seedance", "回退Qwen+Manim+中文TTS", "比赛演示前完成额度和连通性检查"),
        ("手写复杂公式识别存在误差", "图片批改可靠性下降", "视觉模型输出不确定性并要求补充题干", "增加OCR交叉校验与人工确认"),
        ("大型PDF/PPT全量解析成本高", "入库和检索变慢", "先索引后按需解析", "建立后台增量解析与缓存"),
        ("外部平台链接有效性变化", "推荐可能失效", "限定来源、验证时间和反馈淘汰", "周期巡检和来源健康评分"),
        ("画像早期证据稀疏", "个性化置信度较低", "显示待积累和证据可信度", "通过诊断测验快速冷启动"),
    ], font_size=8)

    heading(doc, "3 接口设计", 1)
    heading(doc, "3.1 用户接口", 2)
    add_table(doc, ["界面", "主要操作", "系统反馈"], [
        ("首页/AI伴学", "输入问题、选择课程、粘贴/拖拽附件、停止/重试", "SSE流式回答、引用、生成进度、成果卡片"),
        ("课程中心", "选择软件工程导论、浏览章节、播放PPT、进入练习", "章节树、资料预览、播放控制和学习状态"),
        ("AI资料中心", "查看推荐、收藏、反馈、刷新、生成资源", "画像匹配原因、资源类型、来源与有效性"),
        ("错题本", "按学科打开本子、查看错因、生成练习、重做提交", "分类标签、题目、解析、衍生练习和画像更新"),
        ("学情档案", "刷新画像、查看能力趋势、知识网络和下一步建议", "八维画像、证据状态、课程掌握和学习路径"),
        ("小智助手", "点击胶囊、接收提醒、语音交互", "与画像/规划Agent绑定的场景化提示；做题页不出现"),
    ], font_size=8)
    heading(doc, "3.2 外部接口", 2)
    add_table(doc, ["接口", "协议/数据", "用途", "保护与回退"], [
        ("DeepSeek兼容接口", "HTTPS / OpenAI Chat Completions", "普通文本理解与回答", "服务端密钥；超时重试；可切备用模型"),
        ("阿里云百炼Qwen", "HTTPS / Chat、VL、异步任务", "视觉理解、结构规划、严谨产物生成", "服务端密钥；模型回退；结构校验"),
        ("通义万相", "HTTPS异步文生图", "无文字场景插图", "仅匹配插图意图，不用于流程图/科学图"),
        ("Seedance", "HTTPS异步视频任务", "教学视频生成", "轮询、超时、额度错误转Manim链路"),
        ("Edge TTS", "HTTPS语音合成", "中文旁白", "音色回退、重试；结果与视频合成为AAC"),
        ("公开学习来源", "受控目录/搜索元数据", "个性化外部资源推荐", "限定国内教育来源、去重、验证与反馈"),
        ("PostgreSQL", "TCP/SQL", "结构化数据", "连接池、事务、Alembic迁移"),
        ("Redis/Celery", "TCP/队列", "长任务与状态", "任务租约、重试和结果后端"),
    ], font_size=8)
    heading(doc, "3.3 内部接口", 2)
    endpoints = [
        ("POST", "/api/v1/ai/chat/stream", "发起多智能体流式对话", "SSE事件流"),
        ("POST", "/api/v1/ai/attachments", "上传图片或文件", "附件标识与元数据"),
        ("GET", "/api/v1/ai/context/courses", "取得对话课程上下文", "课程列表"),
        ("POST", "/api/v1/rag/upload/preview|commit", "资料预览与入库", "切片预览/入库结果"),
        ("POST", "/api/v1/resource-generation/runs", "创建异步资源生成任务", "202 + run_id"),
        ("GET", "/api/v1/resource-generation/runs/{id}", "查询任务和步骤", "状态、进度、错误"),
        ("POST", "/api/v1/resource-workshop/images/analyze", "解析作业图片", "识别内容与建议"),
        ("GET", "/api/v1/resource-hub/recommendations", "取得个性化推荐", "站内外资源列表"),
        ("GET", "/api/v1/resource-hub/wrong-book", "按学科获取错题本", "分类与错题"),
        ("POST", "/api/v1/resource-hub/wrong-book/practice", "根据错题生成练习", "练习资源"),
        ("GET", "/api/v1/profile/digital-twin", "读取学习数字分身", "画像、能力和协同说明"),
        ("POST", "/api/v1/profile/analyze", "触发画像分析", "新画像版本"),
        ("GET", "/api/v1/learning-report/me", "取得学情报告", "趋势、证据、建议"),
        ("GET", "/api/v1/knowledge-graph/courses/{id}", "课程知识图谱", "节点与边"),
        ("GET", "/api/v1/healthz|readyz", "存活与就绪检查", "组件健康状态"),
    ]
    add_table(doc, ["方法", "路径", "用途", "响应"], endpoints, font_size=7)
    heading(doc, "3.4 接口通用约定", 2)
    for t in [
        "鉴权：除登录、公开健康检查外，业务接口携带Bearer JWT；对象查询同时校验user_id归属。",
        "错误：REST使用HTTP状态码和结构化detail；流式请求使用error事件并保留已生成内容。",
        "幂等：入库采用稳定file_id；生成任务使用run_id和步骤状态；反馈接口以业务主键更新。",
        "上传：服务器验证扩展名、MIME、大小和安全路径，附件通过受控下载接口访问。",
        "流式：事件包含phase、message、progress、artifact和done/error，前端按事件类型更新状态。",
    ]: bullet(doc, t)

    heading(doc, "4 运行设计", 1)
    heading(doc, "4.1 运行模块组合", 2)
    add_table(doc, ["运行场景", "模块组合", "主要状态"], [
        ("文本答疑", "Chat UI→AI Chat API→Supervisor→Retrieval/Tutor→Safety→SSE", "received/routing/retrieving/generating/reviewing/done"),
        ("图片题目", "Attachment→Vision Client→Tutor/Grading→课程证据→Safety", "uploaded/analyzing/recognized/answering/done"),
        ("资源包生成", "Workshop→Resource Run→Agent步骤→Artifact Service→Package", "queued/running/succeeded/failed/canceled"),
        ("视频生成", "意图路由→Seedance；失败则Qwen→Manim→TTS→PyAV", "submitted/polling/rendering/muxing/completed"),
        ("错题练习", "Wrong Book→Quiz Service→Grading→Evidence→Profile/Planner", "created/answering/submitted/graded/profile_updated"),
        ("画像刷新", "Evidence聚合→Profile Agent→版本化快照→学习路径建议", "collecting/analyzing/committed/synchronized"),
        ("资料推荐", "画像薄弱点→站内排序+外部发现→验证去重→反馈", "refreshing/ranked/delivered/feedback"),
    ], font_size=8)
    heading(doc, "4.2 运行控制", 2)
    numbered(doc, "后端启动时加载环境配置、建立数据库连接、验证迁移与必要目录，并暴露healthz/readyz。")
    numbered(doc, "前端加载用户身份、会话、课程上下文、当前学习任务和小智提醒；任何单一非核心请求失败不阻塞页面。")
    numbered(doc, "同步请求由FastAPI在并发预算内执行；长资源任务写入run/step记录后异步执行，前端轮询或接收事件。")
    numbered(doc, "Agent工具调用受角色白名单约束；Supervisor控制最大协作轮数，防止无限循环和重复调用。")
    numbered(doc, "模型产物经过结构、事实和安全检查后持久化；失败保留步骤错误和可恢复点。")
    numbered(doc, "系统停止时结束新任务接收，已提交任务由队列策略决定完成、重试或恢复。")
    heading(doc, "4.3 运行时间", 2)
    add_table(doc, ["类型", "目标/典型策略", "超时与控制"], [
        ("页面与普通API", "本地数据优先在1秒量级返回", "数据库索引、分页、缓存"),
        ("流式文本", "尽快返回阶段事件和首个文本片段", "每用户SSE并发2；模型超时和重试"),
        ("图片理解", "受图片大小、OCR和模型服务影响", "单文件≤配置上限；视觉超时180秒"),
        ("PPT/图表", "结构生成后本地排版，持续反馈进度", "教学产物超时；失败记录步骤"),
        ("视频", "异步执行，通常显著长于文本", "Seedance轮询/150秒；本地渲染回退"),
        ("知识库入库", "批量、增量、可重复", "大文件按需解析；稳定ID防重"),
    ], font_size=8)
    heading(doc, "4.4 并发、事务与状态一致性", 2)
    body(doc, "系统把外部模型调用与数据库事务分离：先以短事务创建任务和步骤，释放连接后执行耗时调用，再以短事务提交状态与产物，避免长事务占用连接。资源运行记录包含活动状态限制和尝试租约，防止同一任务被重复执行。画像更新采用版本号和事件记录；答题提交、错题更新和学习证据尽可能在同一业务事务中完成。")

    heading(doc, "5 系统数据结构设计", 1)
    heading(doc, "5.1 逻辑结构设计要点", 2)
    add_picture(doc, figures["data"], "图5-1 核心数据对象关系（概要级）")
    data_rows = [
        ("user/student", "用户和学生身份", "UUID、账号、角色、基本资料", "一对一/一对多关联业务对象"),
        ("student_profile", "版本化数字分身", "八维值、优势、薄弱点、画像版本、证据游标", "user_id唯一或按用户取最新"),
        ("learning_evidence", "画像与评估证据", "来源、知识点、得分、置信度、时间", "关联用户、课程或资源"),
        ("chatthread/conversation_message", "会话与消息", "角色、内容、会话、附件、时间", "线程一对多消息"),
        ("agent_task/chatartifact", "智能体执行与产物", "Agent、阶段、状态、结果、错误、资源地址", "关联会话和用户"),
        ("course/resource", "课程和学习资料", "课程、类型、标题、路径、知识点、难度", "课程一对多资源"),
        ("course_knowledge_node/edge", "知识图谱", "节点类型、名称、掌握度、先修边", "课程内有向关系"),
        ("question/quiz_attempt", "题目与作答", "题干、选项、答案、提交、评分", "资源一对多题目/尝试"),
        ("wrong_question", "错题档案", "学科、标签、错因、错误次数、掌握状态", "用户+题目维度"),
        ("learning_path/task", "路径与当前任务", "阶段、步骤、状态、排序、截止时间", "用户拥有动态路径"),
        ("resource_generation_run/step", "生成工作流", "类型、状态、进度、尝试、错误、租约", "运行一对多步骤"),
        ("generated_resource_package", "资源包", "标题、类型、文件清单、预览、来源", "由运行生成并归属用户"),
        ("external_resource/recommendation", "外部资源及推荐", "供应方、URL、摘要、验证、匹配原因、反馈", "外部资源与个性化推荐关联"),
    ]
    add_table(doc, ["结构", "含义", "主要数据项", "关系"], data_rows, font_size=7)
    body(doc, "2026-07-20开发验证库快照包含：126条学习资源、67个课程知识节点、106条知识关系、100条学习证据、46条个性化推荐、15条错题、11次测验尝试、240条Agent任务和2个学生画像。快照用于证明设计已落地，不作为容量上限。")
    heading(doc, "5.2 物理结构设计要点", 2)
    add_table(doc, ["存储", "数据", "访问/索引", "安全与备份"], [
        ("PostgreSQL", "用户、课程、画像、证据、任务、题目、推荐", "UUID主键；user_id/course_id/status/time等索引；事务访问", "账号最小权限；定期逻辑备份；Alembic迁移"),
        ("Chroma向量库", "课程切片和文档向量", "按集合、file_id、课程和相似度检索", "可由原文重建；目录访问受后端控制"),
        ("文件存储", "附件、原始资料、PPT、图片、视频、资源包", "UUID文件名或受控相对路径；元数据入库", "禁止路径穿越；下载鉴权；比赛发布时清理临时文件"),
        ("Redis", "队列、任务结果和短期状态", "任务键、TTL、消费者", "非权威数据；可从PostgreSQL恢复任务状态"),
        ("浏览器", "JWT和必要UI状态", "Pinia/内存或受控本地存储", "不保存外部API密钥和敏感画像全文"),
    ], font_size=8)
    heading(doc, "5.3 数据结构与程序的关系", 2)
    add_table(doc, ["数据结构", "聊天", "RAG/课程", "资源生成", "练习", "画像", "路径推荐"], [
        ("会话/消息", "CRUD", "R", "R", "", "R", ""), ("课程/资源/知识图", "R", "CRUD", "R", "R", "R", "R"),
        ("题目/尝试/错题", "R", "R", "C", "CRUD", "R", "R"), ("学习证据", "C", "C", "C", "C", "CRUD", "R"),
        ("学生画像", "R/U", "R", "R", "R", "CRUD", "R"), ("路径/任务", "R", "R", "C", "U", "R", "CRUD"),
        ("生成运行/产物", "C/R", "R", "CRUD", "R", "R", "R"), ("外部资源/推荐", "R", "C/R", "", "", "R", "CRUD"),
    ], font_size=8)
    body(doc, "注：C/R/U/D分别表示新增、读取、修改和删除。所有用户数据访问都必须带用户边界；跨模块只通过Service或API访问，不允许前端直接连接数据库。")
    heading(doc, "5.4 数据生命周期与隐私", 2)
    for t in [
        "采集最小化：只采集完成学习服务所需的对话、答题、资源使用和显式目标，不推断敏感身份属性。",
        "来源可追踪：学习证据记录来源类型、业务对象和时间；外部资源记录provider、URL、验证时间和元数据。",
        "画像可纠偏：保存画像版本和更新事件；低置信度维度显示“待积累”，不包装成确定事实。",
        "文件隔离：附件采用服务端生成标识和受控下载；临时解析文件按任务生命周期清理。",
        "密钥隔离：DeepSeek、百炼、Seedance等密钥只存在于本机或部署环境的.env，不写入数据库、文档或前端包。",
        "备份恢复：PostgreSQL定期备份；原始课程资料保留；向量库和生成索引可由原始文件重新构建。",
    ]: bullet(doc, t)

    heading(doc, "6 系统出错处理设计", 1)
    heading(doc, "6.1 出错信息", 2)
    errors = [
        ("401/403", "登录失效、密钥无效或对象无权限", "重新登录；检查服务端密钥和对象归属；不向前端回显密钥"),
        ("AI_CONCURRENCY_LIMIT", "同一用户并发AI任务超过预算", "前端提示等待/取消；后端释放完成或超时的令牌"),
        ("MODEL_TIMEOUT/CONNECTION", "模型网络超时或连接失败", "有限重试、切备用模型或返回可重试状态"),
        ("INSUFFICIENT_CREDITS", "Seedance等外部服务额度不足", "提示运营处理；视频自动回退本地Manim+TTS"),
        ("UNSUPPORTED_ATTACHMENT", "文件类型、大小或内容不受支持", "拒绝入库并说明允许类型和上限"),
        ("VISION_FORBIDDEN", "视觉模型无权限或配置不正确", "检查Qwen-VL模型权限；不以无关文本假装识图"),
        ("RAG_EVIDENCE_EMPTY", "知识库没有足够证据", "明确证据不足，建议上传资料或扩大受控检索"),
        ("ARTIFACT_RENDER_FAILED", "PPT、图表、Manim或音视频合成失败", "保留步骤日志；可恢复/重试；切换简化模板"),
        ("DB_SCHEMA_MISMATCH", "代码字段与数据库迁移不一致", "readyz标记异常；执行Alembic修复迁移；验证物理列"),
        ("EXTERNAL_RESOURCE_INVALID", "外部链接失效或来源不受信", "隐藏/替换资源，记录反馈并降低来源评分"),
        ("CONTENT_REVIEW_REJECTED", "事实、来源或安全审查未通过", "不交付原产物；修改提示、补充证据后重新生成"),
    ]
    add_table(doc, ["错误标识", "含义", "处理方法"], errors, font_size=8)
    heading(doc, "6.2 补救措施", 2)
    heading(doc, "6.2.1 后备技术", 3)
    body(doc, "数据库采用周期性备份；课程原始资料与结构化入库文件分离保存；向量库可以由原始资料重建；生成物和元数据同时记录，重要演示产物在提交前离线备份。环境配置提供.env.example但不包含真实密钥。")
    heading(doc, "6.2.2 降效技术", 3)
    body(doc, "系统按能力设置降效路径：模型不可用时保留本地课程浏览与练习；Seedance不可用时使用Qwen分镜、Manim和TTS；图像生成不适合严谨图表时切换Matplotlib或确定性DFD绘图；联网资源不可用时仅返回站内知识库；画像证据不足时提供通用路径并标明低置信度。")
    heading(doc, "6.2.3 恢复与再启动", 3)
    body(doc, "资源生成run和step保存状态、尝试次数、错误和租约。进程重启后可识别未完成任务并恢复或重新入队；幂等写入避免重复产物。数据库迁移在部署阶段执行并通过readyz验证。SSE断开不取消后端已提交的异步任务，用户可通过run_id重新查询。")
    heading(doc, "6.3 系统维护设计", 2)
    add_table(doc, ["维护点", "设计安排"], [
        ("配置", "Pydantic Settings集中管理；示例环境变量分能力分组；敏感值不写日志。"),
        ("数据库", "Alembic版本迁移；幂等修复迁移处理历史漂移；迁移后执行列和ORM查询验证。"),
        ("健康检查", "healthz检查进程存活，readyz检查数据库、关键目录和可选依赖。"),
        ("日志审计", "记录请求阶段、模型供应方、耗时、任务状态和错误摘要，不记录完整密钥。"),
        ("质量测试", "Service单元测试、API集成测试、前端类型检查、生产构建和边界检查。"),
        ("课程扩展", "通过课程数据、入库脚本、知识节点和Agent契约扩展，不修改核心编排。"),
        ("模型扩展", "通过provider配置与服务适配器接入，能力路由不依赖单一供应商。"),
    ], font_size=8)
    heading(doc, "6.4 监控与验收检查点", 2)
    for t in [
        "可用性：前后端、数据库、课程资源、RAG、文本模型和视觉模型分别检查，不以单一绿色状态掩盖局部故障。",
        "性能：记录首事件时间、总耗时、模型重试、生成步骤耗时和文件大小。",
        "质量：抽检课程答案引用、图片题目识别、错题衍生练习、PPT排版和视频音轨。",
        "画像：核对每个维度是否有证据、版本是否递增、低证据状态是否诚实展示。",
        "安全：检查未授权下载、路径穿越、危险MIME、提示注入、敏感输出和密钥泄漏。",
        "赛前：预热演示账号与知识库，检查第三方额度；准备离线回退产物和一键启动说明。",
    ]: bullet(doc, t)

    heading(doc, "附录A 赛题要求追踪矩阵", 1)
    trace = [
        ("不少于6维动态画像", "八维AI认知模型、证据可信度、版本事件", "student_profile_agent / DigitalTwinPanel", "已实现"),
        ("多智能体架构", "13类角色、主管编排、工具白名单和状态轨迹", "chat_engine/chat_runtime", "已实现"),
        ("至少5类个性化资源", "文档、PPT、题目、图表、图片、视频、代码案例", "resource_generation/teaching_artifact", "已实现"),
        ("个性化路径和推送", "路径、当前任务、站内外推荐、反馈排序", "learning_path/resource_recommendation", "已实现"),
        ("智能辅导", "RAG+多模态问答+图解/视频", "ai_chat/vision/teaching artifact", "已实现"),
        ("学习效果评估", "答题证据、错题、能力趋势、画像更新", "quiz/learning_report/profile", "已实现"),
        ("完整专业课程知识库", "软件工程导论127个原始资料文件及结构化入库", "原始资料/ + course_kb", "已实现"),
        ("防幻觉与安全", "证据优先、输出守卫、安全审查、来源边界", "content_quality/safety/output_guard", "已实现并持续完善"),
        ("流式/进度体验", "SSE、生成run/step、停止重试、卡片成果", "ChatLayout/resource_generation", "已实现"),
        ("开源与AI工具说明", "依赖清单、许可证、AI Coding边界", "本说明书附录B", "已说明"),
    ]
    add_table(doc, ["赛题要求", "智屿设计响应", "实现证据", "状态"], trace, font_size=7)

    heading(doc, "附录B 开源软件、外部AI与AI Coding说明", 1)
    body(doc, "本项目遵循“自主业务设计+合规复用开源基础设施”的原则。下表为主要依赖，具体版本以requirements.txt和package.json为准；正式提交前由团队再次核对许可证文本并随作品保留NOTICE。")
    add_table(doc, ["名称", "用途", "来源/协议说明"], [
        ("Vue / Vite / Pinia", "前端框架、构建和状态管理", "官方开源项目，主要为MIT协议"),
        ("Arco Design Vue", "前端组件库", "字节跳动开源，MIT协议"),
        ("FastAPI / Starlette", "后端Web与异步接口", "官方开源项目，MIT/BSD类协议"),
        ("LangGraph / LangChain", "智能体状态图与模型工具封装", "官方开源项目，按各仓库许可证使用"),
        ("PostgreSQL", "关系数据库", "PostgreSQL License"),
        ("Chroma", "向量检索", "Apache-2.0"),
        ("Mermaid", "知识图、流程图前端渲染", "MIT协议"),
        ("Manim", "教学动画确定性渲染", "MIT协议"),
        ("python-pptx / Matplotlib / PyAV", "PPT、图表和音视频合成", "按各官方开源许可证使用"),
        ("DeepSeek", "普通文本模型服务", "按服务平台条款调用，密钥由团队持有"),
        ("阿里云百炼Qwen/VL/万相", "结构规划、视觉理解和场景图片", "按阿里云服务条款调用"),
        ("Seedance", "教学视频生成", "按API服务条款调用；失败可回退本地链路"),
    ], font_size=8)
    body(doc, "AI Coding说明：项目开发过程中使用OpenAI Codex辅助进行代码检索、局部实现建议、重构、测试编写与文档生成。工具不替代团队的软件需求决策、架构设计、课程内容审核和最终验收；所有写入仓库的修改由团队成员检查，并通过编译、测试或运行验证。真实API密钥、账号口令和学生隐私数据不作为AI Coding提示内容进入提交文档。参赛提交时应按组委会最新口径补充工具版本、使用日期与团队审核人。")

    heading(doc, "附录C 部署与启动摘要", 1)
    add_table(doc, ["步骤", "操作"], [
        ("1. 环境", "安装Python 3.11+、Node.js 18+、PostgreSQL 14+；可选Redis与Manim。"),
        ("2. 配置", "复制code/.env.example为code/.env，填写数据库和模型服务配置；不得提交真实密钥。"),
        ("3. 迁移", "在code/backend目录运行：..\\.venv\\Scripts\\python.exe -m alembic upgrade head。"),
        ("4. 后端", "运行：..\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8001。"),
        ("5. 前端", "在code/education/course目录安装依赖后运行npm run dev。"),
        ("6. 验证", "访问/api/v1/healthz、/api/v1/readyz，登录后验证课程、聊天、附件、错题、画像与资源生成。"),
    ], font_size=8)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(36)
    r = p.add_run("—— 文档结束 ——"); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)

    core = doc.core_properties
    core.title = "智屿概要设计说明书"
    core.subject = "第十五届中国软件杯A3赛题参赛项目"
    core.author = "智屿项目团队"
    core.keywords = "多智能体, 个性化学习, 多模态, 学生画像, 资源生成"
    core.comments = "依据赛题官网、概要设计规范和项目当前实现生成"
    doc.save(DOCX)
    print(DOCX)
    for name, path in figures.items(): print(name, path)


if __name__ == "__main__":
    build_document()
