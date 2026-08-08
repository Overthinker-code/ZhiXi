from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "智屿智能教育平台_软件需求规格说明书.md"
OUTPUT_DIR = ROOT / "output"
OUTPUT = OUTPUT_DIR / "智屿智能教育平台_软件需求规格说明书.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    run.bold = header
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    heading_settings = {
        "Title": ("黑体", 30, "25355B"),
        "Subtitle": ("黑体", 20, "4F46E5"),
        "Heading 1": ("黑体", 16, "17233D"),
        "Heading 2": ("黑体", 14, "27364F"),
        "Heading 3": ("黑体", 12, "334155"),
    }
    for name, (font_name, size, color) in heading_settings.items():
        style = styles[name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("智屿智能教育平台 · 软件需求规格说明书")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    footer.add_run(" 页")


def add_rich_text(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            run.font.color.rgb = RGBColor(67, 56, 202)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_cover(doc: Document) -> int:
    paragraphs = SOURCE.read_text(encoding="utf-8").splitlines()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    run = title.add_run("智 屿")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(79, 70, 229)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("智能教育平台")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(26)

    doc_title = doc.add_paragraph()
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title.paragraph_format.space_before = Pt(38)
    run = doc_title.add_run("软件需求规格说明书")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(37, 53, 91)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.add_run("文档编号：ZHIYU-SRS-001    版本：V1.0")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(8)
    for i, (label, value) in enumerate(
        [
            ("参赛团队", "____________________"),
            ("项目负责人", "____________________"),
            ("指导教师", "____________________"),
            ("编制日期", "2026 年 7 月 19 日"),
        ]
    ):
        set_cell_text(table.cell(i, 0), label, header=True)
        set_cell_text(table.cell(i, 1), value)
        set_cell_shading(table.cell(i, 0), "EEF2FF")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(36)
    run = note.add_run("依据 GB/T 9385—2008 编制")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_page_break()

    # Markdown cover consumes the first title, subtitle and metadata block.
    for idx, line in enumerate(paragraphs):
        if line.strip() == "## 文档修订记录":
            return idx
    return 0


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    rows = [rows[0]] + rows[2:]
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), value, header=r_idx == 0)
            if r_idx == 0:
                set_cell_shading(table.cell(r_idx, c_idx), "E8EEFF")
            elif r_idx % 2 == 0:
                set_cell_shading(table.cell(r_idx, c_idx), "F8FAFC")
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6FA")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(51, 65, 85)


def render_body(
    doc: Document,
    lines: list[str],
    start: int,
    *,
    source_path: Path = SOURCE,
) -> None:
    idx = start
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[idx + 1]):
            table_lines = [line]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(doc, table_lines)
            continue
        if not stripped or stripped == "---":
            idx += 1
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            caption, raw_path = image_match.groups()
            image_path = Path(raw_path)
            if not image_path.is_absolute():
                image_path = source_path.parent / image_path
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.add_run().add_picture(str(image_path), width=Cm(16.0))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.first_line_indent = Cm(0)
                run = cp.add_run(caption)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(71, 85, 105)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(f"[{caption}：图片文件不存在 {raw_path}]")
                run.font.color.rgb = RGBColor(190, 24, 93)
            idx += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(0)
            set_cell = OxmlElement("w:shd")
            set_cell.set(qn("w:fill"), "F4F5FF")
            p._p.get_or_add_pPr().append(set_cell)
            add_rich_text(p, stripped[2:])
            idx += 1
            continue
        match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if match:
            level = len(match.group(1))
            text = match.group(2)
            if level == 1 and re.match(r"^\d+\s", text):
                doc.add_page_break()
            p = doc.add_heading(text, level=level)
            p.paragraph_format.first_line_indent = Cm(0)
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or numbered:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.8)
            add_rich_text(p, (bullet or numbered).group(1))
            idx += 1
            continue
        p = doc.add_paragraph()
        add_rich_text(p, stripped.replace("  ", " "))
        idx += 1


def add_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "智屿智能教育平台软件需求规格说明书"
    props.subject = "项目相关材料——软件需求规格说明书"
    props.author = "智屿项目组"
    props.keywords = "智屿, 智能教育, 软件需求规格说明书, 多智能体, 学生画像"
    props.comments = "依据 GB/T 9385—2008 编制"


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    start = add_cover(doc)
    render_body(doc, lines, start)
    add_properties(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
