#!/usr/bin/env python3
"""Build a submission-ready DOCX draft from the competition Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BRAND = "4F46E5"
TEXT = RGBColor(23, 33, 58)
MUTED = RGBColor(102, 112, 133)
SANS = "Arial Unicode MS"
SERIF = "Arial Unicode MS"


def set_font(run, name: str, size: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = TEXT


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_inline_markdown(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, SANS, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Menlo", 9)
            run.font.color.rgb = RGBColor(67, 56, 202)
        else:
            run = paragraph.add_run(part)
            set_font(run, SERIF)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, SANS, 9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    tail = paragraph.add_run(" 页")
    set_font(tail, SANS, 9)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (
        ("Heading 1", 18, RGBColor(49, 46, 129)),
        ("Heading 2", 14, RGBColor(67, 56, 202)),
        ("Heading 3", 11.5, RGBColor(79, 70, 229)),
    ):
        style = doc.styles[name]
        style.font.name = SANS
        style._element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("智屿 · 第十五届中国软件杯 A3 赛题")
    set_font(hr, SANS, 9, True)
    hr.font.color.rgb = RGBColor(79, 70, 229)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, text)
            for run in p.runs:
                set_font(run, SANS, 8.7, r_idx == 0)
                if r_idx == 0:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shade_cell(cell, BRAND if r_idx == 0 else ("F7F8FC" if r_idx % 2 == 0 else "FFFFFF"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build(source: Path, target: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = lines[0].lstrip("# ") if lines else "智屿系统开发说明书"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(85)
    title_p.paragraph_format.space_after = Pt(14)
    title_run = title_p.add_run(title)
    set_font(title_run, SANS, 25, True)
    title_run.font.color.rgb = RGBColor(49, 46, 129)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("第十五届中国软件杯 A3 组参赛材料")
    set_font(sub_run, SANS, 13, True)
    sub_run.font.color.rgb = RGBColor(79, 70, 229)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    meta_run = meta.add_run("项目名称：智屿个性化智能学习平台\n文档状态：初稿\n日期：2026-07-14")
    set_font(meta_run, SANS, 11)
    meta_run.font.color.rgb = MUTED
    doc.add_page_break()

    i = 1
    while i < len(lines) and not lines[i].strip().startswith("##"):
        i += 1
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.right_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                set_font(run, "Menlo", 8.5)
                run.font.color.rgb = RGBColor(55, 65, 81)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [[cell.strip() for cell in stripped.strip("|").split("|")]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            doc.add_heading(heading.group(2), level=level)
            i += 1
            continue
        list_match = re.match(r"^[-*]\s+(.*)$", stripped)
        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if list_match or numbered_match:
            p = doc.add_paragraph(style="List Bullet" if list_match else "List Number")
            add_inline_markdown(p, (list_match or numbered_match).group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped.replace("  ", " "))
        i += 1

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_competition_docx.py SOURCE.md TARGET.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
