from __future__ import annotations

import os
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
ASSETS = OUT / "report_manual_assets"
os.environ.setdefault("MPLCONFIGDIR", str(ROOT / "code/backend/uploads/.matplotlib"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLUE = "315EFB"
DEEP_BLUE = "173A8F"
GRAY = "667085"
LIGHT = "F5F7FF"


def setup_fonts():
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def box(ax, x, y, w, h, title, subtitle="", color="#315EFB", fill="#EEF2FF"):
    p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.018",
                       linewidth=1.5, edgecolor=color, facecolor=fill)
    ax.add_patch(p)
    ax.text(x + w / 2, y + h * .63, title, ha="center", va="center",
            fontsize=12, weight="bold", color="#172B4D")
    if subtitle:
        ax.text(x + w / 2, y + h * .30, subtitle, ha="center", va="center",
                fontsize=9, color="#52607A", wrap=True)


def arrow(ax, a, b, color="#7181A6"):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=12,
                                 linewidth=1.4, color=color))


def make_figures():
    ASSETS.mkdir(parents=True, exist_ok=True)
    setup_fonts()

    # Test strategy
    fig, ax = plt.subplots(figsize=(12, 6.5), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .93, "智屿分层测试策略", ha="center", fontsize=20, weight="bold", color="#173A8F")
    levels = [
        (.12, .13, .76, .13, "竞赛场景与人工验收", "学习闭环、PPT放映、图片解题、资源生成、演示稳定性", "#FFF4E8"),
        (.19, .30, .62, .13, "接口与安全测试", "鉴权、对象级权限、课程RAG、附件/生成物访问、内容安全", "#F3EEFF"),
        (.26, .47, .48, .13, "服务集成测试", "智能体编排、画像、错题、资源包、图表与视频回退链", "#EAF8F6"),
        (.33, .64, .34, .13, "单元与静态质量门", "pytest、Vue TypeScript、生产构建", "#EAF0FF"),
    ]
    for x, y, w, h, t, s, f in levels:
        box(ax, x, y, w, h, t, s, fill=f)
    ax.text(.5, .06, "原则：风险优先 · 证据可复现 · 外部模型采用契约测试与受控冒烟测试", ha="center", fontsize=11, color="#52607A")
    fig.tight_layout(); fig.savefig(ASSETS / "01_test_strategy.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # Test result chart
    fig, ax = plt.subplots(figsize=(11.5, 6), dpi=180)
    labels = ["服务层 pytest", "关键 API", "TypeScript 检查", "生产构建"]
    passed = [115, 26, 1, 1]
    failed = [2, 0, 0, 0]
    y = range(len(labels))
    ax.barh(list(y), passed, color="#315EFB", label="通过")
    ax.barh(list(y), failed, left=passed, color="#F04438", label="未通过")
    for i, (p, f) in enumerate(zip(passed, failed)):
        ax.text(p + f + 1, i, f"{p} 通过 / {f} 未通过", va="center", fontsize=10)
    ax.set_yticks(list(y), labels); ax.invert_yaxis(); ax.set_xlabel("测试/质量门数量")
    ax.set_title("已完成自动化验证结果", fontsize=18, weight="bold", color="#173A8F", pad=16)
    ax.spines[["top", "right"]].set_visible(False); ax.legend(loc="lower right")
    fig.tight_layout(); fig.savefig(ASSETS / "02_test_results.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # User navigation
    fig, ax = plt.subplots(figsize=(12.5, 7), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .94, "智屿用户学习闭环与功能导航", ha="center", fontsize=20, weight="bold", color="#173A8F")
    items = [
        (.04, .59, "课程中心", "软件工程导论\nPPT/笔记/试题"),
        (.24, .59, "AI 伴学", "文本、图片、文件\n引用与流式回答"),
        (.44, .59, "练习与错题本", "学科自动分类\n变式练习"),
        (.64, .59, "学情档案", "八维画像\n能力成长趋势"),
        (.84, .59, "AI 资料中心", "个性化推送\n生成与收藏"),
    ]
    for i, (x, y0, t, s) in enumerate(items):
        box(ax, x, y0, .13, .18, t, s, fill=["#EAF0FF", "#ECFBF8", "#FFF4E8", "#F3EEFF", "#FFF8E8"][i])
        if i < len(items)-1: arrow(ax, (x+.13, y0+.09), (items[i+1][0], y0+.09))
    box(ax, .27, .23, .46, .16, "多智能体 + 学生画像", "学习行为 → 证据 → 画像更新 → 路径/资源 → 新练习 → 再评价", color="#12B8A6", fill="#EAF8F6")
    for x, _, _, _ in items:
        arrow(ax, (x+.065, .59), (.5, .39))
    ax.text(.5, .10, "悬浮小智提供情境提醒；做题界面自动隐藏，避免干扰与不当辅助", ha="center", fontsize=11, color="#52607A")
    fig.tight_layout(); fig.savefig(ASSETS / "03_user_loop.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # Deployment/startup
    fig, ax = plt.subplots(figsize=(12, 5.5), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .91, "本地运行拓扑与启动顺序", ha="center", fontsize=19, weight="bold", color="#173A8F")
    nodes = [
        (.04, "PostgreSQL", "5432\n业务与画像数据"), (.24, "Redis（可选）", "6379\n异步任务/状态"),
        (.44, "FastAPI", "8001\nREST/SSE/WS"), (.64, "Vue 3 + Vite", "5174\n浏览器界面"),
        (.84, "外部能力", "DeepSeek/Qwen\nSeedance/Edge TTS"),
    ]
    for i, (x, t, s) in enumerate(nodes):
        box(ax, x, .42, .13, .23, t, s, fill=["#EEF5FA", "#FFF1F0", "#ECFBF8", "#EAF0FF", "#FFF8E8"][i])
        if i < 4: arrow(ax, (x+.13, .535), (nodes[i+1][0], .535))
    ax.text(.5, .18, "启动前：配置 .env → 执行 Alembic 迁移 → 启动后端 → 启动前端 → 检查 /api/v1/health", ha="center", fontsize=11, color="#52607A")
    fig.tight_layout(); fig.savefig(ASSETS / "04_startup.png", bbox_inches="tight", facecolor="white"); plt.close(fig)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)


def table(doc, headers, rows, widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers): shade(t.rows[0].cells[i], BLUE); cell_text(t.rows[0].cells[i], h, True, "FFFFFF", size)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, False, None, size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri % 2: shade(cells[i], "F7F9FC")
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Cm(width)
    doc.add_paragraph(); return t


def body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(.74); p.paragraph_format.line_spacing = 1.45; p.paragraph_format.space_after = Pt(5)
    p.add_run(text); return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2"); p.paragraph_format.line_spacing = 1.25; p.add_run(text); return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.line_spacing = 1.3; p.add_run(text); return p


def picture(doc, file, caption, width=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(file), width=Inches(width))
    p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GRAY)


def toc(doc):
    p = doc.add_paragraph(); r = p.add_run();
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "目录（在 Word/WPS 中右键更新域）"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r._r.extend([begin, instr, sep, txt, end])


def configure(doc, short_title):
    sec = doc.sections[0]; sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.1); sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]; normal.font.name = "Microsoft YaHei"; normal.font.size = Pt(10.5); normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for name, size, color in [("Title", 28, DEEP_BLUE), ("Heading 1", 18, DEEP_BLUE), ("Heading 2", 14, BLUE), ("Heading 3", 11.5, "344054")]:
        s = doc.styles[name]; s.font.name = "Microsoft YaHei"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color); s._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    header = sec.header.paragraphs[0]; header.text = f"智屿——基于大模型的个性化资源生成与学习多智能体系统　{short_title}"; header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header.runs: r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 "); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld); footer.add_run(" 页")


def cover(doc, title, subtitle):
    doc.add_paragraph("中国软件杯大学生软件设计大赛参赛项目").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("智　屿"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(subtitle); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(5): doc.add_paragraph()
    table(doc, ["文档版本", "编制日期", "文档状态"], [["V1.0", str(date.today()), "提交评审版"]], size=10)
    p = doc.add_paragraph("参赛团队：智屿项目组"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def references(doc):
    rows = [
        ("R1", "中国软件杯大学生软件设计大赛赛题页面", "https://www.cnsoftbei.com/content-3-1286-1.html", "功能、技术、创新与材料要求"),
        ("R2", "测试分析报告编写规范", "用户提供的国家标准模板", "章节结构与内容要求"),
        ("R3", "用户手册编写规范", "用户提供的国家标准模板", "用户手册章节结构"),
        ("R4", "智屿软件需求规格说明书 V1.0", "项目交付文档", "验收基线"),
        ("R5", "智屿概要/详细设计说明书 V1.0", "项目交付文档", "架构、接口与数据设计"),
    ]
    table(doc, ["编号", "资料", "来源", "用途"], rows, [1.2, 5, 6, 5])


def build_test_report():
    doc = Document(); configure(doc, "测试分析报告")
    cover(doc, "测试分析报告", "依据国家标准模板与赛题验收要求编制")
    doc.add_heading("文档修订记录", 1)
    table(doc, ["版本", "日期", "说明", "编制/审核"], [["V1.0", str(date.today()), "首次形成竞赛提交版；纳入真实自动化测试与部署缺口", "智屿项目组"]])
    doc.add_heading("目录", 1); toc(doc); doc.add_page_break()

    doc.add_heading("1 引言", 1)
    doc.add_heading("1.1 编写目的", 2)
    body(doc, "本报告给出智屿系统的测试范围、方法、环境、结果、缺陷与综合评价，为竞赛评审、演示前验收、部署交付和后续回归提供可复现证据。报告遵循用户提供的测试分析报告编写规范，并把赛题中的动态学生画像、多智能体、五类以上资源、个性化路径、多模态辅导、完整课程知识库、安全与流式交互转化为测试目标。")
    doc.add_heading("1.2 项目背景", 2)
    body(doc, "智屿面向高校学生，将课程知识库、学习行为证据、八维动态画像、13类智能体角色和多模态内容生成统一到个人学习闭环中。参赛初始课程为《软件工程导论》，课程资料包含13章笔记、试题、PPT及扩展资料；平台同时支持课程中心、AI伴学、资料中心、错题本、学习路径、学情档案和悬浮小智。")
    doc.add_heading("1.3 定义与缩写", 2)
    table(doc, ["术语", "含义"], [("RAG", "检索增强生成，用课程证据约束回答"), ("SSE", "服务器发送事件，用于流式回答和生成进度"), ("DFD", "数据流图；系统采用确定性绘图链路而非插画模型"), ("学生画像", "由知识、目标、偏好、行为、动力、风格、问题解决、知识网络等维度构成的动态模型"), ("质量门", "必须通过的自动化检查，如 pytest、类型检查和生产构建")])
    doc.add_heading("1.4 参考资料", 2); references(doc)

    doc.add_heading("2 测试概要", 1)
    doc.add_heading("2.1 测试对象与版本", 2)
    table(doc, ["对象", "范围/版本"], [("前端", "Vue 3 + TypeScript + Vite 单页应用"), ("后端", "Python 3.13、FastAPI、SQLModel/SQLAlchemy、Alembic"), ("数据", "PostgreSQL 18.4；课程知识库与生成物文件存储"), ("智能能力", "DeepSeek 文本；Qwen-Max/VL；Wanx；Seedance；Manim/Matplotlib/python-pptx/Edge TTS 回退链"), ("基线", "2026-07-20 工作区版本，V1.0 竞赛提交候选")])
    doc.add_heading("2.2 策略与方法", 2); picture(doc, ASSETS / "01_test_strategy.png", "图2-1 分层测试策略")
    body(doc, "采用风险优先的分层验证：底层执行单元、服务与静态质量检查；中层验证 API、鉴权、对象级授权和持久化；上层以典型学习任务进行端到端与人工验收。计费或受网络配额影响的外部模型不做无约束全量调用，使用契约测试、模拟响应和少量真实冒烟测试结合。")
    doc.add_heading("2.3 赛题要求覆盖", 2)
    table(doc, ["赛题关键点", "验证对象", "判定依据", "结论"], [
        ("≥6维动态画像", "八维画像、证据更新、成长趋势", "画像服务/模式测试、页面构建", "已覆盖"),
        ("多智能体架构", "Supervisor、Tutor、Retrieval、Profile、Planner 等", "编排与服务测试", "已覆盖"),
        ("≥5类资源", "文档、PPT、题目、图表、音视频、外部资源", "资源生成/访问测试", "已覆盖"),
        ("个性化路径与推送", "画像驱动路径、资源推荐、错题变式", "画像、推荐、测验服务测试", "已覆盖"),
        ("多模态辅导", "粘贴/拖拽图片和文件、视觉理解", "附件访问与视觉客户端测试", "已覆盖，受外部配额影响"),
        ("完整专业课程知识库", "软件工程导论 13 章及 PPT/题库", "课程 RAG 与对象授权测试", "已覆盖"),
        ("防幻觉与安全", "引用证据、内容审查、上传限制、权限隔离", "安全路由及访问控制测试", "已覆盖"),
        ("流式/进度交互", "SSE 状态、停止/重试、生成物卡片", "接口与生产构建", "已覆盖"),
    ], size=8)
    doc.add_heading("2.4 环境与测试数据", 2)
    table(doc, ["类别", "配置"], [("操作系统", "Windows 11 开发测试机"), ("运行时", "Python 3.13；Node.js 22.23.1；npm 10.9.8"), ("数据库", "PostgreSQL 18.4 x86_64 Windows"), ("浏览器/前端", "Chromium 系浏览器；Vite 3.2.11"), ("测试数据", "软件工程导论知识库；演示学生、错题、学习证据、推荐记录和生成物"), ("限制", "单机本地环境；外部 API 受网络、余额、并发和模型权限影响；硬件规格未纳入自动采集")])
    doc.add_heading("2.5 测试执行摘要", 2)
    table(doc, ["测试域", "命令/方式", "结果", "耗时"], [
        ("服务层", "pytest app/tests/services --basetemp ...", "115 通过，2 未通过，440 warnings", "23.95s（pytest）/42.03s（总）"),
        ("关键 API A", "附件、生成物、内容安全、课程 RAG", "20/20 通过", "154.99s/173.52s"),
        ("关键 API B", "对象级授权、健康与安全", "6/6 通过", "11.58s/33.66s"),
        ("前端类型", "npm run type:check", "通过", "81.89s"),
        ("前端构建", "npm run build", "5979 modules，成功", "165.33s"),
        ("全量套件", "pytest app/tests -q", "5分钟上限内未完成，转为分域执行", ">300s，未计入通过率"),
    ], size=8)

    doc.add_heading("3 测试结果及发现", 1)
    picture(doc, ASSETS / "02_test_results.png", "图3-1 已完成自动化验证结果")
    doc.add_heading("3.1 后端服务层", 2)
    body(doc, "共完成117项服务层用例，115项通过，服务层通过率为98.29%。两项未通过均已复现并定位为部署依赖缺失：视频确定性栈回退需要系统可执行的 FFmpeg；资源包 PDF 读取需要 ReportLab。首次运行还出现 pytest 临时目录权限错误，改用工作区内 --basetemp 后消除，说明其不属于业务断言失败。")
    doc.add_heading("3.2 API、安全与数据隔离", 2)
    body(doc, "关键接口分组执行26项，全部通过。覆盖聊天附件和生成媒体访问控制、内容安全路由、课程 RAG 权限、学生对象级授权以及健康检查安全策略。结果表明，系统能够阻止跨用户访问私有对象，并对生成物和课程证据实施鉴权。")
    doc.add_heading("3.3 前端质量门", 2)
    body(doc, "TypeScript 类型检查退出码为0；生产构建成功，Vite 转换5979个模块。构建出现 Sass 旧版 JS API、eval 使用和部分压缩后分块大于1000 KiB 的警告，不影响当前功能交付，但会增加维护和首屏加载风险，应在后续版本治理。")
    doc.add_heading("3.4 多模态与教学媒体", 2)
    body(doc, "定向媒体路由测试7项通过，证明 DFD/流程图请求可进入确定性图表链而非通用插画链；已生成并检查《TCP拥塞控制基础知识_中文配音版.mp4》，时长20.8秒，包含H.264视频轨和AAC音频轨。真实云模型仍可能因401/403、并发上限、余额或模型权限失败，前端需展示可操作错误并允许重试或回退。")
    doc.add_heading("3.5 数据库与迁移", 2)
    body(doc, "数据库迁移已执行到028，针对 external_resource 等新增字段完成存在性检查，缺失字段为空，ORM 查询成功。数据库相关验证必须以 Alembic upgrade head 为启动前置，否则可能出现 UndefinedColumn。")

    doc.add_heading("4 对软件功能的结论", 1)
    table(doc, ["功能", "能力与证据", "限制", "结论"], [
        ("账号与权限", "JWT、对象级授权、私有附件/生成物访问测试通过", "密钥和演示密码必须部署时更换", "通过"),
        ("AI伴学", "文本流式回答、课程RAG、引用来源、停止/重试", "模型延迟、并发、网络和配额会影响可用性", "通过"),
        ("图片/文件理解", "拖拽、粘贴、上传、视觉路由与附件鉴权", "需配置有 qwen-vl-max 权限的百炼密钥", "条件通过"),
        ("图表/DFD", "确定性绘图与 Mermaid 渲染；媒体路由用例通过", "复杂图需人工检查语义和布局", "通过"),
        ("视频", "Seedance 主链；Qwen分镜+Manim+TTS回退；带音轨样例已验证", "FFmpeg 未安装时部分回退不可用", "条件通过"),
        ("PPT与课程中心", "PPT生成、课件资源、预览/放映入口", "浏览器兼容和字体替换需演示机复测", "通过"),
        ("错题与练习", "学科/标签分类、错题保留、变式练习服务", "分类质量依赖标签与练习Agent", "通过"),
        ("学情档案", "八维画像、趋势、知识网络、路径建议", "冷启动阶段证据不足会显示待积累", "通过"),
        ("资源推荐", "课程内外部资源、画像驱动排序、来源展示", "外部平台链接有效性需定期更新", "通过"),
    ], size=7.8)

    doc.add_heading("5 分析摘要", 1)
    doc.add_heading("5.1 软件能力", 2)
    body(doc, "智屿已经形成从课程输入、智能辅导、练习评价、证据沉淀、画像更新到路径与资源再生成的闭环。核心优势不是单一聊天窗口，而是将多智能体、课程知识证据和学生画像绑定，并按任务选择确定性工具或生成模型；该设计直接对应赛题的功能与创新要求。")
    doc.add_heading("5.2 缺陷与风险清单", 2)
    table(doc, ["编号", "级别", "问题", "影响", "处置/复测条件"], [
        ("DEP-01", "中", "当前虚拟环境缺少 ReportLab", "PDF资源包读取用例失败", "安装 requirements.txt；复测对应服务用例"),
        ("DEP-02", "中", "系统 PATH 无独立 FFmpeg", "确定性视频栈的特定回退失败", "安装 FFmpeg 并验证 ffmpeg -version"),
        ("PERF-01", "低", "部分前端 chunk >1000KiB", "弱网首屏和缓存效率", "路由懒加载、manualChunks、按需导入"),
        ("MAINT-01", "低", "Sass legacy API/eval 警告", "未来升级兼容与安全审计", "升级依赖并移除非必要 eval"),
        ("TEST-01", "低", "全量套件单次超过5分钟", "CI反馈慢", "按域并行、标记慢测、缓存模型/数据库夹具"),
        ("EXT-01", "中", "外部模型配额/权限/网络不稳定", "视觉、图片、视频能力降级", "启动健康检查、预算控制、降级链和清晰错误提示"),
    ], size=8)
    doc.add_heading("5.3 改进建议与优先级", 2)
    numbered(doc, "演示机先执行依赖自检，补齐 ReportLab 与 FFmpeg，复测2项失败用例并保留日志。")
    numbered(doc, "将服务、API、安全、媒体、慢速外部能力拆分为并行 CI 作业；对外部 API 使用每日受控冒烟测试。")
    numbered(doc, "对前端进行路由级懒加载和大包拆分；建立首屏、SSE首字和生成任务P95耗时基线。")
    numbered(doc, "对多模态识别、课程引用正确率、推荐采纳率和错题变式有效性建立人工标注评测集。")
    doc.add_heading("5.4 综合评价", 2)
    body(doc, "结论：核心业务、安全隔离、前端质量门和关键API达到竞赛演示与阶段交付要求；当前版本建议“有条件通过”。条件为补齐 FFmpeg、ReportLab，并在目标演示机执行启动检查、关键路径冒烟和两项失败用例复测。外部生成服务应准备余额、权限和网络异常时的回退演示方案。")

    doc.add_heading("6 测试资源消耗", 1)
    table(doc, ["资源", "消耗/说明"], [("人员", "1台开发测试机上的自动化执行与人工结果审查；未折算多人并行工时"), ("自动化用例", "服务层117项；关键API26项；媒体路由7项；前端质量门2项"), ("完成质量门耗时", "约496.4秒（分组命令总耗时求和，不含全量超时尝试与人工检查）"), ("全量尝试", "一次300秒未完成；一次较大API组240秒未形成最终摘要，不计入正式通过数"), ("外部API成本", "本轮自动化以契约/模拟为主；真实冒烟受云端计费、余额与限流控制"), ("存储", "测试临时目录、课程知识库、生成PNG/PPTX/MP4；交付前应清理无关临时文件")])

    doc.add_heading("附录A 复测清单", 1)
    for x in ["安装 requirements.txt 中 ReportLab 并确认 import reportlab 成功。", "安装 FFmpeg 并确认 ffmpeg -version 可执行。", "执行 Alembic upgrade head。", "重跑117项服务测试，目标117/117。", "运行 type:check 与生产构建。", "以学生账号完成：登录→课程PPT→图片解题→错题练习→画像更新→资源推送。", "验证 Seedance 不可用时带中文音轨的视频回退。", "确认做题页面不显示悬浮小智。"]: bullet(doc, "□ " + x)
    path = OUT / "智屿测试分析报告_V1.0.docx"; doc.save(path); return path


def build_user_manual():
    doc = Document(); configure(doc, "用户手册")
    cover(doc, "用户手册", "学生端、教师/演示人员与部署管理员适用")
    doc.add_heading("文档修订记录", 1)
    table(doc, ["版本", "日期", "说明", "编制/审核"], [["V1.0", str(date.today()), "首次形成竞赛提交版用户手册", "智屿项目组"]])
    doc.add_heading("目录", 1); toc(doc); doc.add_page_break()

    doc.add_heading("1 引言", 1)
    doc.add_heading("1.1 编写目的", 2)
    body(doc, "本手册帮助学生完成课程学习、AI辅导、图片/文件提问、PPT学习、错题巩固和学情查看；帮助教师或竞赛演示人员准备课程与演示数据；帮助管理员完成安装、配置、启动、备份和故障恢复。")
    doc.add_heading("1.2 背景", 2)
    body(doc, "智屿是面向高校的个性化学习平台，以完整课程知识库为输入，以多智能体协作和动态学生画像为核心。系统针对《软件工程导论》提供笔记、PPT、题目和扩展资源，并可迁移到人工智能、计算机、电子信息等其他课程。")
    doc.add_heading("1.3 定义", 2)
    table(doc, ["术语", "说明"], [("AI伴学", "支持课程证据、多模态输入和生成物交付的对话学习区"), ("小智", "与画像和多智能体联动的悬浮学习助手；做题时隐藏"), ("错题本", "按课程/题目标签自动分类并支持变式练习"), ("学情档案", "八维画像、成长趋势、知识网络和下一步建议"), ("生成物", "PPTX、PNG图表、MP4视频、题目或资源包等可下载成果")])
    doc.add_heading("1.4 参考资料", 2); references(doc)

    doc.add_heading("2 软件用途", 1)
    doc.add_heading("2.1 功能", 2); picture(doc, ASSETS / "03_user_loop.png", "图2-1 用户学习闭环与功能导航")
    table(doc, ["模块", "主要用途", "典型结果"], [("首页/课程中心", "进入课程，学习笔记、PPT、试题和资源", "学习进度、PPT放映"), ("AI伴学", "文本、图片、文件提问；课程证据回答", "流式解释、引用、图表/PPT/视频"), ("练习/错题本", "作答、收集错题、按学科分类、生成变式", "练习反馈与新证据"), ("学情档案", "查看能力、趋势、知识网络和画像总结", "个性化路径与建议"), ("AI资料中心", "检索、收藏、生成和接收外部学习资源", "文档、视频、公开课链接"), ("小智", "提醒、解释和快捷跳转", "与画像弱项和Agent任务联动")])
    doc.add_heading("2.2 性能与时间特性", 2)
    body(doc, "普通页面和数据库查询以本地交互为目标；AI回答采用SSE逐步显示，用户无需等待完整结果。图片、PPT和视频属于异步或长任务，界面显示阶段、百分比、停止和重试。实际耗时由网络、模型服务、文件大小、任务复杂度和演示机性能决定。")
    doc.add_heading("2.3 灵活性", 2)
    bullet(doc, "课程可扩展：用同一知识库规范导入新的高校专业课程。")
    bullet(doc, "模型可路由：文本、视觉、图像、视频分别配置，支持主模型与回退链。")
    bullet(doc, "输入可组合：文字、粘贴图片、拖拽文件和课程上下文可同时使用。")
    bullet(doc, "输出可复用：回答、图表、PPT、视频、题目和引用均可查看或下载。")
    doc.add_heading("2.4 安全与保密", 2)
    body(doc, "用户只能访问本人会话、附件、错题、画像与生成物。API密钥、数据库密码和JWT密钥只写入本机 .env，不得放进截图、报告或 Git。凡曾公开粘贴过的密钥均应在云平台立即作废并重新生成。生产部署应使用HTTPS、32字节以上JWT密钥、强密码、受限CORS/TrustedHost和最小权限数据库账号。")

    doc.add_heading("3 运行环境", 1)
    doc.add_heading("3.1 推荐硬件", 2)
    table(doc, ["项目", "最低建议", "竞赛演示建议"], [("CPU", "4核 x64", "8核及以上"), ("内存", "8GB", "16GB及以上"), ("磁盘", "20GB 可用", "50GB SSD，可容纳知识库与生成视频"), ("网络", "可访问所配置模型API", "稳定有线/高速Wi-Fi并准备离线回退样例"), ("显示", "1366×768", "1920×1080")])
    doc.add_heading("3.2 支持软件", 2)
    table(doc, ["软件", "本项目验证/建议", "用途"], [("Windows", "Windows 11", "开发与演示"), ("Python", "3.13（建议按 requirements 锁定）", "FastAPI后端、绘图、PPT和媒体"), ("Node.js/npm", "22.23.1 / 10.9.8", "Vue前端"), ("PostgreSQL", "18.4", "业务、课程、画像与任务数据"), ("Redis", "可选，建议7.x", "Celery异步队列与状态"), ("FFmpeg", "需可通过PATH调用", "视频编解码与回退合成"), ("浏览器", "最新版Edge/Chrome", "平台使用与PPT预览")])
    doc.add_heading("3.3 数据与文件", 2)
    table(doc, ["类型", "说明/限制"], [("通用上传", "默认 MAX_UPLOAD_SIZE=25MB；以服务器配置为准"), ("预览文件", "单文件最大25MB；压缩包最多800项，并受解压总量限制"), ("视频媒体", "默认最大150MB"), ("课程资料", "PDF、PPT/PPTX、DOC/DOCX、Markdown、TXT、题库等；以页面提示和后端白名单为准"), ("图片", "PNG、JPG/JPEG、WEBP等常用格式；含题目时应清晰、端正、无过度反光"), ("生成物", "PNG、PPTX、MP4、PDF/资源包等；存入后端 uploads 并由鉴权接口访问")])

    doc.add_heading("4 使用过程", 1)
    doc.add_heading("4.1 安装与初始化", 2); picture(doc, ASSETS / "04_startup.png", "图4-1 本地运行拓扑与启动顺序")
    doc.add_heading("4.1.1 准备环境变量", 3)
    body(doc, "复制 code/.env.example 为 code/.env，设置 PostgreSQL 连接、JWT密钥、DeepSeek文本模型和百炼Qwen视觉/结构能力等。不要把真实密钥提交到版本库。示例中的 changethis 仅为占位符。")
    doc.add_heading("4.1.2 初始化并启动后端", 3)
    for line in ["cd C:\\Users\\Eileen\\ZhiYu-main\\code\\backend", "..\\.venv\\Scripts\\python.exe -m alembic upgrade head", "..\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8001"]:
        p = doc.add_paragraph(); p.style = doc.styles["Normal"]; p.paragraph_format.left_indent = Cm(.8); r = p.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9); shade_dummy = None
    body(doc, "PowerShell 中路径与命令之间不能插入空格；正确写法是 ..\\.venv，而不是“..\\ .venv”。看到 Uvicorn running 和健康接口成功后再启动前端。")
    doc.add_heading("4.1.3 启动前端", 3)
    for line in ["cd C:\\Users\\Eileen\\ZhiYu-main\\code\\education\\course", "$env:HUSKY=0; npm ci", "npm run dev"]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(.8); r = p.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9)
    body(doc, "浏览器访问 http://localhost:5174。若 Vite 报 ECONNREFUSED 127.0.0.1:8001，说明后端未启动、已退出或端口不一致。")
    doc.add_heading("4.1.4 登录", 3)
    body(doc, "开发初始化会创建演示学生 student@example.com，当前开发默认密码为 student123456；另有历史兼容邮箱 syudent@example.com。管理员账号由 .env 的 FIRST_SUPERUSER 和 FIRST_SUPERUSER_PASSWORD 决定。正式演示或部署前必须修改默认密码。")

    doc.add_heading("4.2 输入", 2)
    doc.add_heading("4.2.1 输入背景与格式", 3)
    table(doc, ["输入方式", "操作", "建议"], [("文本", "在AI伴学输入框键入问题", "写明课程、任务、输出形式和限制"), ("粘贴图片", "复制截图后在输入框 Ctrl+V", "等待附件卡片出现再发送"), ("拖拽文件", "把文件拖入聊天输入区域", "确认文件名和数量；勿上传敏感资料"), ("选择文件", "点击回形针/上传按钮", "适合PDF、DOCX、PPTX、图片和文本"), ("课程上下文", "从课程中心进入或在对话中指定课程", "优先选择《软件工程导论》以获得课程RAG证据")])
    doc.add_heading("4.2.2 输入实例", 3)
    for x in ["知识问答：请结合《软件工程导论》课程资料解释高内聚、低耦合，并列出引用来源。", "图片解题：粘贴题目截图后输入“识别完整题干，先复述题目，再分步解答并检查答案”。", "数据流图：生成教师工资调整系统的数据流图，输出黑字白底 PNG，不要插图，不要 Mermaid。", "PPT：生成一份软件需求分析入门课件，10页，包含学习目标、案例、练习和总结。", "教学视频：生成一段讲解 TCP 拥塞控制基础知识的视频，要求中文配音、字幕和状态变化示意。"]: bullet(doc, x)
    doc.add_heading("4.2.3 图片识别检查", 3)
    numbered(doc, "发送前确认附件区显示文件，而不只是输入框文字。")
    numbered(doc, "回答应先复述图片题干；若复述错误，停止并上传更清晰截图。")
    numbered(doc, "出现403时检查 QWEN/MULTIMODAL 的 Provider、Model、API Key 和模型权限，不要继续用错误识别结果答题。")

    doc.add_heading("4.3 输出", 2)
    table(doc, ["输出", "表现形式", "主要操作"], [("AI回答", "流式Markdown、公式、代码、引用来源", "停止、重试、展开来源"), ("图表/DFD", "页面渲染或PNG卡片", "预览、下载；核对节点和箭头"), ("PPT课件", "PPTX生成物；课程中心预览/放映", "下载或打开全屏放映"), ("教学视频", "MP4卡片，可能由Seedance或本地教学栈生成", "播放、检查声音、下载"), ("练习", "题目、选项、解析和掌握证据", "作答、提交、加入错题本"), ("画像与路径", "八维指标、趋势、知识网络、下一步建议", "刷新建议、进入课程/任务")])
    doc.add_heading("4.4 主要功能操作", 2)
    doc.add_heading("4.4.1 课程中心与PPT放映", 3)
    numbered(doc, "进入顶部“课程中心”，选择“软件工程导论”。")
    numbered(doc, "在章节/资源列表筛选 PPT 或课件，点击资源卡片。")
    numbered(doc, "进入预览后使用上一页、下一页、页码或全屏放映；浏览器不支持时下载PPTX用本地Office/WPS打开。")
    numbered(doc, "学习行为会形成证据并用于画像与资源推荐；退出前确认进度已保存。")
    doc.add_heading("4.4.2 AI伴学与生成物", 3)
    body(doc, "在AI伴学中明确区分三类需求：知识解释交给文本/RAG；严谨图表进入 Mermaid、Matplotlib 或确定性绘图；无文字场景插画进入 Wanx；视频进入 Seedance，失败时回退到Qwen分镜、Manim和中文TTS。若要求“不要Mermaid”，系统应返回PNG图而不是代码块。")
    doc.add_heading("4.4.3 错题本与变式练习", 3)
    numbered(doc, "进入AI资料中心或学习区的“我的错题本”。")
    numbered(doc, "按课程本子查看；系统根据课程、题目标签和练习Agent判断学科，通用仅作为低置信度回退。")
    numbered(doc, "打开具体错题查看题干、错误次数、答案与解析。")
    numbered(doc, "点击“根据错题生成练习”，等待新练习出现；完成并提交后，结果继续更新画像。")
    doc.add_heading("4.4.4 学情档案与学习路径", 3)
    numbered(doc, "进入“学情档案”查看综合能力、近30日成长、学习投入与优先提升项。")
    numbered(doc, "在数字分身中查看学习阶段、类型、优势、待提升方向、八维模型和知识网络。")
    numbered(doc, "点击“刷新学习建议”或“生成我的学习路径”；新建议基于当前证据，冷启动时先完成课程和练习。")
    doc.add_heading("4.4.5 悬浮小智", 3)
    body(doc, "小智在资料库和AI聊天等学习页面提供提醒、问题解释和快捷任务；提醒由画像弱项、学习节奏和多智能体任务决定。在正式做题页面小智自动隐藏，避免干扰和直接给答案。")

    doc.add_heading("4.5 文件查询与管理", 2)
    bullet(doc, "对话附件：在原消息中查看“已附加N个文件”；生成物在助手消息卡片中下载。")
    bullet(doc, "课程资料：按课程、章节、资源类型或关键词筛选。")
    bullet(doc, "资料中心：查看推荐理由、来源、类型，收藏后可再次访问。")
    bullet(doc, "隐私对象：会话、附件、错题、画像和生成物均按登录用户隔离，不能直接拼接他人URL访问。")

    doc.add_heading("4.6 出错处理与恢复", 2)
    table(doc, ["现象", "原因", "处理"], [
        ("ECONNREFUSED 127.0.0.1:8001", "后端未启动/端口不符", "按4.1.2启动；确认Vite代理为8001"),
        ("数据库 UndefinedColumn", "迁移未执行", "在 backend 执行 alembic upgrade head 后重启"),
        ("401 Invalid API Key", "密钥无效/变量名错误", "在对应平台新建密钥，更新.env并重启；不要把密钥贴入聊天"),
        ("403 Forbidden（视觉）", "百炼模型未开通、密钥无权限或Provider错误", "检查Qwen VL配置、地域端点和模型权限"),
        ("AI concurrency limit exceeded", "并发预算已占满或任务未释放", "停止重复任务，等待后重试；管理员检查并发配置"),
        ("Seedance insufficient credits", "视频账户余额不足", "充值/更换可用账号，或使用本地教学视频回退"),
        ("视频无声音", "仅生成画面或音频合成失败", "要求中文配音；检查Edge TTS、AAC音轨与FFmpeg/PyAV"),
        ("图表生成插画", "任务误路由到通用图像模型", "明确“DFD/流程图、黑白PNG、不要插图/不要Mermaid”；检查artifact路由"),
        ("npm ECONNRESET", "网络/代理/镜像中断", "清理错误代理、切换稳定网络/镜像后重试npm ci"),
        ("husky .git can't be found", "子目录安装脚本找不到仓库", "在仓库内执行；开发安装可临时 $env:HUSKY=0; npm ci"),
        ("PDF资源包失败", "ReportLab未安装", "pip install -r requirements.txt 并复测"),
        ("视频回退找不到ffmpeg", "FFmpeg未安装或PATH无效", "安装后运行 ffmpeg -version，重启终端和后端"),
    ], size=7.5)
    doc.add_heading("4.7 终端与管理员操作", 2)
    table(doc, ["目的", "命令"], [("数据库迁移", "..\\.venv\\Scripts\\python.exe -m alembic upgrade head"), ("启动后端", "..\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8001"), ("前端类型检查", "npm run type:check"), ("前端生产构建", "npm run build"), ("服务层测试", "pytest app/tests/services --basetemp .test-tmp-services"), ("检查FFmpeg", "ffmpeg -version"), ("检查Python", "..\\.venv\\Scripts\\python.exe --version")], size=8)

    doc.add_heading("附录A 学生快速开始", 1)
    for x in ["登录演示学生账号并修改默认密码。", "进入《软件工程导论》，放映一份PPT并阅读一章笔记。", "向AI伴学提一个带课程引用的问题。", "粘贴一张题目图片，确认系统先准确复述题干。", "完成一道练习并查看错题本分类。", "生成变式练习并提交。", "打开学情档案，查看八维画像和下一步路径。", "从资料中心收藏一条与弱项匹配的资源。"]: bullet(doc, "□ " + x)
    doc.add_heading("附录B 竞赛演示前检查", 1)
    for x in ["后端8001、前端5174、PostgreSQL均正常；Redis按配置启用。", "执行数据库迁移；演示学生画像、错题、课程和推荐数据完整。", "DeepSeek/Qwen/Seedance密钥已轮换且余额、权限、地域端点有效。", "FFmpeg、ReportLab、Edge TTS/PyAV可用；准备一份带声音的离线视频回退样例。", "演示图片不含隐私；所有附件均能识别并有访问控制。", "PPT可预览/全屏，字体和页面比例正常。", "做题页小智隐藏；聊天/资料页小智可用。", "准备断网、限流、余额不足时的说明与回退路线。"]: bullet(doc, "□ " + x)
    path = OUT / "智屿用户手册_V1.0.docx"; doc.save(path); return path


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True); make_figures()
    print(build_test_report())
    print(build_user_manual())
