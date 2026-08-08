from __future__ import annotations

import os
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
os.environ.setdefault("MPLCONFIGDIR", str(ROOT / "code" / "backend" / "uploads" / ".matplotlib"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_overview_design_doc import (
    BLUE, DEEP_BLUE, GRAY, add_caption, add_page_number, add_picture,
    add_table, add_toc, body, bullet, configure_matplotlib, heading, numbered,
    page_break, save_architecture, save_data_model, save_deployment,
)


OUT = ROOT / "output"
ASSETS = OUT / "detailed_design_assets"
DOCX = OUT / "智屿详细设计说明书_V1.0.docx"


def dbox(ax, x, y, w, h, title, sub="", color="#315EFB", fill="#F7F9FF", size=11):
    p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=.012,rounding_size=.018",
                       linewidth=1.4, edgecolor=color, facecolor=fill)
    ax.add_patch(p)
    ax.text(x+w/2, y+h*.62, title, ha="center", va="center", fontsize=size,
            weight="bold", color="#172B4D")
    if sub:
        ax.text(x+w/2, y+h*.29, sub, ha="center", va="center", fontsize=size-2,
                color="#52607A", wrap=True)
    return p


def darr(ax, a, b, color="#7181A6", rad=0):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=12,
                                linewidth=1.35, color=color,
                                connectionstyle=f"arc3,rad={rad}"))


def save_program_structure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 9), dpi=180)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .96, "智屿程序系统层次结构", ha="center", fontsize=21, weight="bold", color="#173A8F")
    dbox(ax, .38, .82, .24, .10, "ZHIYU-WEB", "Vue前端 + FastAPI后端", fill="#EAF0FF", size=13)
    row1 = [
        (.03,"P01 AI对话编排"),(.23,"P02 课程RAG"),(.43,"P03 多模态理解"),
        (.63,"P04 资源工作流"),(.83,"P05 教学产物")]
    row2 = [
        (.03,"P06 练习错题"),(.23,"P07 画像评估"),(.43,"P08 路径推荐"),
        (.63,"P09 小智数字人"),(.83,"P10 前端安全运维")]
    for x,t in row1:
        dbox(ax,x,.61,.14,.11,t,"核心服务程序",fill="#ECFBF8",size=10); darr(ax,(.50,.82),(x+.07,.72))
    for x,t in row2:
        dbox(ax,x,.37,.14,.11,t,"闭环与支撑程序",color="#8B5CF6",fill="#F7F2FF",size=10); darr(ax,(.50,.82),(x+.07,.48),rad=.06)
    dbox(ax,.06,.13,.25,.10,"PostgreSQL", "画像·任务·资源·证据",color="#336791",fill="#EEF5FA",size=12)
    dbox(ax,.375,.13,.25,.10,"Chroma + 文件存储", "向量·附件·PPT·生成物",color="#12B8A6",fill="#ECFBF8",size=12)
    dbox(ax,.69,.13,.25,.10,"模型与异步服务", "DeepSeek/Qwen/Seedance/Redis",color="#F59E0B",fill="#FFF8E8",size=12)
    for x in (.18,.50,.82): darr(ax,(x,.37),(x,.23))
    ax.text(.5,.045,"共同约束：JWT用户边界 · 证据优先 · 状态持久化 · 幂等与回退 · 内容安全",ha="center",fontsize=11,color="#52607A")
    fig.tight_layout(); fig.savefig(path,bbox_inches="tight",facecolor="white"); plt.close(fig)


def save_flow(path: Path, title: str, steps: list[tuple[str, str]], fallback: str = "") -> None:
    height = max(6.5, len(steps)*.72+1.8)
    fig, ax = plt.subplots(figsize=(11, height), dpi=180)
    ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off")
    ax.text(.5,.965,title,ha="center",fontsize=18,weight="bold",color="#173A8F")
    top=.88; gap=.72/(max(1,len(steps)-1)); h=min(.09,gap*.62)
    for i,(name,desc) in enumerate(steps):
        y=top-i*gap
        fill=["#EAF0FF","#ECFBF8","#FFF4E8","#F7F2FF"][i%4]
        dbox(ax,.18,y-h/2,.64,h,name,desc,fill=fill,size=11)
        if i<len(steps)-1: darr(ax,(.5,y-h/2),(.5,y-gap+h/2))
    if fallback:
        ax.text(.84,.50,"异常/回退",rotation=90,ha="center",va="center",fontsize=10,weight="bold",color="#B54708")
        dbox(ax,.83,.29,.15,.42,"恢复路径",fallback,color="#F59E0B",fill="#FFF8E8",size=10)
        darr(ax,(.82,.50),(.83,.50),color="#F59E0B")
    fig.tight_layout(); fig.savefig(path,bbox_inches="tight",facecolor="white"); plt.close(fig)


MODULES = [
    {
        "id":"P01", "name":"AI对话与多智能体编排程序", "ident":"AI_CHAT_ORCHESTRATOR",
        "files":"app/ai/chat_engine.py；app/ai/chat_runtime.py；api/v1/endpoints/ai_chat.py",
        "description":"面向AI伴学会话的核心常驻服务程序。每个请求创建独立状态，可重入并支持并发；通过SSE输出阶段与文本。程序把用户问题、课程、画像和附件整理为统一状态，由Supervisor选择专业Agent，最后执行引用、安全和可见内容整理。",
        "features":["会话与消息持久化","规则优先与Supervisor模型双路由","13类Agent节点及工具白名单","RAG证据注入与严格课程模式","流式阶段、文本、产物和错误事件","最终答案清洗、引用规范化和任务指标"],
        "performance":"普通请求尽快返回phase事件；每用户SSE并发上限2，同步AI并发1；上下文按配置裁剪，默认单条消息最多8000字符。图执行限制协作轮次，防止Agent循环。",
        "inputs":[("message","string，1..配置上限","学生问题或指令","浏览器"),("session_id","UUID/字符串，可选","会话标识","Chat UI"),("course_context","对象，可选","课程、章节、知识点","课程中心"),("attachments","数组，受所有权校验","图片或文件引用","附件程序"),("mode/tools","枚举/对象","辅导模式和允许工具","前端配置"),("profile_context","服务端对象","画像、目标、偏好","画像程序")],
        "outputs":[("SSE phase","事件对象","路由、检索、生成、审查进度"),("answer delta","UTF-8文本片段","面向学生的可见回答"),("citations","结构化数组","课程或文件证据"),("artifact","资源对象","图片、PPT、视频或练习入口"),("trace/metrics","结构化对象","Agent跳转、耗时和模型信息"),("error","受控错误事件","错误码、可重试建议")],
        "algorithms":["意图优先级：显式视频/图片/PPT/严谨图表意图先走能力路由，避免被通用Mermaid或文本规则截获；其余请求先规则匹配，无法确定时调用Supervisor。","图状态包含messages、selected_agent、route_trace、rag_context、worker_outputs和final_answer。Supervisor只做拆解与调度，专业Agent完成任务，finalize节点统一汇总。","严格课程模式下，若RAG未返回有效课程证据，则不允许把通用模型知识伪装为课程结论；输出证据不足与最小补充建议。","输出前移除内部JSON、thinking标签、工具日志和Agent内部名；引用标识与citation数组交叉核验后再发送。"],
        "steps":[("接收与鉴权","验证用户、会话、并发预算"),("构造统一状态","合并问题、课程、画像和附件"),("能力与Agent路由","显式媒体优先，随后规则/主管判断"),("检索与工具执行","RAG、视觉、资源或专业工具"),("专业Agent处理","Tutor/Grading/Planner等生成结果"),("安全与证据审查","引用、事实边界和内容安全"),("SSE输出并持久化","回答、轨迹、指标与产物")],
        "interfaces":[("上层","POST /api/v1/ai/chat/stream","AIChatStreamRequest → SSE"),("下层","RAGService / VisionClient / ArtifactService","函数调用"),("下层","ChatModelFactory","LangChain兼容模型调用"),("数据","chatthread/conversation_message/agent_task/chatartifact","SQLModel事务")],
        "storage":[("内存状态","单请求图状态与流式缓冲","请求结束释放"),("数据库","会话、消息、Agent任务、产物元数据","长期"),("附件目录","上传文件与索引","受控生命周期")],
        "comments":["所有路由优先级分支必须注明为何早于通用意图。","对用户可见文本与内部过程分离，禁止输出思维链。","新增Agent时同步更新AgentName、配置、图节点和工具白名单。"],
        "limits":["依赖外部文本模型可用性和上下文窗口。","复杂请求可能跨多次模型调用，耗时高于普通对话。","用户主动停止只终止当前流；已提交异步资源任务需单独取消。"],
        "tests":[("显式媒体路由","视频、PNG DFD、PPT提示词","进入正确能力分支"),("附件越权","引用其他用户附件","403"),("课程严格模式","无RAG证据","明确证据不足"),("并发限制","同用户超预算","429/受控错误"),("流式轨迹","正常问答","phase顺序与done完整")],
        "unresolved":"需要继续压测高并发SSE连接与超长多轮会话；生产环境应引入集中式分布式并发令牌。",
        "fallback":"模型超时→重试/备用模型\n流断开→保留消息与任务\n无证据→边界提示",
    },
    {
        "id":"P02", "name":"课程知识库与RAG程序", "ident":"COURSE_RAG_SERVICE",
        "files":"app/services/rag_service.py；document_processor.py；vector_store_factory.py",
        "description":"负责课程原始资料的预览、解析、切片、入库、权限过滤和混合检索。程序为非常驻业务服务对象，可重入；向量库为权威检索索引，原始文件与结构化资料为可重建来源。",
        "features":["PDF/DOCX/PPTX/Markdown等文档解析","preview/commit/cancel两阶段入库","课程与用户作用域隔离","向量相似度与中文字符词法召回","参考文件列表、删除和知识库重置","稳定file_id与重复入库控制"],
        "performance":"普通查询只在候选集合内计算；词法回退文档量受RAG_LEXICAL_MAX_DOCUMENTS限制，默认2500。大文件采用预览和按需解析，避免启动阶段全量处理。",
        "inputs":[("file","二进制+MIME，≤25MB默认","课程或用户资料","课程中心/资料中心"),("query","UTF-8文本","检索问题","Agent"),("scope","system/course/user","可见范围","服务端上下文"),("course_id/user_id","UUID，可选","访问边界","鉴权上下文"),("top_k","正整数","召回数量","调用程序")],
        "outputs":[("preview","流式切片摘要","待确认的解析结果"),("commit result","文件ID和切片数","入库确认"),("documents","文本、分数、元数据","RAG证据"),("files","参考资料元数据数组","资料管理界面")],
        "algorithms":["入库采用预览—确认模型：解析结果先写临时缓存，用户确认后才写入向量库；取消则删除预览缓存。","检索先构造course/user/scope过滤条件，向量召回后补充词法匹配。词法使用规范化字符与查询词重叠评分，适配中英文混排知识点。","结果按文档主键和切片位置去重，并再次执行可见性检查；用户只能管理自己上传的文件，系统课程资料只读。","软件工程导论使用稳定file_id，全量脚本重复运行时先删除旧版本，避免重复切片。"],
        "steps":[("接收文件或查询","类型、大小与身份检查"),("解析与标准化","抽取文本、页码、标题和元数据"),("切片/查询向量化","生成chunk或query embedding"),("作用域过滤","课程、系统、用户和所有权"),("向量+词法混合召回","计算分数并合并候选"),("去重与证据包装","返回出处、片段和适用边界")],
        "interfaces":[("上层","/api/v1/rag/upload/preview|commit|query","REST"),("下层","DocumentProcessor","文件解析"),("下层","VectorStoreFactory/Chroma","向量写入与查询"),("数据","原始资料目录/预览缓存","文件访问")],
        "storage":[("Chroma","文本、向量与作用域元数据","可从原文重建"),("原始资料","PPT、教材、笔记和题目","长期保留"),("预览缓存","待提交解析结果","短期/可取消")],
        "comments":["过滤条件修改必须同步权限测试。","切片应记录file_id、页码、课程和来源。","解析失败不得写入半成品向量。"],
        "limits":["扫描PDF可能无法直接抽取文字。","嵌入模型变化后需重建向量。","约497MB课程资料不适合每次启动全量解析。"],
        "tests":[("课程隔离","跨课程查询","不返回其他课程切片"),("文件越权","删除他人文件","403"),("重复入库","相同稳定ID","不产生重复切片"),("词法回退","短中文知识点","返回相关候选"),("预览取消","取消token","缓存被清理")],
        "unresolved":"后续增加扫描教材OCR队列、增量索引版本和向量模型迁移工具。",
        "fallback":"向量服务异常→词法检索\n解析失败→仅建资料索引\n证据为空→提示上传资料",
    },
    {
        "id":"P03", "name":"多模态附件与视觉理解程序", "ident":"MULTIMODAL_INPUT_SERVICE",
        "files":"api/v1/endpoints/ai_chat.py；services/vision_client.py；document_processor.py",
        "description":"负责聊天中的复制粘贴、拖拽和文件选择上传，以及图片的真实视觉识别。附件先落入用户隔离目录并登记索引，模型调用前重新校验所有权。图片内容、补充文字、课程和历史上下文共同构成视觉请求。",
        "features":["拖拽/粘贴/选择统一上传","MIME、扩展名、大小和最小图像尺寸检查","附件所有权索引与受控下载","Qwen-VL主备模型调用","可选OCR文本交叉证据","视觉失败明确返回，禁止无关猜测"],
        "performance":"图片默认上限20MB，通用附件上限25MB；视觉超时默认180秒。每个请求只发送必要图片和裁剪后的上下文，避免Base64导致请求体膨胀。",
        "inputs":[("UploadFile","二进制","图片/PDF/文档","ChatComposer"),("AttachmentRef","file_id/name/type","已上传附件引用","对话请求"),("prompt","文本","学生补充指令","对话请求"),("course context","对象","课程与章节","服务端"),("vision config","模型/API地址/密钥","能力配置",".env")],
        "outputs":[("attachment meta","file_id、名称、MIME、大小","前端附件卡"),("image context","识别文字、图形和不确定性","Tutor Agent"),("download stream","受控二进制","文件预览"),("vision error","权限/配置/识别错误","用户可见边界说明")],
        "algorithms":["上传时服务端生成file_id和安全文件名，不信任客户端路径；索引记录owner_id，读取和下载前再次核对。","图像解码后检查尺寸和格式；调用视觉模型时按主模型、备用模型顺序尝试。HTTP 401/403不被包装成成功识别。","视觉消息采用image_url/data URL与文本共同输入；题目图片必须把识别题干、课程知识点和学生指令联合判断。","OCR只作为可选交叉证据；当OCR与视觉结论冲突时标记不确定并请求用户补充。"],
        "steps":[("客户端捕获附件","拖拽、粘贴或选择"),("上传校验","大小、类型、名称和用户"),("隔离存储与索引","生成file_id并登记owner"),("图像解码/OCR","验证可读性并提取辅助文字"),("视觉模型调用","Qwen-VL主备模型"),("上下文融合","课程、提示与识别结果"),("交给辅导/批改Agent","输出答案或不确定性")],
        "interfaces":[("上层","POST /api/v1/ai/attachments","multipart/form-data"),("上层","GET /api/v1/ai/attachments/{id}","鉴权下载"),("下层","call_vision_model/build_chat_image_context","Python函数"),("数据","uploads/chat_attachments + index","文件/JSON")],
        "storage":[("附件目录","用户上传二进制","按清理策略"),("附件索引","file_id、owner、MIME和路径","持久"),("请求内存","Base64/视觉载荷","调用后释放")],
        "comments":["不得根据文件名推断图片内容。","所有本地路径必须resolve后验证在允许目录内。","视觉异常应保留供应商状态码摘要但不泄露密钥。"],
        "limits":["复杂手写体、低清照片和密集公式可能误识别。","视觉API受账号权限、额度和网络影响。","超大PPT/PDF不进入单次视觉请求。"],
        "tests":[("上传成功","合法PNG","返回file_id"),("附件越权","他人file_id","403"),("伪造MIME","扩展名与内容不符","拒绝"),("视觉403","无模型权限","明确配置错误"),("图片题问答","题目图+课程","回答围绕真实题干")],
        "unresolved":"增加图像预处理、旋转校正、公式OCR和多页文档视觉采样策略。",
        "fallback":"主视觉模型失败→备用VL\nOCR可用→提供辅助证据\n仍失败→要求补充题干",
    },
    {
        "id":"P04", "name":"个性化资源生成工作流程序", "ident":"RESOURCE_GENERATION_WORKFLOW",
        "files":"services/resource_generation_service.py；endpoints/resource_generation.py；models/resource_run.py",
        "description":"负责把主题、课程、学生画像和资源类型转化为可追踪的异步资源生成运行。每次运行由run和多个step组成，支持幂等、单用户活动任务约束、尝试租约、取消、恢复、质量审查和资源包交付。",
        "features":["画像和课程上下文构造","多资源并行生成","Agent执行轨迹与引用","质量门禁与一次重做","run/step状态持久化","取消、恢复、租约和幂等键","资源包清单与摘要校验"],
        "performance":"耗时任务返回202与run_id；同一用户默认只允许一个requested/running任务。并行生成受线程/模型并发预算控制，单项失败不覆盖其他成功产物。",
        "inputs":[("subject/topic","1..160字符","课程与知识点","资源工坊"),("resource_types","枚举数组","文档、题目、图等","用户"),("profile","服务端对象","薄弱点、偏好、难度","画像程序"),("course evidence","片段/引用","课程依据","RAG"),("idempotency_key","≤128字符，可选","重复提交控制","客户端")],
        "outputs":[("run","id、status、current_step","进度查询"),("steps","角色、模型、摘要、耗时、错误","可追踪执行"),("package","manifest与artifact列表","下载/预览"),("path event","新增学习节点","路径程序"),("error","错误码、失败步骤、恢复建议","前端")],
        "algorithms":["请求正规化后计算request_digest；同用户+idempotency_key唯一，重复请求返回原运行。部分唯一索引限制同一用户同时仅有一个requested/running运行。","执行器为每次尝试生成active_attempt_id并增加attempt_sequence，写入lease_expires_at；只有持有有效租约的尝试可以推进状态。","生成上下文包含课程、主题、画像信号和RAG证据摘要；不同资源类型可并行生成，之后执行结构、占位符、协议标记、主题一致性和安全质量门禁。","质量门禁失败时允许受控重做；仍失败则记录具体reason，不把低质量产物伪装为成功。成功后写manifest、哈希和资源关联，再更新学习路径。"],
        "steps":[("创建运行","校验、摘要、幂等和活动任务约束"),("获取执行租约","attempt_id、序号和过期时间"),("构造共享上下文","课程证据、画像和资源要求"),("多Agent并行生成","按资源类型调用专业能力"),("审查与受控重做","结构、主题、来源和安全"),("写入资源包","文件、manifest、哈希和预览"),("更新路径与状态","succeeded/failed/canceled")],
        "interfaces":[("上层","POST/GET /resource-generation/runs","REST 202/查询"),("下层","ResourceGenerationService.generate","业务调用"),("下层","TeachingArtifactService/模型客户端","资源类型适配"),("数据","resource_generation_run/step/package","PostgreSQL+文件")],
        "storage":[("run表","请求、状态、尝试、租约和错误","长期审计"),("step表","输入输出摘要、引用、模型、耗时","长期"),("package目录","artifact与manifest","用户资源生命周期")],
        "comments":["状态改变必须验证active_attempt_id。","外部调用不得包在长数据库事务中。","manifest写入后计算哈希并校验路径。"],
        "limits":["外部模型并行数受账号配额限制。","本地PPT/视频生成占用CPU和磁盘。","恢复只能从已持久化步骤边界继续。"],
        "tests":[("幂等重复提交","相同键与请求","返回同run"),("活动任务冲突","已有running","拒绝第二任务"),("租约过期","旧attempt提交","不允许推进"),("质量门禁","包含占位符/跑题","重做或失败"),("产物完整性","manifest哈希不符","拒绝下载")],
        "unresolved":"生产环境进一步接入Celery分布式worker、任务优先级和资源配额计费。",
        "fallback":"步骤失败→记录并重试\n主模型失败→备用/确定性模板\n重启→按run/lease恢复",
    },
    {
        "id":"P05", "name":"教学产物渲染与媒体路由程序", "ident":"TEACHING_ARTIFACT_RENDERER",
        "files":"services/teaching_artifact_service.py；bailian_service.py；seedance_video_service.py；teaching_video_audio.py",
        "description":"把模型输出的结构计划转化为可验证的PPT、图表、DFD、图片和有声视频。程序采用能力路由而非统一文生图：严谨图进入确定性绘图，场景插图才进入万相，视频优先Seedance并保留本地动画回退。",
        "features":["Qwen结构化JSON规划","python-pptx课件排版","Matplotlib科学图表","黑字白底DFD确定性绘制","Mermaid知识图前端渲染","万相场景插图","Seedance视频与Manim+TTS回退"],
        "performance":"PPT/图表在本地生成；视频异步且默认轮询间隔10秒、超时150秒。图片默认1024×1024，视频默认720p/16:9。媒体大小受20MB图片、150MB视频限制。",
        "inputs":[("topic/request","文本","主题与格式要求","资源/聊天程序"),("storyboard/plan","结构化JSON","标题、步骤、公式、要点","Qwen"),("artifact kind","枚举","PPT/chart/DFD/image/video","能力路由"),("media config","尺寸、时长、比例","生成参数","配置")],
        "outputs":[("PPTX","Office文档","可播放课件"),("PNG","黑白图表或场景图","预览/下载"),("Mermaid code","受限语法","前端渲染"),("MP4","H.264+AAC","有声教学视频"),("artifact metadata","URL、大小、预览、voice","对话卡片")],
        "algorithms":["意图分类先判断显式视频，再判断图片/严谨图表，避免“生成视频”被Mermaid关键词截获。DFD请求且明确黑白图片时直接使用确定性布局，不调用插画模型。","PPT由Qwen返回title/slides/bullets/note结构，python-pptx统一母版、字号和分页；模型不直接输出二进制文件。","科学图表由结构计划生成Matplotlib代码/参数并限制允许操作；场景插图使用万相，但提示词禁止品牌水印和密集文字。","视频优先创建Seedance任务并轮询。出现额度不足或非堆栈视频生成失败时，Qwen生成3..7步分镜，Manim渲染H.264，Edge TTS生成中文语音，PyAV合成为AAC音轨；旁白更长时延长最后一帧。"],
        "steps":[("识别资源意图","视频/严谨图/插图/PPT"),("生成结构计划","Qwen JSON并做字段校验"),("选择确定性或生成式工具","按准确性要求路由"),("渲染产物","pptx/Matplotlib/Manim/外部任务"),("质量检查","文件头、尺寸、音轨和主题"),("持久化与卡片化","安全路径、URL和元数据")],
        "interfaces":[("上层","TeachingArtifactService.generate_*","同步服务函数"),("外部","DashScope/Seedance/Edge TTS","HTTPS"),("本地","python-pptx/Matplotlib/Manim/PyAV","Python库/子进程"),("数据","uploads/generated_artifacts/images","文件")],
        "storage":[("generated_artifacts","PPTX、MP4和其他产物","长期/用户可下载"),("generated_images","PNG/JPEG","长期/受控访问"),("临时目录","Manim脚本、音频和中间帧","完成后清理")],
        "comments":["严谨图和场景插图必须使用不同路由。","外部返回URL先下载到本地再交付，避免临时链接失效。","媒体输出必须检查magic header和至少一个有效流。"],
        "limits":["Seedance额度不足会回退，画面风格与云端视频不同。","Manim依赖字体和本地渲染环境。","文生图不保证中文文字正确，因此不用于密集标注图。"],
        "tests":[("DFD路由","黑字白底且不要Mermaid","生成PNG而非插画"),("视频优先级","TCP教学视频提示","不进入Mermaid"),("Seedance 402","额度不足","回退Manim"),("音轨校验","回退视频","包含H.264和AAC"),("PPT结构","5..12页计划","可打开且有标题/要点")],
        "unresolved":"增加公式动画模板、视频字幕轨、产物缩略图和不同课程视觉主题。",
        "fallback":"Seedance→Qwen+Manim+TTS\n文生图不适合→确定性绘图\n模型JSON异常→修复/模板",
    },
    {
        "id":"P06", "name":"练习、批改与学科错题本程序", "ident":"QUIZ_WRONG_BOOK_SERVICE",
        "files":"services/quiz_service.py；api/v1/endpoints/quiz.py；models/quiz.py",
        "description":"负责题目生成、质量审查、答题提交、逐题判定、错题同步、按学科归档和错题衍生练习。程序以QuizAttempt为评分事实，以WrongQuestion保存累计错误状态，并向画像与路径提供学习证据。",
        "features":["课程题库和LLM生成双来源","题目结构与学科规则质量校验","自动评分和逐题解析","错题唯一约束、错误次数与掌握状态","按subject/knowledge_point智能分类","基于错因的变式练习","再次提交后更新掌握状态"],
        "performance":"一次测验通常5..20题；数据库按user_id/resource_id/question_id建立索引。错题练习生成使用有限题量，避免大模型长时间输出。",
        "inputs":[("resource_id","UUID","练习资源","课程/错题本"),("answers","question_id→answer","学生答案","Quiz UI"),("wrong question ids","UUID数组","目标错题","错题本"),("subject/tags","文本","学科和知识点","题目/资源"),("favorite/mastered","布尔","错题状态","用户操作")],
        "outputs":[("quiz","不含标准答案的题目数组","做题界面"),("attempt","分数、正确数、错知识点","学习记录"),("question results","对错、正确答案、解析","批改界面"),("wrong book","按学科分组的本子与错题","错题本"),("practice resource","衍生练习题","再次作答")],
        "algorithms":["提交时逐题比较规范化答案，计算correct_count和score=correct_count/total_questions，记录wrong_knowledge_points。标准答案只在提交后输出。","_sync_wrong_questions按(user_id,question_id)唯一约束更新：首次错误创建，重复错误wrong_count+1并更新时间；重新掌握后可更新mastered。","学科分类优先使用课程/资源subject和知识标签；若是“通用”等泛化值，则由Practice Agent结合题干、知识点和课程上下文推断，并限制为可解释学科名称。","错题衍生练习把原题、错误知识点和错因作为约束，要求生成同知识点、不同表述或条件的变式题；生成后执行题干、选项、答案、难度与主题一致性校验。"],
        "steps":[("读取练习并隐藏答案","验证资源和用户"),("学生提交答案","格式与题目集合校验"),("逐题评分与解析","计算分数和错知识点"),("同步错题本","分类、计数与掌握状态"),("写入学习证据","供画像与路径使用"),("生成变式练习","错因约束+质量复核"),("再次作答","形成新证据闭环")],
        "interfaces":[("上层","/resource-hub/quizzes/*","REST"),("上层","/resource-hub/wrong-book/*","REST"),("下层","QuizService/模型客户端","函数调用"),("数据","question/quiz_attempt/wrong_question/resource","SQLModel")],
        "storage":[("question","题干、选项、答案、解析和标签","课程/生成资源"),("quiz_attempt","答案、分数和错知识点","长期证据"),("wrong_question","错误次数、收藏、掌握状态","用户维度")],
        "comments":["答案不得在未提交的公共题目DTO中出现。","数据库范式类题目有专用逻辑校验，避免概念错误。","错题分类必须保留推断依据或原始标签。"],
        "limits":["主观题自动评分需要量表，当前重点为客观题。","标签缺失时学科推断可能不稳定。","题库质量最终仍需教师抽检。"],
        "tests":[("答案隐藏","获取练习","响应无answer"),("重复错题","同题再次答错","wrong_count递增不重复建行"),("学科分类","软件工程题+通用标签","归入软件工程"),("变式练习","选择错题生成","新题同知识点且非原题复制"),("越权提交","他人资源/尝试","403/404")],
        "unresolved":"增加主观题评分量表、知识点覆盖率约束和教师题目审核工作台。",
        "fallback":"模型生成失败→课程精选题库\n单题不合格→局部修复\n仍失败→返回可重试错误",
    },
    {
        "id":"P07", "name":"学生画像与学习效果评估程序", "ident":"STUDENT_PROFILE_AGENT",
        "files":"services/student_profile_agent.py；learning_report_service.py；models/student_profile.py",
        "description":"把对话、练习、资源访问、收藏与记忆画像聚合为版本化AI学习数字分身。定量掌握度只由有分数的QuizAttempt更新；浏览和聊天只影响行为、偏好和置信度，不会虚假提升掌握度。",
        "features":["八维画像与总分","掌握度历史平滑更新","资源偏好指数平滑","学习阶段判定","优势/薄弱点提取","个人知识网络","画像版本、证据游标与最近更新","向规划/资源/评估Agent共享"],
        "performance":"按用户增量处理QuizAttempt；evidence_cursor最多保留最近200个attempt_id，避免重复计算。画像同步为短事务，单用户唯一行。",
        "inputs":[("QuizAttempt","score、错知识点","定量掌握证据","练习程序"),("ConversationMessage","用户轮次","行为证据","聊天程序"),("LearningActivity/Favorite","访问与收藏","偏好证据","资源程序"),("memory profile","目标、风格、反馈","长期记忆","Memory服务"),("User/Student/UD","专业、学校","背景","用户中心")],
        "outputs":[("StudentProfile","八维、优势、薄弱、摘要、版本","学情档案"),("dimension scores","0..100整数","雷达/进度图"),("knowledge graph","nodes/edges","知识网络画像"),("agent links","协同用途说明","规划/资源/评估"),("last updates","最多5条","动态更新区域")],
        "algorithms":["掌握度使用指数历史平滑：new=0.7×previous+0.3×observation，并截断到[0,1]。答题错误知识点以observation=0参与更新；浏览和聊天不进入该公式。","偏好融合使用alpha=0.2：merged=(1-alpha)×current+alpha×observation，observation由显式推荐反馈亲和度映射，最多保留8种偏好。","阶段按证据数判定：evidence≥20为能力提升期，6..19为课程强化阶段，否则为画像形成期。弱项为掌握度<0.55，强项为≥0.72。","综合能力overall=0.38×知识+0.22×行为+0.30×正确率+0.10×反馈。八维分别结合目标是否明确、偏好置信度、行为量、正确率和知识覆盖计算。","只有关键字段发生变化时profile_version递增；已处理attempt_id进入游标，保证同步幂等。"],
        "steps":[("加载旧画像与证据游标","取得已处理attempt"),("聚合新增练习","更新知识掌握与错知识点"),("聚合对话/访问/反馈","更新行为和偏好"),("计算阶段与八维分数","强弱项、综合能力"),("构建知识网络与AI摘要","节点、边和建议"),("比较变更并递增版本","写入更新时间和游标"),("同步给协作Agent","路径、推荐和评估")],
        "interfaces":[("上层","GET /profile/digital-twin","公共画像DTO"),("上层","POST /profile/analyze|update","触发同步"),("下层","UserMemoryProfileService/LearningReport","函数调用"),("数据","student_profile/profile_update_event/learning_evidence","SQLModel")],
        "storage":[("student_profile","单用户物化快照和JSON维度","长期"),("profile_update_event","对话画像信号与证据","审计"),("evidence_cursor","已处理attempt与来源版本","随画像保存")],
        "comments":["掌握度只能由评分证据改变。","所有分数必须clamp并注明量纲。","新增维度需同步DIMENSION_LABELS、前端与测试。"],
        "limits":["冷启动证据不足时画像可信度低。","当前综合权重为工程先验，仍需更多真实学生校准。","对话信号不能替代正式能力测评。"],
        "tests":[("掌握平滑","旧0.5新1.0","结果0.65"),("浏览行为","只有资源访问","掌握度不提升"),("幂等同步","同attempt重复刷新","不重复更新"),("阶段边界","证据5/6/20","阶段正确"),("版本更新","无字段变化","版本不递增")],
        "unresolved":"扩大真实学生样本校准权重，增加画像维度置信区间和学生主动纠偏入口。",
        "fallback":"无新增证据→保留旧画像\n证据不足→显示待积累\n异常值→clamp并记录",
    },
    {
        "id":"P08", "name":"学习路径与个性化资源推荐程序", "ident":"LEARNING_PATH_RECOMMENDER",
        "files":"services/learning_path_service.py；resource_recommendation_service.py；recommendation_ranking_service.py",
        "description":"根据画像薄弱点、练习差距、目标、知识库主题和偏好生成学习路径，并统一排序站内、生成和受控外部资源。排序核心为本地可解释算法，不依赖模型黑盒分数。",
        "features":["弱项优先路径节点","资源运行后自动补充路径","站内/生成/外部候选统一排序","中文字符n-gram BM25","练习与掌握差距加权","模态、难度和来源权威度加权","MMR多样性重排","收藏、忽略和显式反馈闭环"],
        "performance":"候选集合小规模内存排序；BM25使用字符一元/二元词，避免中文分词依赖。外部来源每provider和topic均限制数量，并设置新鲜度和相关性门槛。",
        "inputs":[("profile context","弱点、目标、风格、偏好","画像程序"),("practice/mastery gaps","topic→gap","学习报告"),("candidates","站内/生成/外部资源","资源库"),("feedback","收藏、忽略、评价","用户"),("course graph","知识节点与先修关系","知识图谱")],
        "outputs":[("LearningPath","subject、summary、ordered nodes","学习路径"),("recommendations","资源、分数证据、推荐理由","资料中心"),("external metadata","provider、URL、验证信息","外部资源卡"),("feedback signal","主题/模态亲和度","画像偏好")],
        "algorithms":["学习路径先取最多4个weak_points，首项为in_progress，其余pending；再追加最多3条recommended_actions。资源生成成功时，若topic不存在则追加“学习并核验topic”节点并记录before/after事件。","文本相关性采用字符unigram+bigram BM25，k1=1.2、b=0.75，归一化到[0,1]。基础分=0.08+0.52×lexical。","练习差距命中加0.18..0.30，掌握差距加0.10..0.22，薄弱点加0.16，目标加0.11，课程主题加0.07；模态、反馈、难度、权威来源分别贡献受限加分。","外部资源必须通过国内HTTPS来源策略与标题具体主题门槛。最后使用lambda=0.60的MMR兼顾相关性和内容多样性，推荐理由从最高权重证据生成。"],
        "steps":[("聚合画像与学习报告","弱点、差距、目标和偏好"),("构造学习路径","弱项优先和前置顺序"),("收集候选资源","站内、生成与受控外部来源"),("BM25与证据加权","相关性、差距、模态和权威度"),("来源门禁与去重","HTTPS、国内来源、主题匹配"),("MMR多样性排序","避免同类资源占满"),("交付并收集反馈","收藏/忽略影响后续偏好")],
        "interfaces":[("上层","GET /learning-path/me","路径DTO"),("上层","GET/POST /resource-hub/recommendations/*","推荐与反馈"),("下层","ExternalResourceDiscoveryService","受控目录发现"),("数据","learning_path/update_event/recommendation/external_resource","SQLModel")],
        "storage":[("learning_path","用户有序节点JSON","动态覆盖"),("personalized_resource_recommendation","分数、理由、状态和反馈","长期"),("external_resource","provider、URL、摘要、验证时间","可刷新")],
        "comments":["排序权重修改必须更新解释文本和单测。","外部URL不能仅凭元数据标签信任。","路径节点写入需保留source_run_id。"],
        "limits":["冷启动只能使用课程主题和通用难度。","外部平台页面可能失效或改变。","当前权重是可解释启发式，尚未采用大规模在线学习。"],
        "tests":[("弱项优先","两个weak points","第一节点in_progress"),("中文相关性","事务隔离查询","相关标题高分"),("来源门禁","非受控域名","不公开推荐"),("多样性","多个同模态候选","MMR保留其他形式"),("反馈闭环","收藏视频","视频偏好上升")],
        "unresolved":"增加A/B实验、路径先修约束求解和外部资源周期性可用性巡检。",
        "fallback":"无画像→课程通用路径\n外部发现失败→仅站内资源\n无候选→生成资源建议",
    },
    {
        "id":"P09", "name":"小智悬浮助手与数字人程序", "ident":"XIAOZHI_DIGITAL_HUMAN",
        "files":"components/float-ai/ClassroomQuickChat.vue；digital_human_assistant.py；digital_human_service.py；digital_human_tts.py",
        "description":"把多智能体的提醒、建议和语音能力封装为轻量悬浮胶囊，并提供文本/PPT到数字人视频的任务接口。小智读取画像、当前任务和最近建议，但在练习做题场景按路由策略隐藏，保证测评独立性。",
        "features":["悬浮胶囊与场景可见性","画像驱动主动提醒","最近建议与自动播报","Edge TTS音色重试","文本/PPT数字人任务","任务所有权与签名产物URL","MuseTalk/Wav2Lip运行适配"],
        "performance":"提醒接口返回短文本；语音与数字人渲染为异步/子进程任务。TTS设置超时、重试和备用音色；同一页面避免频繁自动播报。",
        "inputs":[("route","前端路由名","决定是否显示","Vue Router"),("profile/task","画像和当前任务","提醒依据","后端"),("text/ppt","文本或PPT文件","数字人内容","用户"),("voice_id","音色标识","语音选择","配置/用户"),("auto_speak","布尔与冷却时间","主动提醒策略","前端")],
        "outputs":[("capsule state","显示/隐藏/动画状态","页面UI"),("suggestion","短提醒文本","学生"),("audio","MP3/WAV","语音播放"),("video job","task_id/status","数字人生成进度"),("signed URL","限时产物地址","授权下载")],
        "algorithms":["前端根据路由meta和明确的练习/测验路径计算showXiaozhi；做题页无论画像如何均返回false，资料库和AI聊天可显示。","提醒内容由近期建议、当前任务和画像薄弱点排序，并设置时间冷却，避免同一提示重复打断。","TTS按首选音色、配置音色和备用音色顺序尝试，每种音色有限重试；只接受大小超过阈值的有效音频文件。","数字人任务为用户登记owner sidecar；查询和下载时校验任务归属，产物URL使用签名ticket并带过期时间。"],
        "steps":[("页面场景判断","练习页隐藏，学习页可显示"),("读取画像与当前任务","仅取必要提醒信号"),("生成/选择短建议","去重、优先级和冷却"),("展示胶囊或语音","用户可关闭"),("可选数字人任务","TTS、唇形和视频渲染"),("归属校验后交付","签名URL和状态")],
        "interfaces":[("前端","ClassroomQuickChat.vue/default-layout.vue","组件与路由"),("API","/classroom-assistant/speak|auto-speak|recent-suggestions","REST"),("服务","DigitalHumanService/DigitalHumanTTS","函数/子进程"),("存储","digital-human jobs/artifacts ownership","文件")],
        "storage":[("前端状态","展开、冷却和最近提示","会话级"),("语音/视频目录","TTS与渲染产物","任务生命周期"),("owner sidecar","任务与用户映射","长期/随产物")],
        "comments":["练习页隐藏是强规则，不得被模型建议覆盖。","自动播报必须尊重用户静音和浏览器权限。","产物访问必须验证签名与owner。"],
        "limits":["本地数字人依赖模型文件和运行环境。","浏览器自动播放策略可能阻止语音。","主动提醒需防止过度打扰。"],
        "tests":[("场景隐藏","进入quiz/practice","胶囊不存在"),("学习场景","资料中心/聊天","胶囊可用"),("TTS回退","首选音色失败","使用备用音色"),("任务越权","访问他人task","拒绝"),("签名过期","旧ticket","拒绝下载")],
        "unresolved":"增加用户可配置提醒频率、更多动态形象状态和基于学习节奏的打扰成本模型。",
        "fallback":"语音失败→仅文字提醒\n数字人不可用→静态小智\n无画像→通用学习建议",
    },
    {
        "id":"P10", "name":"前端交互、安全与运行保障程序", "ident":"WEB_SECURITY_OPERATIONS",
        "files":"education/course/src；core/http_security.py；endpoints/health.py；core/config.py",
        "description":"由Vue单页客户端与FastAPI中间件/健康检查共同组成。前端负责路由、状态、流式渲染、附件交互和资源卡片；后端负责请求体限制、安全响应头、AI并发预算、可信主机、配置验证和组件就绪检查。",
        "features":["Vue Router与Pinia状态","SSE增量渲染和停止/重试","Markdown/公式/Mermaid/成果卡片","附件拖拽和粘贴","请求大小与AI并发中间件","JWT/CORS/TrustedHost安全边界","healthz/readyz与配置告警","Alembic迁移和数据库漂移修复"],
        "performance":"Vite按路由分包；流式回答只追加增量内容。最大请求体约26MB，WebSocket消息约3MB；AI并发预算按用户身份清理和限制，防止资源耗尽。",
        "inputs":[("HTTP/SSE/WS","请求、事件和消息","浏览器/后端"),("route meta","权限、布局和场景","Vue Router"),("JWT","Bearer令牌","身份认证"),("settings","环境变量","数据库、模型和安全配置"),("health probes","GET请求","部署平台")],
        "outputs":[("页面状态","加载、空态、成功、失败","用户界面"),("HTTP headers","CSP类安全头/缓存策略","浏览器"),("429/413/4xx","受控安全错误","客户端"),("health report","组件存活/就绪状态","运维"),("logs","请求、模型、任务和异常摘要","开发运维")],
        "algorithms":["ChatLayout根据SSE事件类型维护phase、answer、artifact和error；停止操作关闭流并保留已生成内容，重试复用原请求但创建新执行。","RequestSizeLimitMiddleware在读取完整请求前检查Content-Length，并在流式接收过程中累计字节，超过上限返回413。","AIRequestBudgetMiddleware按用户/身份和请求类型分为SSE或同步预算，获取令牌后用finally释放；定期清理长期不活跃身份。","readyz检查数据库、关键目录和必要配置；Alembic迁移是部署前置步骤。对于历史数据库列漂移，幂等修复迁移检查物理列后补齐。"],
        "steps":[("浏览器路由与鉴权","加载用户和页面组件"),("提交请求","Axios/SSE/WS统一API前缀"),("安全中间件","主机、大小、并发与令牌"),("业务处理与增量事件","后端服务程序"),("前端状态归并","Markdown/卡片/进度"),("健康、日志与告警","运维观察和错误恢复")],
        "interfaces":[("前端","src/api/*.ts → /api/v1","Axios/SSE"),("中间件","ASGI scope/receive/send","Starlette"),("认证","login/JWT/current_user","FastAPI依赖"),("运维","/healthz /readyz","HTTP")],
        "storage":[("Pinia/内存","页面和会话UI状态","短期"),("浏览器受控存储","令牌/偏好","最小化"),("后端日志","错误与运行摘要","按轮转策略"),("Alembic版本表","数据库版本","长期")],
        "comments":["前端不得包含模型API密钥。","所有文件链接必须使用后端受控URL。","中间件异常必须返回结构化响应而非原始堆栈。"],
        "limits":["开发模式Vite代理依赖8001后端。","单机内存并发预算不跨多实例共享。","前端错误边界仍需覆盖所有异步组件。"],
        "tests":[("生产构建","npm run build","成功且无边界违规"),("请求过大","超过上限","413"),("AI并发","超过预算","429且令牌可释放"),("健康检查","数据库异常","readyz非就绪"),("数据库迁移","旧schema","upgrade后ORM查询成功")],
        "unresolved":"生产部署增加反向代理、TLS、集中日志、Redis共享限流和浏览器端端到端测试。",
        "fallback":"SSE断开→保留已生成内容\n后端未启动→统一连接提示\n非核心模块失败→页面局部降级",
    },
]


def set_doc_styles(doc: Document) -> None:
    sec=doc.sections[0]
    sec.top_margin=Cm(2.3); sec.bottom_margin=Cm(2.1); sec.left_margin=Cm(2.45); sec.right_margin=Cm(2.25)
    sec.header_distance=Cm(.9); sec.footer_distance=Cm(.9)
    normal=doc.styles["Normal"]; normal.font.name="Microsoft YaHei"; normal.font.size=Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    for name,size,color in [("Title",28,DEEP_BLUE),("Heading 1",17,DEEP_BLUE),("Heading 2",13.5,BLUE),("Heading 3",11,"344054")]:
        st=doc.styles[name]; st.font.name="Microsoft YaHei"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    doc.styles["Heading 1"].paragraph_format.space_before=Pt(16)
    doc.styles["Heading 2"].paragraph_format.space_before=Pt(10)
    header=sec.header.paragraphs[0]; header.text="智屿——基于大模型的个性化资源生成与学习多智能体系统　详细设计说明书"; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in header.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    add_page_number(sec.footer.paragraphs[0])


def module_chapter(doc: Document, chapter: int, m: dict, fig: Path) -> None:
    heading(doc, f"{chapter} {m['name']}（{m['id']}）设计说明", 1)
    heading(doc, f"{chapter}.1 程序描述", 2)
    body(doc,m["description"])
    add_table(doc,["属性","设计"],[
        ("程序标识符",m["ident"]),("程序编号",m["id"]),("主要代码",m["files"]),
        ("运行方式","服务对象/请求处理程序；可重入；按模块支持同步、异步或流式"),("所属系统","智屿Web应用"),
    ],font_size=8)
    heading(doc,f"{chapter}.2 功能",2)
    for f in m["features"]: bullet(doc,f)
    add_table(doc,["输入","处理","输出"],[("用户请求、课程/画像上下文及程序专用数据","校验→核心算法→持久化/审查→状态更新","DTO、SSE事件、学习证据或文件产物")],font_size=8)
    heading(doc,f"{chapter}.3 性能",2); body(doc,m["performance"])
    heading(doc,f"{chapter}.4 输入项",2); add_table(doc,["名称","类型/范围","含义","来源"],m["inputs"],font_size=7)
    heading(doc,f"{chapter}.5 输出项",2); add_table(doc,["名称","格式","用途"],m["outputs"],font_size=7)
    heading(doc,f"{chapter}.6 算法",2)
    for i,a in enumerate(m["algorithms"],1): numbered(doc,f"{a}")
    heading(doc,f"{chapter}.7 流程逻辑",2); add_picture(doc,fig,f"图{chapter}-1 {m['name']}流程逻辑",width=5.8)
    heading(doc,f"{chapter}.8 接口",2); add_table(doc,["方向","接口/对象","调用方式"],m["interfaces"],font_size=7)
    heading(doc,f"{chapter}.9 存储分配",2); add_table(doc,["存储对象","内容","生命周期"],m["storage"],font_size=7)
    heading(doc,f"{chapter}.10 注释设计",2)
    for c in m["comments"]: bullet(doc,c)
    body(doc,"模块首部应说明职责、并发/事务特性和安全边界；复杂分支注明业务原因而非复述代码；配置变量注明默认值、范围和是否敏感；异常分支说明是否可重试、是否回退以及对已持久化状态的影响。")
    heading(doc,f"{chapter}.11 限制条件",2)
    for x in m["limits"]: bullet(doc,x)
    heading(doc,f"{chapter}.12 测试计划",2); add_table(doc,["测试项","输入/条件","预期结果"],m["tests"],font_size=7)
    body(doc,"单元测试由开发人员编写并在提交前运行；API权限和状态测试使用隔离测试数据库；媒体能力在具备依赖的环境中执行集成测试；涉及外部付费API时使用受控测试账号或桩响应，同时保留至少一次真实端到端验证。")
    heading(doc,f"{chapter}.13 尚未解决的问题",2); body(doc,m["unresolved"])


def build() -> None:
    OUT.mkdir(exist_ok=True); ASSETS.mkdir(exist_ok=True); configure_matplotlib()
    common={"structure":ASSETS/"00_program_structure.png","architecture":ASSETS/"00_architecture.png","deployment":ASSETS/"00_deployment.png","data":ASSETS/"00_data_model.png"}
    save_program_structure(common["structure"]); save_architecture(common["architecture"]); save_deployment(common["deployment"]); save_data_model(common["data"])
    flow_paths=[]
    for i,m in enumerate(MODULES,1):
        p=ASSETS/f"{i:02d}_{m['id'].lower()}_flow.png"; save_flow(p,m["name"]+"流程",m["steps"],m["fallback"]); flow_paths.append(p)

    doc=Document(); set_doc_styles(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(60)
    r=p.add_run("第十五届“中国软件杯”大学生软件设计大赛\nA3 赛题参赛项目"); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(BLUE); r.bold=True
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(35)
    r=p.add_run("智　屿"); r.font.size=Pt(38); r.bold=True; r.font.color.rgb=RGBColor.from_string(DEEP_BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("基于大模型的个性化资源生成与学习多智能体系统").font.size=Pt(17)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(38)
    r=p.add_run("详细设计说明书"); r.font.size=Pt(30); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(75)
    p.add_run("文档版本：V1.0\n编制单位：智屿项目团队\n完成日期：2026 年 7 月 20 日").font.size=Pt(12)
    page_break(doc)
    heading(doc,"文档控制",1)
    add_table(doc,["项目","内容"],[
        ("文档名称","智屿详细设计说明书"),("赛题","A3—基于大模型的个性化资源生成与学习多智能体系统开发"),
        ("版本/状态","V1.0 / 提交评审版"),("编制/审核","智屿项目团队 / 待团队负责人签署"),
        ("配套文档","智屿概要设计说明书V1.0、数据库设计说明书、需求说明书、测试说明书"),
        ("安全说明","文档不包含真实API密钥、账号口令和可识别学生隐私"),
    ],font_size=9)
    heading(doc,"修改情况记录",2)
    add_table(doc,["版本","日期","修改人","说明","批准人"],[
        ("V0.1","2026-07-19","智屿项目团队","按代码结构形成模块清单","—"),
        ("V1.0","2026-07-20","智屿项目团队","完成10个核心程序详细设计与赛题追踪","待签署"),
    ],font_size=8)
    heading(doc,"目录",1); add_toc(doc.add_paragraph()); body(doc,"在WPS/Word中打开后更新目录域即可显示最终页码。")
    page_break(doc)

    heading(doc,"1 引言",1)
    heading(doc,"1.1 编写目的",2)
    body(doc,"本文档在《智屿概要设计说明书》基础上，进一步规定各核心程序的职责、输入输出、算法、流程、接口、存储、限制和单元测试计划，为编码复核、系统集成、缺陷定位、比赛验收和后续扩展提供可执行依据。预期读者为开发、测试、运维成员、指导教师和竞赛评审专家。")
    heading(doc,"1.2 背景",2)
    add_table(doc,["背景项","内容"],[
        ("软件名称","智屿——基于大模型的个性化资源生成与学习多智能体系统"),("任务提出者","第十五届中国软件杯A3赛题，出题企业为科大讯飞股份有限公司"),
        ("开发者","智屿项目团队"),("用户","高校学生、教师、课程建设者和系统管理员"),
        ("运行单位","参赛演示环境及具备常规Web运行条件的高校教学环境"),("课程样板","软件工程导论完整知识库与文档集"),
    ],font_size=8)
    body(doc,"赛题要求不少于6维动态画像、多智能体协作生成至少5类个性化资源、个性化路径和精准推送，并鼓励多模态辅导与学习效果评估。智屿以13类Agent、八维数字分身、完整课程RAG、资源工作流和错题闭环实现上述目标。")
    heading(doc,"1.3 定义",2)
    add_table(doc,["术语","定义"],[
        ("程序","本说明书中的可独立描述、测试和维护的服务模块、前端模块或后台执行单元。"),("Agent","具有角色、上下文、工具权限和状态的智能体节点。"),
        ("Supervisor","多智能体协作主管，只负责拆解、路由和汇总。"),("RAG","检索增强生成，以课程或文档证据约束模型回答。"),
        ("SSE","服务器发送事件，用于流式阶段、文本和产物通知。"),("Run/Step","一次资源生成运行及其可追踪子步骤。"),
        ("学习证据","支持掌握度、偏好或风险判断的可追溯记录。"),("能力路由","依据资源准确性和模态选择模型或确定性工具链。"),
        ("幂等","重复提交相同业务请求不会产生重复副作用。"),("租约","异步执行尝试在限定时间内推进任务状态的所有权。"),
    ],font_size=8)
    heading(doc,"1.4 参考资料",2)
    refs=[
        "[1] 第十五届中国软件杯A3赛题《基于大模型的个性化资源生成与学习多智能体系统开发》，2026-04-01，https://www.cnsoftbei.com/content-3-1286-1.html。",
        "[2] 用户提供的《详细设计说明书编写规范》。",
        "[3] 《智屿概要设计说明书V1.0》。",
        "[4] ZhiYu-main项目源代码、数据库迁移、测试用例、requirements.txt与package.json，核对日期2026-07-20。",
        "[5] 软件工程导论原始资料集与课程入库说明，共127个文件、约497MB。",
    ]
    for x in refs: bullet(doc,x)

    heading(doc,"2 程序系统的结构",1)
    heading(doc,"2.1 总体层次结构",2)
    body(doc,"智屿采用前端交互层、API与业务服务层、多智能体编排层、能力适配层和数据层。详细设计按10个核心程序划分；程序之间通过显式Service调用、REST/SSE接口和数据库事件对象协作。")
    add_picture(doc,common["structure"],"图2-1 程序系统层次结构",width=6.5)
    add_picture(doc,common["architecture"],"图2-2 系统总体技术架构",width=6.5)
    heading(doc,"2.2 程序清单",2)
    add_table(doc,["编号","标识符","程序名称","主要职责"],[(m["id"],m["ident"],m["name"],"；".join(m["features"][:3])) for m in MODULES],font_size=7)
    heading(doc,"2.3 公共调用与数据约定",2)
    for x in [
        "用户边界：所有用户对象读写必须以current_user为边界，不能仅信任客户端传入user_id。",
        "事务边界：数据库短事务与外部模型调用分离；先写任务状态，外部调用完成后再提交结果。",
        "错误约定：REST使用HTTP状态码和detail；SSE使用error事件；内部日志保留错误摘要但屏蔽密钥。",
        "时间与标识：业务主键优先UUID或服务端生成字符串；时间使用带时区UTC，前端按本地时区显示。",
        "文件约定：文件名由服务端生成，路径resolve后必须位于允许根目录；下载经鉴权端点。",
        "证据约定：课程结论优先携带citation；学习画像定量值只由适当证据更新。",
    ]: bullet(doc,x)
    heading(doc,"2.4 物理部署关系",2); add_picture(doc,common["deployment"],"图2-3 物理部署与依赖",width=6.4)
    heading(doc,"2.5 核心数据关系",2); add_picture(doc,common["data"],"图2-4 核心数据对象关系",width=6.4)

    for idx,(m,fig) in enumerate(zip(MODULES,flow_paths,strict=True),start=3):
        module_chapter(doc,idx,m,fig)

    heading(doc,"附录A 核心算法参数表",1)
    add_table(doc,["算法/控制","参数","当前值","设计含义"],[
        ("掌握度平滑","history weight","0.70","历史占70%，新评分占30%"),("偏好融合","alpha","0.20","避免一次反馈剧烈改变偏好"),
        ("综合能力","knowledge/behavior/accuracy/feedback","0.38/0.22/0.30/0.10","兼顾知识、投入、表现和反馈"),
        ("推荐BM25","k1/b","1.2/0.75","适配小候选集文本相关性"),("推荐MMR","lambda","0.60","相关性优先并保留多样性"),
        ("外部相关门槛","relevance/title","0.16/0.08","阻止标题弱相关资源"),("AI并发","SSE/sync每用户","2/1","避免单用户耗尽模型配额"),
        ("上传大小","通用/图片/视频","25MB/20MB/150MB","控制内存与外部请求体"),("画像阶段","证据数","<6 / 6..19 / ≥20","形成期/强化阶段/能力提升期"),
    ],font_size=7)
    heading(doc,"附录B 赛题功能—程序追踪矩阵",1)
    add_table(doc,["赛题要求","主程序","协同程序","实现要点"],[
        ("不少于6维动态画像","P07","P01/P06/P08","八维画像、证据游标、版本更新"),("多智能体协同","P01","P02..P09","13类Agent、主管编排和工具边界"),
        ("至少5类资源","P04/P05","P01/P02/P07","文档、PPT、题目、图、视频、代码"),("个性化路径与推送","P08","P06/P07","差距、目标、偏好和MMR排序"),
        ("智能辅导","P01/P03","P02/P05","图片+文本+课程证据+图解视频"),("学习效果评估","P06/P07","P08","评分证据、错题、八维趋势和路径重排"),
        ("完整专业课程知识库","P02","P01/P06/P08","软件工程导论127个原始文件"),("防幻觉与内容安全","P01/P02/P05","P10","证据优先、结构门禁、安全审查、回退"),
    ],font_size=7)
    heading(doc,"附录C 代码与测试证据索引",1)
    add_table(doc,["类别","位置","说明"],[
        ("后端入口","code/backend/app/main.py、app/api/main.py","FastAPI与路由装配"),("多智能体","app/ai/chat_engine.py、chat_runtime.py","状态图、13类角色和流式输出"),
        ("课程知识库","app/services/rag_service.py、backend/data/course_kb","入库、检索与课程资料"),("资源生成","resource_generation_service.py、teaching_artifact_service.py","运行、质量门禁和媒体产物"),
        ("学习闭环","quiz_service.py、student_profile_agent.py、learning_path_service.py","错题、画像与路径"),("前端","code/education/course/src","聊天、课程、资料、错题和学情页面"),
        ("数据库","app/models、app/alembic/versions","SQLModel与版本迁移"),("测试","app/tests/services、app/tests/api/routes","路由、安全、证据、媒体和画像测试"),
    ],font_size=8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
    r=p.add_run("—— 文档结束 ——"); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    core=doc.core_properties; core.title="智屿详细设计说明书"; core.subject="第十五届中国软件杯A3赛题"; core.author="智屿项目团队"; core.keywords="详细设计,多智能体,RAG,学生画像,资源生成"
    doc.save(DOCX); print(DOCX)


if __name__ == "__main__": build()
