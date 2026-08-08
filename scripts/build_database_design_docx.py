from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_srs_docx import (
    add_properties as add_generic_properties,
    configure_page,
    configure_styles,
    render_body,
    set_cell_shading,
    set_cell_text,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "智屿智能教育平台_数据库设计说明书.md"
ASSET_DIR = ROOT / "docs" / "assets" / "database-design"
OUTPUT_DIR = ROOT / "output"
OUTPUT = OUTPUT_DIR / "智屿智能教育平台_数据库设计说明书.docx"

COLORS = {
    "indigo": "#4F46E5",
    "blue": "#2563EB",
    "cyan": "#0891B2",
    "green": "#059669",
    "amber": "#D97706",
    "rose": "#E11D48",
    "slate": "#334155",
    "light": "#F8FAFC",
}


def setup_plot() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Noto Sans CJK SC",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def box(ax, xy, width, height, title, body="", color=COLORS["indigo"], fontsize=12):
    x, y = xy
    patch = FancyBboxPatch(
        (x, y), width, height,
        boxstyle="round,pad=0.012,rounding_size=0.025",
        linewidth=1.7, edgecolor=color, facecolor="#FFFFFF",
    )
    ax.add_patch(patch)
    ax.text(x + width / 2, y + height * 0.64, title, ha="center", va="center",
            fontsize=fontsize, fontweight="bold", color=color)
    if body:
        ax.text(x + width / 2, y + height * 0.30, body, ha="center", va="center",
                fontsize=fontsize - 2, color=COLORS["slate"], linespacing=1.35)
    return patch


def arrow(ax, start, end, color="#94A3B8", text=None, rad=0.0):
    patch = FancyArrowPatch(
        start, end, arrowstyle="-|>", mutation_scale=13,
        linewidth=1.4, color=color,
        connectionstyle=f"arc3,rad={rad}",
    )
    ax.add_patch(patch)
    if text:
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        ax.text(mx, my + 0.025, text, ha="center", va="bottom", fontsize=8, color=COLORS["slate"])


def save(fig, name):
    fig.savefig(ASSET_DIR / name, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def draw_storage_architecture():
    fig, ax = plt.subplots(figsize=(12, 6.2))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(0.5, 0.96, "智屿混合数据存储架构", ha="center", va="center",
            fontsize=18, fontweight="bold", color=COLORS["slate"])
    box(ax, (0.35, 0.78), 0.30, 0.12, "Vue 3 学生端 / 教师端", "课程 · 对话 · 练习 · 画像 · 资源", COLORS["indigo"])
    box(ax, (0.35, 0.56), 0.30, 0.12, "FastAPI 业务与数据访问层", "鉴权 · 事务 · 多智能体编排 · 文件处理", COLORS["blue"])
    arrow(ax, (0.50, 0.78), (0.50, 0.68), text="HTTPS / SSE")
    box(ax, (0.05, 0.24), 0.26, 0.18, "PostgreSQL", "权威业务数据\n关系、约束、事务、审计", COLORS["blue"])
    box(ax, (0.37, 0.24), 0.26, 0.18, "Chroma 向量库", "课程/附件分片与向量\n语义检索，不作事实主库", COLORS["green"])
    box(ax, (0.69, 0.24), 0.26, 0.18, "文件存储", "PPT · PDF · 图片 · 视频\n预览文件与 AI 生成物", COLORS["amber"])
    for x in (0.18, 0.50, 0.82):
        arrow(ax, (0.50, 0.56), (x, 0.42), color=COLORS["cyan"])
    ax.text(0.5, 0.10, "统一资源标识与用户权限把三类存储关联为可追踪的数据整体",
            ha="center", fontsize=11, color=COLORS["slate"])
    save(fig, "01_storage_architecture.png")


def draw_learning_flow():
    fig, ax = plt.subplots(figsize=(13, 4.8))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(0.5, 0.92, "个性化学习闭环数据流", ha="center", fontsize=18,
            fontweight="bold", color=COLORS["slate"])
    items = [
        ("学习行为", "对话 / 阅读 / 作答", COLORS["cyan"]),
        ("原子证据", "learning_evidence", COLORS["blue"]),
        ("画像更新", "student_profile\nprofile_update_event", COLORS["indigo"]),
        ("规划与推荐", "learning_path\nrecommendation", COLORS["green"]),
        ("学习行动", "资源 / 讲解 / 变式练习", COLORS["amber"]),
    ]
    xs = [0.03, 0.225, 0.42, 0.615, 0.81]
    for x, (title, body, color) in zip(xs, items):
        box(ax, (x, 0.42), 0.16, 0.25, title, body, color, fontsize=11)
    for i in range(4):
        arrow(ax, (xs[i] + 0.16, 0.545), (xs[i + 1], 0.545), color="#64748B")
    arrow(ax, (0.89, 0.40), (0.11, 0.40), color=COLORS["rose"], text="新结果再次形成学习证据", rad=0.28)
    ax.text(0.5, 0.18, "原则：证据不可伪造 · 更新可解释 · 请求可幂等 · 数据按用户隔离",
            ha="center", fontsize=11, color=COLORS["slate"])
    save(fig, "02_learning_data_flow.png")


def draw_er_diagram():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(0.5, 0.97, "智屿核心实体关系图（简化）", ha="center", fontsize=18,
            fontweight="bold", color=COLORS["slate"])
    nodes = {
        "user": (0.04, 0.76, "user", "用户身份", COLORS["indigo"]),
        "course": (0.27, 0.76, "course", "高校课程", COLORS["blue"]),
        "resource": (0.50, 0.76, "resource", "课件/笔记/试题", COLORS["cyan"]),
        "question": (0.73, 0.76, "question", "题目", COLORS["amber"]),
        "attempt": (0.73, 0.52, "quiz_attempt", "一次作答", COLORS["amber"]),
        "wrong": (0.73, 0.28, "wrong_question", "用户错题", COLORS["rose"]),
        "evidence": (0.50, 0.28, "learning_evidence", "原子学习证据", COLORS["blue"]),
        "profile": (0.27, 0.28, "student_profile", "当前画像快照", COLORS["indigo"]),
        "path": (0.04, 0.28, "learning_path", "个性化路径", COLORS["green"]),
        "run": (0.04, 0.52, "resource_generation_run", "多智能体生成运行", COLORS["green"]),
        "step": (0.27, 0.52, "resource_generation_step", "规划/生成/质检步骤", COLORS["green"]),
        "knowledge": (0.50, 0.52, "course_knowledge_node", "课程知识节点", COLORS["cyan"]),
    }
    w, h = 0.18, 0.12
    for _, (x, y, title, body, color) in nodes.items():
        box(ax, (x, y), w, h, title, body, color, fontsize=9)
    def center_right(key):
        x, y, *_ = nodes[key]; return x + w, y + h / 2
    def center_left(key):
        x, y, *_ = nodes[key]; return x, y + h / 2
    def center_top(key):
        x, y, *_ = nodes[key]; return x + w / 2, y + h
    def center_bottom(key):
        x, y, *_ = nodes[key]; return x + w / 2, y
    arrow(ax, center_right("user"), center_left("course"), text="学习/授权")
    arrow(ax, center_right("course"), center_left("resource"), text="1:N")
    arrow(ax, center_right("resource"), center_left("question"), text="1:N")
    arrow(ax, center_bottom("question"), center_top("attempt"), text="提交")
    arrow(ax, center_bottom("attempt"), center_top("wrong"), text="答错")
    arrow(ax, center_left("wrong"), center_right("evidence"), text="复练证据")
    arrow(ax, center_left("evidence"), center_right("profile"), text="聚合更新")
    arrow(ax, center_left("profile"), center_right("path"), text="驱动规划")
    arrow(ax, center_top("run"), center_bottom("user"), text="属于")
    arrow(ax, center_right("run"), center_left("step"), text="1:N")
    arrow(ax, center_right("step"), center_left("knowledge"), text="产出/关联")
    arrow(ax, center_top("knowledge"), center_bottom("resource"), text="知识挂接")
    arrow(ax, center_bottom("knowledge"), center_top("evidence"), text="定位知识点")
    ax.text(0.5, 0.10, "注：图为比赛核心链路简化视图；完整数据库共 60 张表（含迁移表）",
            ha="center", fontsize=10, color=COLORS["slate"])
    save(fig, "03_core_er_diagram.png")


def generate_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    setup_plot()
    draw_storage_architecture()
    draw_learning_flow()
    draw_er_diagram()


def configure_db_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "智屿智能教育平台数据库设计说明书"
    props.subject = "中国软件杯 A3 赛题项目数据库设计"
    props.author = "智屿项目组"
    props.keywords = "智屿, 中国软件杯, 数据库设计, PostgreSQL, 学生画像, 多智能体"
    props.comments = "依据项目真实模型与 Alembic 迁移编制"


def add_cover(doc: Document, lines: list[str]) -> int:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(58)
    run = title.add_run("智 屿")
    run.bold = True; run.font.name = "黑体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(34); run.font.color.rgb = RGBColor(79, 70, 229)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("智能教育平台"); run.bold = True; run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体"); run.font.size = Pt(26)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    run = p.add_run("数据库设计说明书"); run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体"); run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(37, 53, 91)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("A3—基于大模型的个性化资源生成与学习多智能体系统开发")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("文档编号：ZHIYU-DBDD-001    版本：V1.0")

    table = doc.add_table(rows=4, cols=2); table.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("参赛团队", "____________________"),
        ("项目负责人", "____________________"),
        ("指导教师", "____________________"),
        ("编制日期", "2026 年 7 月 20 日"),
    ]):
        set_cell_text(table.cell(i, 0), label, header=True)
        set_cell_text(table.cell(i, 1), value)
        set_cell_shading(table.cell(i, 0), "EEF2FF")
    doc.add_page_break()
    for idx, line in enumerate(lines):
        if line.strip() == "## 文档修订记录":
            return idx
    return 0


def main() -> None:
    generate_diagrams()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    start = add_cover(doc, lines)
    render_body(doc, lines, start, source_path=SOURCE)
    configure_db_properties(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
